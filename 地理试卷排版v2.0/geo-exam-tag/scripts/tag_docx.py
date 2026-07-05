# -*- coding: utf-8 -*-
"""
地理试卷打标脚本
对清洗后的地理高考真题 Word 文档进行结构化打标。

用法:
    python tag_docx.py --input cleaned.docx --output tagged.json

输出:
    tagged.json       结构化打标结果
    images/           提取的图片文件
    tag_log.txt       打标日志
"""

import argparse
import json
import os
import re
import sys
from copy import deepcopy

# 图片占位符正则（用于解析手动插入的占位符）
CHINESE_IMAGE_PLACEHOLDER = re.compile(r'【图片[：:]\s*(\S+?)(?:\s*[-—]\s*.+?)?】')

from docx import Document
from lxml import etree

# 确保能导入同目录下的 utils
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from utils import (
    qn, setup_logger, get_paragraph_text, has_drawing, get_drawings,
    get_embed_rid, get_media_info, get_image_relationships, read_media_file,
    table_to_dict, extract_exam_info,
    is_section_header, is_choice_section, is_non_choice_section, is_fill_section,
    is_question_stem, is_option, is_sub_option, is_sub_question,
    is_instruction, is_exam_info, is_non_choice_title, is_question_text,
    extract_instruction, contains_sub_option, split_sub_options,
    get_question_number, parse_options,
    SECTION_PATTERN, QUESTION_NUMBER_PATTERN, OPTION_PATTERN,
    SUB_QUESTION_PATTERN, SUB_OPTION_PATTERN,
)


def tag_docx(input_path, output_path, log_path=None, images_dir=None):
    """执行试卷打标。

    Args:
        input_path: 输入 docx 文件路径（清洗后）
        output_path: 输出 JSON 文件路径
        log_path: 日志文件路径（可选）
        images_dir: 图片提取目录（可选）
    """
    input_path = os.path.abspath(input_path)
    output_path = os.path.abspath(output_path)

    if not os.path.exists(input_path):
        print(f'错误: 输入文件不存在: {input_path}')
        return False

    if log_path is None:
        log_path = os.path.join(os.path.dirname(output_path), 'tag_log.txt')

    if images_dir is None:
        images_dir = os.path.join(os.path.dirname(output_path), 'images')

    os.makedirs(images_dir, exist_ok=True)

    logger = setup_logger(log_path)

    logger.info('=' * 60)
    logger.info('地理试卷打标开始')
    logger.info(f'输入文件: {input_path}')
    logger.info(f'输出文件: {output_path}')
    logger.info(f'图片目录: {images_dir}')
    logger.info('=' * 60)

    try:
        doc = Document(input_path)
        logger.info(f'文档加载成功，共 {len(doc.paragraphs)} 个段落，{len(doc.tables)} 个表格')

        # 获取媒体文件信息和关系映射
        media_info = get_media_info(input_path)
        rels = get_image_relationships(input_path)

        # 构建段落-表格的有序序列
        # python-docx 的 doc.paragraphs 和 doc.tables 是分开的
        # 需要通过 XML 遍历来保持原始顺序
        elements = _get_ordered_elements(doc)
        logger.info(f'有序元素: {len(elements)} 个（段落+表格）')

        # ========== 全局不确定段落收集器 ==========
        uncertain_paragraphs = []

        # ========== 第一步：提取考试信息 ==========
        exam_info = _extract_exam_info(elements, logger)
        logger.info(f'考试信息: {exam_info}')

        # ========== 第一步补充：提取考前内容（注意事项等） ==========
        pre_exam_content = _extract_pre_exam_content(elements, logger)

        # ========== 第二步：识别分区边界 ==========
        sections = _identify_sections(elements, logger, uncertain_paragraphs)
        logger.info(f'识别到 {len(sections)} 个分区')
        logger.info(f'不确定段落: {len(uncertain_paragraphs)} 个')

        # ========== 第三步：解析每个分区的题组和题目 ==========
        image_counter = [0]  # 图片计数器（用列表以便在函数内修改）
        parsed_sections = []

        for sec_idx, section in enumerate(sections, 1):
            parsed_section = _parse_section(
                section, sec_idx, doc, input_path, media_info, rels, images_dir,
                image_counter, logger, uncertain_paragraphs
            )
            parsed_sections.append(parsed_section)

        # ========== 生成文档全文Markdown（供AI审查） ==========
        document_markdown = _generate_document_markdown(elements)
        logger.info(f'文档Markdown已生成: {len(document_markdown)} 字符')

        # ========== 组装最终 JSON ==========
        result = {
            'exam_info': exam_info,
            'pre_exam_content': pre_exam_content,
            'sections': parsed_sections,
            'uncertain_paragraphs': uncertain_paragraphs,
            'document_markdown': document_markdown,
        }

        # 统计信息
        total_questions = sum(
            len(q)
            for s in parsed_sections
            for g in s.get('question_groups', [])
            for q in [g.get('questions', [])]
        )
        logger.info(f'总题数: {total_questions}')
        logger.info(f'提取图片: {image_counter[0]} 张')

        # 保存 JSON
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

        logger.info(f'JSON 已保存: {output_path}')

        logger.info('')
        logger.info('=' * 60)
        logger.info('打标完成！')
        logger.info('=' * 60)

        return True

    except Exception as e:
        logger.error(f'打标过程中出错: {e}', exc_info=True)
        print(f'错误: {e}')
        return False


def _get_ordered_elements(doc):
    """获取文档中段落和表格的有序列表。
    
    返回列表，每个元素为:
    {'type': 'paragraph', 'element': p_element, 'text': text, 'has_image': bool, 'para_obj': paragraph}
    或
    {'type': 'table', 'table_obj': table, 'table_index': index}
    """
    elements = []
    body = doc.element.body

    # 创建段落和表格的映射
    para_map = {}  # XML元素 -> python-docx Paragraph对象
    for p in doc.paragraphs:
        para_map[p._element] = p

    table_map = {}
    for i, t in enumerate(doc.tables):
        table_map[t._element] = (i, t)

    # 遍历 body 的直接子元素
    for child in body:
        if child.tag == qn('w:p'):
            p_obj = para_map.get(child)
            if p_obj is not None:
                text = get_paragraph_text(child).strip()
                has_img = has_drawing(child)
                elements.append({
                    'type': 'paragraph',
                    'element': child,
                    'text': text,
                    'has_image': has_img,
                    'para_obj': p_obj,
                })
        elif child.tag == qn('w:tbl'):
            info = table_map.get(child)
            if info:
                elements.append({
                    'type': 'table',
                    'table_obj': info[1],
                    'table_index': info[0],
                })

    return elements


def _extract_exam_info(elements, logger):
    """从文档前几个段落提取考试信息。"""
    for elem in elements[:5]:
        if elem['type'] != 'paragraph':
            continue
        text = elem['text']
        if is_exam_info(text) and '地理' in text.replace(' ', ''):
            exam_name, subject = extract_exam_info(text)
            return {'exam_name': exam_name, 'subject': subject}

    # 如果没找到，尝试第一个非空段落
    for elem in elements[:5]:
        if elem['type'] == 'paragraph' and elem['text']:
            exam_name, subject = extract_exam_info(elem['text'])
            return {'exam_name': exam_name, 'subject': subject}

    logger.warning('未找到考试信息，使用默认值')
    return {'exam_name': '', 'subject': '地理'}


# 考前内容关键词
_NOTICE_KEYWORDS = ['注意事项', '考生注意事项', '答题注意事项']
_STUDENT_INFO_KEYWORDS = ['姓名', '班级', '准考证', '考号', '座位号', '考场号', '考场', '装订']


def _is_pre_exam_content(text):
    """判断段落是否为考前内容（注意事项、考生信息栏等）。
    
    识别规则：
    - 含"注意事项"关键词
    - 含考生信息关键词（姓名、班级、准考证号等）且非题干
    - 注意事项后的编号列表项（如"1.作答前..."）
    """
    if not text:
        return False
    # 含注意事项关键词
    if any(kw in text for kw in _NOTICE_KEYWORDS):
        return True
    # 含考生信息关键词（但排除题干，题干以数字+点开头且较长）
    if any(kw in text for kw in _STUDENT_INFO_KEYWORDS):
        # 排除题干（题干通常较长且含问句）
        if len(text) < 80:
            return True
    return False


def _extract_pre_exam_content(elements, logger):
    """提取分区标题之前的考前内容（注意事项、考生信息栏等）。
    
    返回考前内容列表，每项为 {'type': 'notice'|'student_info', 'text': '...'}
    跳过考试名称和科目名称（已在exam_info中处理）。
    """
    pre_content = []
    in_notice = False  # 是否处于注意事项编号列表中

    for elem in elements:
        if elem['type'] != 'paragraph':
            continue
        text = elem['text'].strip()
        if not text:
            continue

        # 遇到分区标题，结束提取
        if is_section_header(text):
            break

        # 跳过考试信息（已在exam_info中处理）
        if is_exam_info(text) and ('地理' in text.replace(' ', '') or '考试' in text):
            continue
        # 跳过纯科目名称（如"地理试题"、"地  理"）
        cleaned = text.replace(' ', '').replace('　', '')
        if cleaned in ('地理', '地理试题', '地理试卷'):
            continue

        # 注意事项标题
        if any(kw in text for kw in _NOTICE_KEYWORDS):
            pre_content.append({'type': 'notice_title', 'text': text})
            in_notice = True
            logger.debug(f'  考前内容(注意事项标题): {text[:40]}')
            continue

        # 注意事项后的编号列表项（如"1.作答前..."）
        if in_notice:
            # 以"数字."开头的列表项
            import re
            if re.match(r'^\d+[.．、]', text):
                pre_content.append({'type': 'notice_item', 'text': text})
                logger.debug(f'  考前内容(注意事项条目): {text[:40]}')
                continue
            # 空行或其他内容可能表示注意事项结束
            if not text:
                in_notice = False
                continue

        # 考生信息栏
        if _is_pre_exam_content(text):
            pre_content.append({'type': 'student_info', 'text': text})
            logger.debug(f'  考前内容(考生信息): {text[:40]}')
            continue

    logger.info(f'考前内容: {len(pre_content)} 项')
    return pre_content


def _identify_sections(elements, logger, uncertain_paragraphs=None):
    """识别分区边界，返回分区列表。
    
    每个分区: {'type': 'paragraph'|'table', 'section_type': '选择题'|'非选择题'|'填空题'|'未知', 'elements': [...]}
    
    当段落看起来像分区标题但无法匹配标准正则时（如"第I卷"），
    记录到uncertain_paragraphs供AI审查。
    """
    sections = []
    current_section = {
        'section_type': '未知',
        'section_title': '',
        'elements': [],
    }

    section_started = False
    para_index = 0

    for elem in elements:
        if elem['type'] == 'paragraph':
            text = elem['text']
            para_index += 1

            # 检查是否为分区标题
            if is_section_header(text):
                # 保存之前的分区
                if section_started and current_section['elements']:
                    sections.append(current_section)

                # 确定分区类型（先检查非选择题，因为"非选择题"包含"选择题"）
                if is_non_choice_section(text):
                    section_type = '非选择题'
                elif is_choice_section(text):
                    section_type = '选择题'
                elif is_fill_section(text):
                    section_type = '填空题'
                else:
                    section_type = '未知'

                current_section = {
                    'section_type': section_type,
                    'section_title': text,
                    'elements': [],
                }
                section_started = True
                logger.debug(f'  分区: {section_type} - {text[:40]}')
                continue
            else:
                # 检查是否疑似分区标题但无法匹配（如"第I卷""Part 1"等）
                if _looks_like_section_header(text):
                    if uncertain_paragraphs is not None:
                        uncertain_paragraphs.append({
                            'paragraph_index': para_index,
                            'text': text,
                            'context_before': _get_context_text(elements, para_index - 2, 1),
                            'context_after': _get_context_text(elements, para_index, 1),
                            'has_image': elem.get('has_image', False),
                            'image_names': [],
                            'script_guess': '疑似分区标题',
                            'possible_types': ['section_header', 'instruction', 'material'],
                        })
                        logger.debug(f'  [不确定] 疑似分区标题: {text[:40]}')

        # 添加元素到当前分区
        if section_started:
            current_section['elements'].append(elem)
        else:
            # 分区标题之前的元素（考试信息等），跳过
            pass

    # 保存最后一个分区
    if section_started and current_section['elements']:
        sections.append(current_section)

    return sections


def _looks_like_section_header(text):
    """判断文本是否疑似分区标题但无法被标准正则匹配。
    
    检测以下模式：
    - "第I卷""第II卷""第1卷"等
    - "Part 1""Part I"等
    - "一、选择题"的变体（如"一.选择题""一、 选择题"）
    - 短文本（<20字）且包含"题""卷""部分"等关键词
    """
    if not text or len(text) > 30:
        return False
    
    # "第X卷"模式
    if re.match(r'^第\s*[IⅠIIⅡ12一三四五六七八九十]+\s*卷', text):
        return True
    
    # "Part X"模式
    if re.match(r'^Part\s+[IVX1-9]', text, re.IGNORECASE):
        return True
    
    # 包含"卷""部分"且较短
    if ('卷' in text or '部分' in text) and len(text) < 20:
        return True
    
    # 中文数字+非顿号开头的疑似分区标题
    if re.match(r'^[一二三四五六七八九十][.．、\s]', text) and len(text) < 20:
        return True
    
    return False


def _get_context_text(elements, start_idx, count):
    """从elements列表中获取指定位置的上下文文本。
    
    Args:
        elements: 有序元素列表
        start_idx: 起始索引
        count: 获取的元素数量
    """
    texts = []
    for i in range(start_idx, min(start_idx + count, len(elements))):
        if i < 0 or i >= len(elements):
            continue
        elem = elements[i]
        if elem['type'] == 'paragraph':
            t = elem.get('text', '')
            if t:
                texts.append(t[:50])
    return ' | '.join(texts)


def _generate_document_markdown(elements):
    """将文档元素列表转换为Markdown格式文本，供AI审查。
    
    保留加粗、下划线等格式标记，用Markdown语法表示。
    """
    lines = []
    for elem in elements:
        if elem['type'] == 'table':
            lines.append(f'[表格]')
            continue
        
        if elem['type'] != 'paragraph':
            continue
        
        text = elem.get('text', '')
        has_img = elem.get('has_image', False)
        
        if not text and not has_img:
            continue
        
        if has_img and not text:
            lines.append(f'[图片]')
        elif has_img and text:
            lines.append(f'{text} [图片]')
        else:
            lines.append(text)
    
    return '\n'.join(lines)


def _parse_section(section, sec_idx, doc, docx_path, media_info, rels, images_dir, image_counter, logger, uncertain_paragraphs=None):
    """解析一个分区，返回结构化数据。"""
    section_type = section['section_type']
    section_title = section['section_title']
    elements = section['elements']

    logger.info(f'  解析分区 {sec_idx}: {section_type} ({len(elements)} 个元素)')

    # 识别题组
    groups = _identify_question_groups(elements, section_type, logger)

    # 解析每个题组
    parsed_groups = []
    for group in groups:
        parsed_group = _parse_question_group(
            group, section_type, doc, docx_path, media_info, rels,
            images_dir, image_counter, logger, uncertain_paragraphs, sec_idx
        )
        parsed_groups.append(parsed_group)

    return {
        'section_id': sec_idx,
        'section_title': section_title,
        'section_type': section_type,
        'question_groups': parsed_groups,
    }


def _identify_question_groups(elements, section_type, logger):
    """将分区内的元素分组为题组。
    
    题组边界判断：
    - 选择题：遇到材料段落（非题干、非选项）开始新题组
    - 非选择题：每个题号开始新题组
    
    返回: [[elem1, elem2, ...], [elem1, ...], ...]
    """
    groups = []
    current_group = []
    prev_question_num = None

    for elem in elements:
        if elem['type'] == 'paragraph':
            text = elem['text']

            # 空段落：跳过
            if not text and not elem['has_image']:
                continue

            # 检查是否为题干
            q_info = get_question_number(text) if is_question_stem(text) else None

            if q_info:
                q_num = q_info[0]

                # 判断是否需要开始新题组
                if section_type == '非选择题':
                    # 非选择题：每个题号都是新题组
                    if current_group:
                        groups.append(current_group)
                    current_group = [elem]
                    prev_question_num = q_num
                else:
                    # 选择题：题号不连续时开始新题组
                    if prev_question_num is not None and q_num != prev_question_num + 1:
                        if current_group:
                            groups.append(current_group)
                        current_group = [elem]
                    else:
                        current_group.append(elem)
                    prev_question_num = q_num
            else:
                # 非题干元素（材料、选项、图片、表格等）
                # 如果当前组为空，或者前一个元素是选项（说明新材料的开始）
                if not current_group:
                    current_group.append(elem)
                elif current_group[-1]['type'] == 'paragraph':
                    last_text = current_group[-1].get('text', '')
                    # 如果前一个是选项，且当前是材料，开始新题组
                    if is_option(last_text) and text and not is_option(text):
                        # 检查是否真的是材料（而不是选项的延续）
                        if not is_sub_option(text):
                            groups.append(current_group)
                            current_group = [elem]
                            prev_question_num = None
                        else:
                            current_group.append(elem)
                    else:
                        current_group.append(elem)
                else:
                    current_group.append(elem)
        else:
            # 表格元素
            current_group.append(elem)

    # 保存最后一个题组
    if current_group:
        groups.append(current_group)

    logger.debug(f'    识别到 {len(groups)} 个题组')
    return groups


def _parse_question_group(group_elements, section_type, doc, docx_path, media_info, rels,
                          images_dir, image_counter, logger, uncertain_paragraphs=None, sec_idx=0):
    """解析一个题组，返回结构化数据。"""
    materials = []
    instruction = ''
    questions = []
    tables_data = []

    # 临时存储
    current_material = {'segments': []}
    current_question = None
    current_options = {}
    current_sub_options = []
    has_question_started = False  # 标记是否已遇到第一个题干

    # 辅助函数：保存当前题目
    def save_current_question():
        nonlocal current_question, current_options, current_sub_options
        if current_question:
            if current_question['question_type'] == '选择题':
                current_question['options'] = current_options if current_options else None
                current_question['sub_options'] = current_sub_options if current_sub_options else None
                current_question['sub_questions'] = None
            else:
                # 非选择题
                current_question['options'] = None
                current_question['sub_options'] = None
                # 确保sub_questions列表存在
                if 'sub_questions' not in current_question or current_question['sub_questions'] is None:
                    current_question['sub_questions'] = []
                # 清理子问题中的空sub_options列表
                for sq in current_question['sub_questions']:
                    if not sq.get('sub_options'):
                        sq['sub_options'] = None
            questions.append(current_question)
        current_question = None
        current_options = {}
        current_sub_options = []

    # 辅助函数：保存当前材料
    def save_current_material():
        nonlocal current_material, instruction
        if current_material['segments']:
            # 从材料segments中提取引导语（查找最后一个text segment）
            for seg in reversed(current_material['segments']):
                if seg['type'] == 'text' and seg['content']:
                    mat_text, instr = extract_instruction(seg['content'])
                    seg['content'] = mat_text
                    if instr and not instruction:
                        instruction = instr
                    break
            materials.append(dict(current_material))
        current_material = {'segments': []}

    # 辅助函数：提取图片（prefix 控制命名前缀）
    # 返回 [{"name": img_name, "width_cm": ..., "height_cm": ...}, ...]
    def extract_image(elem, prefix='material'):
        nonlocal image_counter
        drawings = get_drawings(elem['element'])
        image_infos = []
        for drawing in drawings:
            rid = get_embed_rid(drawing)
            if not rid:
                continue
            media_path = rels.get(rid)
            if not media_path:
                continue
            info = media_info.get(media_path, {})
            file_size = info.get('size', 0)
            # 跳过小图片（可能是残留的标点图片）
            if file_size < 512:
                continue
            image_counter[0] += 1
            ext = info.get('ext', '.png')
            img_name = f'{prefix}_{image_counter[0]:03d}{ext}'
            img_path = os.path.join(images_dir, img_name)
            img_data = read_media_file(docx_path, media_path)
            if img_data:
                with open(img_path, 'wb') as f:
                    f.write(img_data)
                img_info = {
                    'name': img_name,
                    'width_cm': info.get('width_cm'),
                    'height_cm': info.get('height_cm'),
                }
                image_infos.append(img_info)
                logger.debug(f'    提取图片: {img_name} ({file_size}字节, {img_info.get("width_cm")}x{img_info.get("height_cm")}cm)')
        return image_infos

    # 辅助函数：解析含图片的选项段落
    # 遍历 run 序列，将图片关联到最近的选项字母，生成 {{IMAGE:...}} 占位符
    def parse_options_with_images(elem):
        """解析选项段落，提取图片并生成占位符。
        返回 {letter: content_with_placeholders} 或 None。
        """
        import re as _re

        p_element = elem['element']
        # 按 run 顺序收集文本和图片
        segments = []  # [(type, content), ...] type='text'|'image'
        for child in p_element:
            if child.tag == qn('w:r'):
                # 检查 run 中的文本
                run_text = ''
                for t in child.iter(qn('w:t')):
                    if t.text:
                        run_text += t.text
                if run_text:
                    segments.append(('text', run_text))
                # 检查 run 中的图片
                for drawing in child.iter(qn('w:drawing')):
                    rid = get_embed_rid(drawing)
                    if not rid:
                        continue
                    media_path = rels.get(rid)
                    if not media_path:
                        continue
                    info = media_info.get(media_path, {})
                    file_size = info.get('size', 0)
                    if file_size < 512:
                        continue
                    # 提取图片，命名需要知道选项字母，先记录rid稍后处理
                    segments.append(('image_rid', rid))

        if not any(s[0] == 'image_rid' for s in segments):
            # 没有图片，用普通解析
            return parse_options(elem['text'])

        # 有图片：按选项字母分组
        options = {}
        current_letter = None
        current_text = ''

        for seg_type, seg_content in segments:
            if seg_type == 'text':
                # 检查文本中是否有选项字母标记
                remaining = seg_content
                while remaining:
                    # 查找 A./B./C./D. 标记
                    match = _re.search(r'([ABCD])\.\s*', remaining)
                    if match:
                        # 保存之前选项的文本
                        if current_letter:
                            options[current_letter] = current_text
                        # 开始新选项
                        before_match = remaining[:match.start()]
                        if before_match and current_letter:
                            options[current_letter] += before_match
                        current_letter = match.group(1)
                        current_text = ''
                        remaining = remaining[match.end():]
                    else:
                        current_text += remaining
                        break
            elif seg_type == 'image_rid':
                # 提取图片，用选项字母命名
                if current_letter:
                    prefix = f'option_{current_letter}'
                    img_names = _extract_image_by_rid(seg_content, prefix)
                    if img_names:
                        for img_name in img_names:
                            current_text += '{{IMAGE:' + img_name + '}}'
                else:
                    # 图片在选项字母之前，暂时归为材料图片
                    img_names = _extract_image_by_rid(seg_content, 'material')
                    # 这些图片可能是题干的一部分

        # 保存最后一个选项
        if current_letter:
            options[current_letter] = current_text

        return options if options else None

    # 辅助函数：通过 rid 提取图片
    def _extract_image_by_rid(rid, prefix):
        nonlocal image_counter
        media_path = rels.get(rid)
        if not media_path:
            return []
        info = media_info.get(media_path, {})
        file_size = info.get('size', 0)
        if file_size < 512:
            return []
        image_counter[0] += 1
        ext = info.get('ext', '.png')
        img_name = f'{prefix}_{image_counter[0]:03d}{ext}'
        img_path = os.path.join(images_dir, img_name)
        img_data = read_media_file(docx_path, media_path)
        if img_data:
            with open(img_path, 'wb') as f:
                f.write(img_data)
            logger.debug(f'    提取选项图片: {img_name} ({file_size}字节)')
            return [img_name]
        return []

    for elem in group_elements:
        if elem['type'] == 'table':
            # 表格 → append table segment
            table_dict = table_to_dict(elem['table_obj'])
            table_id = f'table_{len(tables_data) + 1}'
            table_role = 'question' if has_question_started else 'material'
            tables_data.append({'id': table_id, 'role': table_role, **table_dict})
            current_material['segments'].append({
                'type': 'table',
                'table_id': table_id,
                'role': table_role,
            })
            continue

        if elem['type'] != 'paragraph':
            continue

        text = elem['text']
        has_img = elem['has_image']

        # 空段落
        if not text and not has_img:
            continue

        # 纯图片段落 → append image segment(s)
        if has_img and not text:
            img_infos = extract_image(elem)
            for img_info in img_infos:
                current_material['segments'].append({
                    'type': 'image',
                    'name': img_info['name'],
                    'width_cm': img_info.get('width_cm'),
                    'height_cm': img_info.get('height_cm'),
                })
            continue

        # 检查是否为题干
        q_info = get_question_number(text) if is_question_stem(text) else None

        if q_info:
            # 保存前一个题目
            save_current_question()

            q_num = q_info[0]
            q_content = q_info[1]

            # 检查是否为非选择题标题
            if is_non_choice_title(text):
                # 非选择题标题
                # 保存之前的材料
                save_current_material()

                current_question = {
                    'question_number': q_num,
                    'question_type': '非选择题',
                    'stem': q_content,  # 如"阅读图文材料，完成下列要求。"
                    'sub_questions': [],
                }
            else:
                # 选择题题干
                # 如果之前有材料且没有题目，说明材料属于这个题组
                save_current_material()

                current_question = {
                    'question_number': q_num,
                    'question_type': section_type if section_type != '未知' else '选择题',
                    'stem': q_content,
                }
            has_question_started = True
            continue

        # 检查是否为选项
        if is_option(text):
            if has_img:
                # 选项含图片：提取图片并生成占位符
                parsed_opts = parse_options_with_images(elem)
            else:
                parsed_opts = parse_options(text)
            if parsed_opts:
                current_options.update(parsed_opts)
            continue

        # 检查是否为子选项（①②③④）
        if is_sub_option(text) or contains_sub_option(text):
            # 拆分同行多子选项，如"①苹果②香蕉③橙子" → ["①苹果", "②香蕉", "③橙子"]
            sub_parts = split_sub_options(text)
            if not sub_parts:
                sub_parts = [text]
            
            if current_question and current_question['question_type'] == '非选择题':
                # 非选择题：将①②③子选项附加到最近的子问题
                if current_question.get('sub_questions'):
                    current_question['sub_questions'][-1].setdefault('sub_options', []).extend(sub_parts)
                else:
                    current_sub_options.extend(sub_parts)
            else:
                # 选择题：保持原有行为
                current_sub_options.extend(sub_parts)
            continue

        # 检查是否为子问题（（1）（2））
        if is_sub_question(text):
            match = SUB_QUESTION_PATTERN.match(text)
            if match:
                sub_id = text[:match.start(1)].strip()
                sub_text = match.group(1)
                if current_question and current_question.get('sub_questions') is not None:
                    current_question['sub_questions'].append({
                        'sub_id': sub_id,
                        'text': sub_text,
                        'sub_options': [],
                    })
                continue

        # 检查是否为引导语（独立段落的引导语）
        if is_instruction(text):
            instruction = text
            continue

        # ===== 不确定段落追踪 =====
        # 当段落不匹配任何已知模式（非题干/选项/子选项/子问题/引导语）时
        # 记录到uncertain_paragraphs供AI审查
        # 同时仍按现有逻辑归入材料（保证内容不丢失）
        if uncertain_paragraphs is not None and text:
            # 获取前后段落文本作为上下文
            idx_in_group = group_elements.index(elem) if elem in group_elements else -1
            before_text = ''
            after_text = ''
            if idx_in_group > 0:
                prev = group_elements[idx_in_group - 1]
                before_text = prev.get('text', '')[:50] if prev['type'] == 'paragraph' else '[表格]'
            if idx_in_group >= 0 and idx_in_group < len(group_elements) - 1:
                nxt = group_elements[idx_in_group + 1]
                after_text = nxt.get('text', '')[:50] if nxt['type'] == 'paragraph' else '[表格]'

            uncertain_paragraphs.append({
                'paragraph_index': len(uncertain_paragraphs) + 1,
                'section_id': sec_idx,
                'section_type': section_type,
                'text': text[:200],
                'context_before': before_text,
                'context_after': after_text,
                'has_image': has_img,
                'image_names': [],
                'script_guess': '材料' if not is_question_text(text) else '子问题(未匹配动词)',
                'possible_types': ['material', 'sub_question', 'instruction', 'option', 'sub_option'],
            })

        # 其他文字处理（使用segments数组）
        # 对于非选择题：需要区分材料和子问题
        # 对于选择题：归为材料
        
        # 多材料标记检测：遇到"材料一""材料二"等标记时，保存当前材料并开始新材料
        MATERIAL_MARKER_PATTERN = re.compile(r'^材料[一二三四五六七八九十\d]+[：:、\s]')
        is_new_material = MATERIAL_MARKER_PATTERN.match(text)
        
        if current_question and current_question['question_type'] == '非选择题':
            # 非选择题的文字：判断是材料还是子问题
            if is_question_text(text):
                # 子问题（无编号，但含问题动词）
                current_question['sub_questions'].append({
                    'sub_id': '',
                    'text': text,
                    'sub_options': [],
                })
            else:
                # 材料
                if is_new_material and current_material['segments']:
                    # 遇到新材料标记，保存之前的材料，开始新材料
                    save_current_material()
                
                # append text segment
                current_material['segments'].append({
                    'type': 'text',
                    'content': text,
                })
                
                # 如果段落包含图片 → append image segment(s)
                if has_img:
                    img_infos = extract_image(elem)
                    for img_info in img_infos:
                        current_material['segments'].append({
                            'type': 'image',
                            'name': img_info['name'],
                            'width_cm': img_info.get('width_cm'),
                            'height_cm': img_info.get('height_cm'),
                        })
                # 兼容：解析中文图片占位符（当图片已被extract_images.py删除时）
                elif CHINESE_IMAGE_PLACEHOLDER.search(text):
                    for img_match in CHINESE_IMAGE_PLACEHOLDER.finditer(text):
                        img_name = img_match.group(1).strip()
                        current_material['segments'].append({
                            'type': 'image',
                            'name': img_name,
                            'width_cm': None,
                            'height_cm': None,
                        })
                    logger.debug(f'  从占位符解析图片: {[m.group(1) for m in CHINESE_IMAGE_PLACEHOLDER.finditer(text)]}')
        else:
            # 选择题的材料
            if is_new_material and current_material['segments']:
                # 遇到新材料标记，保存之前的材料，开始新材料
                save_current_material()
            
            # append text segment
            current_material['segments'].append({
                'type': 'text',
                'content': text,
            })
            
            # 如果段落包含图片 → append image segment(s)
            if has_img:
                img_infos = extract_image(elem)
                for img_info in img_infos:
                    current_material['segments'].append({
                        'type': 'image',
                        'name': img_info['name'],
                        'width_cm': img_info.get('width_cm'),
                        'height_cm': img_info.get('height_cm'),
                    })
            # 兼容：解析中文图片占位符（当图片已被extract_images.py删除时）
            elif CHINESE_IMAGE_PLACEHOLDER.search(text):
                for img_match in CHINESE_IMAGE_PLACEHOLDER.finditer(text):
                    img_name = img_match.group(1).strip()
                    current_material['segments'].append({
                        'type': 'image',
                        'name': img_name,
                        'width_cm': None,
                        'height_cm': None,
                    })
                logger.debug(f'  从占位符解析图片(选择题): {[m.group(1) for m in CHINESE_IMAGE_PLACEHOLDER.finditer(text)]}')

    # 保存最后的题目和材料
    save_current_question()
    save_current_material()

    # 生成 group_id
    if questions:
        first_q = questions[0]['question_number']
        last_q = questions[-1]['question_number']
        if first_q == last_q:
            group_id = str(first_q)
        else:
            group_id = f'{first_q}-{last_q}'
    else:
        group_id = 'unknown'

    return {
        'group_id': group_id,
        'materials': materials,
        'instruction': instruction,
        'tables': tables_data,
        'questions': questions,
    }


def main():
    parser = argparse.ArgumentParser(
        description='地理试卷打标脚本 - 解析清洗后的试卷文档，输出结构化 JSON',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例:
  python tag_docx.py --input cleaned.docx --output tagged.json
  python tag_docx.py --input cleaned.docx --output tagged.json --images-dir ./images
        '''
    )
    parser.add_argument('--input', '-i', required=True, help='输入 docx 文件路径（清洗后）')
    parser.add_argument('--output', '-o', required=True, help='输出 JSON 文件路径')
    parser.add_argument('--log', '-l', help='日志文件路径（默认与输出文件同目录）')
    parser.add_argument('--images-dir', help='图片提取目录（默认与输出文件同目录的 images/）')

    args = parser.parse_args()

    success = tag_docx(args.input, args.output, args.log, args.images_dir)

    if not success:
        sys.exit(1)


if __name__ == '__main__':
    main()
