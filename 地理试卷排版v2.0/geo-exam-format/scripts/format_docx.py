# -*- coding: utf-8 -*-
"""
地理试卷排版脚本
基于打标 JSON 和模板，生成排版好的 Word 文档。

用法:
    python format_docx.py --json tagged.json --template assets/template.dotx --output formatted.docx

输出:
    formatted.docx      排版后的文档
    format_log.txt      排版日志
"""

import argparse
import json
import os
import re
import shutil
import sys
import tempfile
from copy import deepcopy
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn as docx_qn
from lxml import etree

try:
    from PIL import Image as PILImage
except ImportError:
    PILImage = None

# 页面尺寸常量（A4，与模板边距一致）
# 纸张 21cm x 29.7cm，边距 上下2.54cm 左右1.9cm
PAGE_CONTENT_WIDTH_CM = 21.0 - 1.9 - 1.9      # 版心宽度 17.2cm
PAGE_CONTENT_HEIGHT_CM = 29.7 - 2.54 - 2.54   # 版心高度 24.62cm


# =====================================================================
# 日志
# =====================================================================

import logging

def setup_logger(log_path):
    logger = logging.getLogger('geo_exam_format')
    logger.setLevel(logging.DEBUG)
    logger.handlers = []

    fh = logging.FileHandler(log_path, mode='w', encoding='utf-8')
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(logging.Formatter('%(asctime)s [%(levelname)s] %(message)s'))
    logger.addHandler(fh)

    sh = logging.StreamHandler()
    sh.setLevel(logging.INFO)
    sh.setFormatter(logging.Formatter('%(message)s'))
    logger.addHandler(sh)

    return logger


# =====================================================================
# 模板加载（dotx -> docx 转换）
# =====================================================================

def load_template(template_path):
    """加载 dotx 模板文件。
    
    python-docx 不支持直接打开 .dotx（content-type 不匹配），
    需要先复制为临时 .docx 并修改 [Content_Types].xml 中的类型。
    返回 Document 对象。
    """
    import zipfile
    
    # 创建临时 docx 文件（从 dotx 复制并修改 content-type）
    temp_docx = template_path + '.tmp.docx'
    
    with zipfile.ZipFile(template_path, 'r') as zin:
        with zipfile.ZipFile(temp_docx, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                data = zin.read(item)
                # 修改 content-type：template -> document
                if item == '[Content_Types].xml':
                    data = data.replace(
                        b'template.main+xml',
                        b'document.main+xml'
                    )
                zout.writestr(item, data)
    
    try:
        doc = Document(temp_docx)
        # 清空模板中的所有内容（保留样式定义和 sectPr）
        body = doc.element.body
        for child in list(body):
            if child.tag != docx_qn('w:sectPr'):
                body.remove(child)
        return doc
    finally:
        if os.path.exists(temp_docx):
            os.remove(temp_docx)


# =====================================================================
# 样式应用工具
# =====================================================================

# 模板样式ID映射
STYLE_MAP = {
    '考试名称': '考试名称',
    '科目名称': '科目名称',
    '题型标题': '题型标题',
    'Body Text': 'Body Text',
    'Normal': 'Normal',
    'Table Grid': 'Table Grid',
    '注意事项内容': '注意事项内容',
    '第几卷': '第几卷',
    '选项': '选项',
}


def apply_style(paragraph, style_name, logger=None):
    """给段落应用样式。如果样式不存在，使用 Normal。"""
    try:
        paragraph.style = style_name
        if logger:
            logger.debug(f'    样式应用成功: {style_name}')
    except KeyError:
        if logger:
            logger.warning(f'    样式不存在: {style_name}, 回退到Normal')
        try:
            paragraph.style = 'Normal'
        except KeyError:
            pass


def add_run_with_text(paragraph, text):
    """给段落添加文本 run。"""
    run = paragraph.add_run(text)
    return run


def set_paragraph_alignment(paragraph, alignment):
    """设置段落对齐方式。"""
    paragraph.alignment = alignment


# =====================================================================
# 图片处理
# =====================================================================

IMAGE_PLACEHOLDER_PATTERN = re.compile(r'\{\{IMAGE:([^}]+)\}\}')

# 图片占位符正则（兼容AI/脚本打标输出的两种占位符格式）
# 格式1: 【图片：filename - description】
# 格式2: {{IMAGE:filename}}
# 提取文件名，忽略描述部分
CHINESE_IMAGE_PLACEHOLDER = re.compile(r'【图片[：:]\s*(\S+?)(?:\s*[-—]\s*.+?)?】')


def get_image_size_cm(image_path):
    """读取图片的物理尺寸（厘米）。
    
    通过 PIL 读取像素尺寸和 DPI，换算为厘米。
    无 DPI 信息时默认 96 DPI。
    返回 (width_cm, height_cm)，读取失败返回 None。
    """
    if PILImage is None:
        return None
    try:
        with PILImage.open(image_path) as im:
            px_w, px_h = im.size
            dpi = im.info.get('dpi', (96, 96))
            dpi_x = dpi[0] if dpi[0] else 96
            dpi_y = dpi[1] if dpi[1] else 96
            w_cm = px_w / dpi_x * 2.54
            h_cm = px_h / dpi_y * 2.54
            return w_cm, h_cm
    except Exception:
        return None


def compute_image_display_width(image_path):
    """计算图片的合适显示宽度（EMU），保证完整显示在页面内。
    
    规则：
    - 宽度不超过 12cm，高度不超过 8cm（临时限制，后续按图片内容单独处理）
    - 等比缩放（锁定长宽比），只缩小不放大（避免放大模糊）
    - 通过只设置 width 实现：Word 按原始宽高比自动计算高度，保证不拉伸
    返回 EMU 单位的宽度值，失败返回 12cm。
    """
    max_w_cm = 12.0
    max_h_cm = 8.0

    size = get_image_size_cm(image_path)
    if size is None:
        return Cm(max_w_cm)

    orig_w, orig_h = size
    if orig_w <= 0 or orig_h <= 0:
        return Cm(max_w_cm)

    # 等比缩放：取宽度限制和高度限制中较小的缩放比例
    scale = 1.0
    if orig_w > max_w_cm:
        scale = min(scale, max_w_cm / orig_w)
    if orig_h > max_h_cm:
        scale = min(scale, max_h_cm / orig_h)

    # 只返回宽度，Word 会按图片原始宽高比自动确定高度，从而锁定长宽比
    display_w = orig_w * scale
    return Cm(display_w)


def add_image_to_paragraph(paragraph, image_path, width=None):
    """在段落中添加图片。width 为 None 时自动计算合适宽度。"""
    if not os.path.exists(image_path):
        # 图片文件不存在，添加占位文字
        run = paragraph.add_run(f'[图片缺失: {os.path.basename(image_path)}]')
        return
    
    if width is None:
        width = compute_image_display_width(image_path)
    
    run = paragraph.add_run()
    run.add_picture(image_path, width=width)


def add_image_centered(doc, image_path, logger):
    """添加居中的图片段落，自动缩放到页面范围内。"""
    p = doc.add_paragraph()
    set_paragraph_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)
    apply_style(p, 'Normal')
    width = compute_image_display_width(image_path)
    add_image_to_paragraph(p, image_path, width=width)
    logger.debug(f'  添加图片: {os.path.basename(image_path)} (显示宽度={width.cm:.2f}cm)')
    return p


# =====================================================================
# 表格处理
# =====================================================================

def add_table(doc, table_data, logger, role='material'):
    """添加表格。

    格式规范：
    - 表头行：中文黑体、英文/数字Times New Roman、加粗、居中
    - 内容行：中文宋体、英文/数字Times New Roman、单倍行距、居中
    - 表头行跨页重复
    - 表格宽度适应版心（17.2cm）

    Args:
        role: 'material' 表示材料表格，'question' 表示题目表格。
              两者表头均为黑体，内容均为宋体；差异在于排版上下文。
    """
    rows = table_data.get('rows', 0)
    cols = table_data.get('cols', 0)
    data = table_data.get('data', [])

    if rows == 0 or cols == 0 or not data:
        return None

    table = doc.add_table(rows=rows, cols=cols)
    apply_style(table, 'Table Grid')

    # 设置表格宽度为版心宽度
    table.autofit = False
    table.allow_autofit = False
    table.width = Cm(PAGE_CONTENT_WIDTH_CM)

    # 填充数据
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            if j < cols:
                cell = table.cell(i, j)
                # 清除单元格默认内容
                cell.text = ''
                paragraph = cell.paragraphs[0]
                # 居中对齐
                set_paragraph_alignment(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
                # 单倍行距
                paragraph.paragraph_format.line_spacing = 1.0
                # 使用混合字体添加文本：中文用黑体/宋体，英文数字用Times New Roman
                cell_str = str(cell_text)
                if i == 0:
                    # 表头行：黑体加粗
                    add_text_mixed_fonts(paragraph, cell_str,
                                        cn_font='黑体', en_font='Times New Roman')
                    for run in paragraph.runs:
                        run.bold = True
                else:
                    # 内容行：宋体
                    add_text_mixed_fonts(paragraph, cell_str,
                                        cn_font='宋体', en_font='Times New Roman')

    # 设置表头行重复
    if rows > 1:
        tbl = table._tbl
        tblPr = tbl.find(docx_qn('w:tblPr'))
        if tblPr is None:
            tblPr = etree.SubElement(tbl, docx_qn('w:tblPr'))
        tblHeader = tblPr.find(docx_qn('w:tblHeader'))
        if tblHeader is None:
            tblHeader = etree.SubElement(tblPr, docx_qn('w:tblHeader'))
        tblHeader.set(docx_qn('w:val'), 'true')

    role_label = '材料' if role == 'material' else '题目'
    logger.debug(f'  添加表格({role_label}): {rows}行 x {cols}列, 宽度={PAGE_CONTENT_WIDTH_CM}cm')
    return table


# =====================================================================
# 选项排版
# =====================================================================

# 版心宽度（cm）：21 - 1.9 - 1.9 = 17.2cm
# 中文字符约 0.4cm/字（四号字），英文约 0.2cm/字符
# 版心宽度约可容纳 43 个中文字符

PAGE_WIDTH_CHARS = 43  # 版心可容纳的中文字符数

# 模板"选项"样式中的制表位位置（cm），用于选项排版对齐
# 来源：模板 dotx 中样式 ID=13 "选项" 的 w:tabs 定义
TAB_STOPS_4 = [Cm(4.54), Cm(8.98), Cm(13.43)]  # 四选项一行的制表位
TAB_STOPS_2 = [Cm(8.98), Cm(9.0)]  # 两选项一行：两个相近制表位，配合两个Tab使用


def set_tab_stops(paragraph, tab_positions):
    """在段落上设置制表位。

    Args:
        paragraph: python-docx Paragraph 对象
        tab_positions: 制表位位置列表（EMU 单位，如 Cm(4.54) 的返回值）
    """
    pPr = paragraph._element.get_or_add_pPr()
    # 移除已有的制表位
    for existing in pPr.findall(docx_qn('w:tabs')):
        pPr.remove(existing)
    # 添加新的制表位
    tabs = etree.SubElement(pPr, docx_qn('w:tabs'))
    for pos in tab_positions:
        tab = etree.SubElement(tabs, docx_qn('w:tab'))
        tab.set(docx_qn('w:val'), 'left')
        # pos 是 EMU 单位，需要转换为 twips（1 cm = 567 twips, 1 EMU = 1/914400 inch, 1 inch = 1440 twips）
        # EMU to twips: twips = EMU * 1440 / 914400 = EMU / 635
        twips = str(int(pos / 635))
        tab.set(docx_qn('w:pos'), twips)


def clear_run_fonts(paragraph):
    """清除段落中所有 run 的字体覆盖，让样式统一控制字体。"""
    for run in paragraph.runs:
        rPr = run._element.find(docx_qn('w:rPr'))
        if rPr is not None:
            # 删除 rFonts 和 sz 元素，保留 bold/italic 等
            for tag in ['w:rFonts', 'w:sz', 'w:szCs']:
                for el in rPr.findall(docx_qn(tag)):
                    rPr.remove(el)


def _set_run_font(run, cn_font=None, en_font=None, size_pt=None, force_cn_all=False):
    """设置 run 的字体（中文字体 / 西文字体 / 字号）。

    - force_cn_all=True 时，将 ascii/hAnsi/eastAsia 三个字体槽全部设为 cn_font，
      用于中文标点（如弯引号“”‘’）确保用中文字体渲染。
    """
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(docx_qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, docx_qn('w:rFonts'))
    if force_cn_all and cn_font:
        # 三个字体槽全部用中文字体，确保引号等中文标点不被英文字体渲染
        rFonts.set(docx_qn('w:ascii'), cn_font)
        rFonts.set(docx_qn('w:hAnsi'), cn_font)
        rFonts.set(docx_qn('w:eastAsia'), cn_font)
    else:
        if en_font:
            rFonts.set(docx_qn('w:ascii'), en_font)
            rFonts.set(docx_qn('w:hAnsi'), en_font)
        if cn_font:
            rFonts.set(docx_qn('w:eastAsia'), cn_font)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


# 中文标点字符集（弯引号、直角引号等）——这些字符需要强制用中文字体
_CN_PUNCT_CHARS = set('“”‘’〈〉《》「」『』【】〔〕〖〗〘〙〚〛…—·、。，；：？！·')


def _classify_char(ch):
    """字符分类：'en'（数字/英文字母）或 'cn'（中文/标点/引号等）。

    - ASCII 数字和字母 → 'en'
    - ASCII 单引号 '（U+0027）→ 'en'（用于经纬度如 64°48'N，保持英文字体）
    - Unicode 上标/下标字符（U+2070~U+208F）→ 'en'（化学方程式如CaCO₃，使用Times New Roman）
    - 中文弯引号 ""'' 及其他中文标点 → 'cn'（强制中文字体）
    - 其余（中文、中文标点等）→ 'cn'
    """
    # ASCII 数字和字母归为西文
    if ch.isascii() and (ch.isdigit() or ch.isalpha()):
        return 'en'
    # ASCII 单引号（经纬度场景）保持英文
    if ch == "'":
        return 'en'
    # Unicode 上标/下标字符归为西文（化学方程式、数学公式）
    # U+2070~U+207F: 上标 ⁰ⁱ²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿ
    # U+2080~U+208F: 下标 ₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎
    code = ord(ch)
    if 0x2070 <= code <= 0x208F:
        return 'en'
    # 其余（中文、中文标点、中文引号等）归为中文
    return 'cn'


def _is_cn_punct(ch):
    """判断字符是否为需要强制中文字体的中文标点（主要是引号）。"""
    return ch in _CN_PUNCT_CHARS


def add_text_mixed_fonts(paragraph, text, cn_font=None, en_font='Times New Roman', size_pt=None):
    """添加文本，按字符类型分段设置字体。

    - 数字和英文字母 → en_font（默认 Times New Roman）
    - ASCII 单引号 ' → en_font（经纬度场景，保持英文）
    - 其他字符（中文、中文标点、中文弯引号等）→ cn_font
    - 中文弯引号 “”‘’ 等通过 force_cn_all 强制三个字体槽都用中文字体，
      确保不被 Word 用英文字体渲染
    """
    if not text:
        return
    # 按字符类型分段
    segments = []
    current_type = _classify_char(text[0])
    current_chars = [text[0]]
    for ch in text[1:]:
        t = _classify_char(ch)
        if t == current_type:
            current_chars.append(ch)
        else:
            segments.append((''.join(current_chars), current_type))
            current_chars = [ch]
            current_type = t
    segments.append((''.join(current_chars), current_type))

    for seg_text, seg_type in segments:
        run = paragraph.add_run(seg_text)
        if seg_type == 'en':
            _set_run_font(run, cn_font=cn_font, en_font=en_font, size_pt=size_pt)
        else:
            # 中文段：若包含中文标点（引号等），强制三个字体槽都用中文字体
            force = any(_is_cn_punct(c) for c in seg_text)
            _set_run_font(run, cn_font=cn_font, en_font=en_font, size_pt=size_pt,
                          force_cn_all=force)


def estimate_text_length(text):
    """估算文本在版心中占用的字符宽度（旧接口，保留兼容）。
    中文1字=1，英文2字=1，数字=0.5
    """
    length = 0
    for ch in text:
        if '\u4e00' <= ch <= '\u9fff':
            length += 1
        elif ch.isdigit():
            length += 0.5
        else:
            length += 0.5
    return length


# 四号字(14pt)字符宽度（cm）：中文/全角约0.494cm，英文/数字约0.247cm
_FONT_SIZE_PT = 14
_CN_CHAR_WIDTH_CM = _FONT_SIZE_PT * 0.0353   # ≈0.494cm
_EN_CHAR_WIDTH_CM = _FONT_SIZE_PT * 0.0353 / 2  # ≈0.247cm


def estimate_text_width_cm(text):
    """精确估算文本在四号字下的显示宽度（cm）。

    - 中文字符（含中文标点、中文引号）→ 0.494cm
    - 英文字母、数字、ASCII单引号（经纬度）→ 0.247cm
    - 用于判断选项是否会换行，确保排版规则选择正确
    """
    width = 0.0
    for ch in text:
        if ch.isascii() and (ch.isdigit() or ch.isalpha()):
            width += _EN_CHAR_WIDTH_CM
        elif ch == "'":
            width += _EN_CHAR_WIDTH_CM
        else:
            width += _CN_CHAR_WIDTH_CM
    return round(width, 2)


def select_option_rule(options):
    """根据选项实际宽度选择排版规则，只返回 1/2/3（对应1×4/2×2/4×1）。

    规则1（1×4）：每个选项（含"A. "前缀）宽度 ≤ 版心宽度/4，
                  且四选项总宽度+3个Tab间隔 ≤ 版心宽度
    规则2（2×2）：每个选项宽度 ≤ 8.0cm
                  （确保A在第一个制表位8.98cm前结束，B从9.0cm开始后在版心内）
    规则3（4×1）：其他情况（每行一个）

    确保不会出现"一行两个+一行一个"等不规则排列。
    """
    page_w = PAGE_CONTENT_WIDTH_CM  # 17.2cm
    # Tab间隔宽度约0.5cm（制表位间距）
    tab_gap = 0.5

    # 计算每个选项含前缀的宽度
    opt_widths = {}
    for letter, text in options.items():
        full_text = f'{letter}. {text}'
        # 去除图片占位符的影响（图片选项走规则5，不在此处理）
        clean = IMAGE_PLACEHOLDER_PATTERN.sub('', full_text)
        opt_widths[letter] = estimate_text_width_cm(clean)

    max_w = max(opt_widths.values()) if opt_widths else 0

    # 规则1判定：单选项不超 版心/4，且总宽度不超版心
    rule1_per_opt = page_w / 4  # ≈4.3cm
    total_w = sum(opt_widths.values()) + tab_gap * 3
    if max_w <= rule1_per_opt and total_w <= page_w:
        return 1

    # 规则2判定：单选项不超 8.0cm
    # 制表位在8.98和9.0cm，A从悬挂缩进0.54cm开始，
    # 需 A宽度 ≤ 8.98-0.54=8.44cm；B从9.0cm开始，需 B宽度 ≤ 17.2-9.0=8.2cm
    # 取较小值8.2，留余量用8.0
    rule2_max = 8.0
    if max_w <= rule2_max:
        return 2

    # 否则规则3：每行一个
    return 3


def format_options(doc, options, sub_options, images_dir, logger):
    """根据选项排版规则排版选择题选项。"""
    if not options:
        return
    
    # 处理子选项（①②③④）
    if sub_options:
        _format_sub_options(doc, sub_options, logger)
    
    # 检查是否包含图片
    has_images = any(IMAGE_PLACEHOLDER_PATTERN.search(v) for v in options.values())
    
    # 选择排版规则
    if has_images:
        rule = 5
    else:
        rule = select_option_rule(options)
    
    logger.debug(f'  选项排版: 规则{rule}')
    
    if rule == 1:
        _format_options_rule1(doc, options, images_dir, logger)
    elif rule == 2:
        _format_options_rule2(doc, options, images_dir, logger)
    elif rule == 3:
        _format_options_rule3(doc, options, images_dir, logger)
    elif rule == 5:
        _format_options_rule5(doc, options, images_dir, logger)


def _format_sub_options(doc, sub_options, logger):
    """排版①②③④子选项。按精确宽度估算分行，避免换行。
    
    每两个子选项之间用2个空格分隔。
    如果一行能放下则放置一行，放不下则自动分行（一行两个或三个）。
    """
    # 按内容宽度自动分行
    current_line = []
    current_width = 0.0
    gap_cm = 0.5  # 2个空格在四号字下的宽度约0.5cm
    max_width = PAGE_CONTENT_WIDTH_CM

    for sub in sub_options:
        sub_w = estimate_text_width_cm(sub)
        if current_line and current_width + gap_cm + sub_w > max_width:
            # 输出当前行
            p = doc.add_paragraph()
            apply_style(p, '选项')
            add_run_with_text(p, '  '.join(current_line))
            clear_run_fonts(p)
            current_line = [sub]
            current_width = sub_w
        else:
            if current_line:
                current_width += gap_cm
            current_line.append(sub)
            current_width += sub_w
    
    # 输出最后一行
    if current_line:
        p = doc.add_paragraph()
        apply_style(p, '选项')
        add_run_with_text(p, '  '.join(current_line))
        clear_run_fonts(p)


def _format_options_rule1(doc, options, images_dir, logger):
    """规则1：短选项，四个一行，Tab分隔。使用"选项"样式（含悬挂缩进+制表位）。"""
    p = doc.add_paragraph()
    apply_style(p, '选项')
    set_tab_stops(p, TAB_STOPS_4)
    
    letters = sorted(options.keys())
    for i, letter in enumerate(letters):
        if i > 0:
            p.add_run('\t')
        text = options[letter]
        _add_option_text(p, letter, text, images_dir)
    clear_run_fonts(p)


def _format_options_rule2(doc, options, images_dir, logger):
    """规则2：中等选项，AB一行，CD一行，两个Tab分隔。使用"选项"样式。"""
    letters = sorted(options.keys())
    
    # AB行
    p1 = doc.add_paragraph()
    apply_style(p1, '选项')
    set_tab_stops(p1, TAB_STOPS_2)
    if 'A' in options:
        _add_option_text(p1, 'A', options['A'], images_dir)
    if 'B' in options:
        p1.add_run('\t\t')  # 两个Tab分隔
        _add_option_text(p1, 'B', options['B'], images_dir)
    clear_run_fonts(p1)
    
    # CD行
    p2 = doc.add_paragraph()
    apply_style(p2, '选项')
    set_tab_stops(p2, TAB_STOPS_2)
    if 'C' in options:
        _add_option_text(p2, 'C', options['C'], images_dir)
    if 'D' in options:
        p2.add_run('\t\t')  # 两个Tab分隔
        _add_option_text(p2, 'D', options['D'], images_dir)
    clear_run_fonts(p2)


def _format_options_rule3(doc, options, images_dir, logger):
    """规则3：长选项，每个独占一行。使用"选项"样式。"""
    letters = sorted(options.keys())
    for letter in letters:
        p = doc.add_paragraph()
        apply_style(p, '选项')
        _add_option_text(p, letter, options[letter], images_dir)
        clear_run_fonts(p)


def _format_options_rule5(doc, options, images_dir, logger):
    """规则5：图片选项，2x2排列，等比缩放。
    
    布局：
        A. [图片]    B. [图片]
        C. [图片]    D. [图片]
    
    所有图片缩放到相同宽度（版心宽度的40%，约6.9cm），
    确保视觉一致性。
    """
    letters = sorted(options.keys())
    # 图片统一宽度：版心宽度的40%
    img_width = Cm(6.9)
    
    # AB行
    p1 = doc.add_paragraph()
    apply_style(p1, '选项')
    set_tab_stops(p1, TAB_STOPS_2)
    if 'A' in options:
        _add_option_text(p1, 'A', options['A'], images_dir, img_width)
    if 'B' in options:
        p1.add_run('\t')
        _add_option_text(p1, 'B', options['B'], images_dir, img_width)
    
    # CD行
    p2 = doc.add_paragraph()
    apply_style(p2, '选项')
    set_tab_stops(p2, TAB_STOPS_2)
    if 'C' in options:
        _add_option_text(p2, 'C', options['C'], images_dir, img_width)
    if 'D' in options:
        p2.add_run('\t')
        _add_option_text(p2, 'D', options['D'], images_dir, img_width)
    
    logger.debug(f'  图片选项: 2x2排列, 图片宽度={img_width.cm:.1f}cm')


def _add_option_text(paragraph, letter, text, images_dir, img_width=None):
    """在段落中添加选项文字，处理图片占位符。
    
    Args:
        paragraph: 段落对象
        letter: 选项字母（A/B/C/D）
        text: 选项文本（可能含 {{IMAGE:filename}} 占位符）
        images_dir: 图片目录
        img_width: 图片宽度（可选，默认6cm）
    """
    paragraph.add_run(f'{letter}. ')
    
    if img_width is None:
        img_width = Cm(6)
    
    # 检查是否包含图片占位符
    parts = IMAGE_PLACEHOLDER_PATTERN.split(text)
    for i, part in enumerate(parts):
        if i % 2 == 0:
            # 普通文本
            if part:
                paragraph.add_run(part)
        else:
            # 图片占位符（part 是文件名）
            img_path = os.path.join(images_dir, part)
            if os.path.exists(img_path):
                run = paragraph.add_run()
                run.add_picture(img_path, width=img_width)
            else:
                paragraph.add_run(f'[图片:{part}]')


# =====================================================================
# 主排版逻辑
# =====================================================================

def format_docx(json_path, template_path, output_path, images_dir=None, log_path=None):
    """执行试卷排版。"""
    json_path = os.path.abspath(json_path)
    template_path = os.path.abspath(template_path)
    output_path = os.path.abspath(output_path)

    if not os.path.exists(json_path):
        print(f'错误: JSON 文件不存在: {json_path}')
        return False

    if not os.path.exists(template_path):
        print(f'错误: 模板文件不存在: {template_path}')
        return False

    if images_dir is None:
        images_dir = os.path.join(os.path.dirname(json_path), 'images')

    if log_path is None:
        log_path = os.path.join(os.path.dirname(output_path), 'format_log.txt')

    logger = setup_logger(log_path)

    logger.info('=' * 60)
    logger.info('地理试卷排版开始')
    logger.info(f'JSON 文件: {json_path}')
    logger.info(f'模板文件: {template_path}')
    logger.info(f'图片目录: {images_dir}')
    logger.info(f'输出文件: {output_path}')
    logger.info('=' * 60)

    try:
        # 加载 JSON
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # 质检报告数据
        quality_data = {
            'exam_name': data.get('exam_info', {}).get('exam_name', ''),
            'sections': 0,
            'question_groups': 0,
            'total_questions': 0,
            'choice_questions': 0,
            'non_choice_questions': 0,
            'images_inserted': 0,
            'tables_inserted': 0,
            'option_rules': {},
            'missing_images': [],
            'fill_in_blank_count': 0,
            'warnings': [],
        }

        # 加载模板
        logger.info('加载模板...')
        doc = load_template(template_path)
        logger.info('模板加载成功')

        # ========== 排版考试信息 ==========
        logger.info('')
        logger.info('--- 排版考试信息 ---')
        _format_exam_info(doc, data.get('exam_info', {}), logger)

        # ========== 排版考前内容（注意事项等） ==========
        _format_pre_exam_content(doc, data.get('pre_exam_content', []), logger)

        # ========== 排版各分区 ==========
        for section in data.get('sections', []):
            quality_data['sections'] += 1
            logger.info(f'')
            logger.info(f'--- 排版分区 {section.get("section_id", "?")}: {section.get("section_type", "?")} ---')
            _format_section(doc, section, section.get('section_type', ''), images_dir, logger, quality_data)

        # ========== 保存 ==========
        logger.info('')
        logger.info('--- 保存文档 ---')
        doc.save(output_path)
        logger.info(f'文档已保存: {output_path}')

        # ========== 生成质检报告 ==========
        _generate_quality_report(quality_data, output_path, logger)

        logger.info('')
        logger.info('=' * 60)
        logger.info('排版完成！')
        logger.info('=' * 60)

        return True

    except Exception as e:
        logger.error(f'排版过程中出错: {e}', exc_info=True)
        print(f'错误: {e}')
        return False


def _format_exam_info(doc, exam_info, logger):
    """排版考试名称和科目名称。"""
    exam_name = exam_info.get('exam_name', '')
    subject = exam_info.get('subject', '地理')

    # 考试名称
    if exam_name:
        p = doc.add_paragraph()
        apply_style(p, '考试名称')
        add_run_with_text(p, exam_name)
        logger.info(f'  考试名称: {exam_name}')

    # 科目名称（两字科目加空格）
    if subject == '地理':
        subject_display = '地  理'
    else:
        subject_display = subject
    p = doc.add_paragraph()
    apply_style(p, '科目名称')
    add_run_with_text(p, subject_display)
    logger.info(f'  科目名称: {subject_display}')


def _format_pre_exam_content(doc, pre_exam_content, logger):
    """排版考前内容（注意事项、考生信息栏等）。
    
    在科目名称之后、分区标题之前输出。
    注意事项使用"注意事项内容"样式（模板中已有定义）。
    """
    if not pre_exam_content:
        return

    for item in pre_exam_content:
        item_type = item.get('type', '')
        text = item.get('text', '')
        if not text:
            continue

        p = doc.add_paragraph()
        # 注意事项标题使用"题型标题"样式（黑体加粗），条目使用"注意事项内容"样式（楷体）
        if item_type == 'notice_title':
            apply_style(p, '题型标题', logger)
            add_run_with_text(p, text)
        elif item_type == 'notice_item':
            apply_style(p, '注意事项内容', logger)
            add_text_mixed_fonts(p, text, cn_font='楷体', en_font='Times New Roman')
        else:
            apply_style(p, 'Body Text', logger)
            add_text_mixed_fonts(p, text, cn_font='宋体', en_font='Times New Roman')
        logger.debug(f'  考前内容({item_type}): {text[:40]}')


def _format_section(doc, section, section_type, images_dir, logger, quality_data):
    """排版一个分区。"""
    section_title = section.get('section_title', '')

    # 分区标题
    if section_title:
        p = doc.add_paragraph()
        apply_style(p, '题型标题')
        add_run_with_text(p, section_title)
        logger.info(f'  分区标题: {section_title[:40]}')

    # 排版题组
    for group in section.get('question_groups', []):
        quality_data['question_groups'] += 1
        _format_question_group(doc, group, section_type, images_dir, logger, quality_data)


def _format_question_group(doc, group, section_type, images_dir, logger, quality_data):
    """排版一个题组。

    选择题顺序：材料（引导语紧跟末尾）→ 题干 + 选项
    非选择题顺序：题干 → 材料 → 子问题
    """
    group_id = group.get('group_id', '')
    logger.info(f'  题组 {group_id}:')

    instruction = group.get('instruction', '')
    materials = group.get('materials', [])
    questions = group.get('questions', [])

    if section_type == '非选择题':
        # 非选择题：题干 → 材料 → 子问题
        # 1. 先输出题干（如"16. 阅读图文材料，完成下列要求。"）
        for question in questions:
            _format_question_stem(doc, question, logger, quality_data)
        # 2. 再输出材料（材料紧跟题干）
        for material in materials:
            _format_material(doc, material, group, images_dir, logger, quality_data)
        # 3. 最后输出子问题
        for question in questions:
            _format_sub_questions(doc, question, logger, quality_data)
    else:
        # 选择题：材料（引导语紧跟末尾，宋体）→ 题干 + 选项
        last_text_para = None
        for material in materials:
            last_text_para = _format_material(doc, material, group, images_dir, logger, quality_data)
        # 引导语紧跟最后一个材料文字段落末尾（宋体）
        if instruction:
            if last_text_para is not None:
                run = last_text_para.add_run(instruction)
                _set_run_font(run, cn_font='宋体', en_font='Times New Roman')
                logger.debug(f'    引导语(紧跟材料): {instruction}')
            else:
                # 无材料文字时，引导语单独成段（宋体）
                p = doc.add_paragraph()
                apply_style(p, 'Body Text')
                run = p.add_run(instruction)
                _set_run_font(run, cn_font='宋体', en_font='Times New Roman')
                logger.debug(f'    引导语: {instruction}')
        # 题干 + 选项
        for question in questions:
            _format_question(doc, question, section_type, images_dir, logger, quality_data)


def _format_material(doc, material, group, images_dir, logger, quality_data, instruction=''):
    """排版材料（按 segments 数组顺序处理）。

    segments 数组支持：text、image、table 三种类型。
    - text: 文字段落（楷体，带首行缩进）
    - image: 图片段落（居中，使用记录的尺寸或自动缩放）
    - table: 表格（黑体表头+宋体内容+居中+单倍行距）
    
    返回最后一个文字段落（便于追加引导语）。
    """
    text_para = None
    segments = material.get('segments', [])
    
    if not segments:
        # 兼容旧格式：如果 segments 为空但有 text/images/tables 字段
        # 使用旧逻辑处理
        text = material.get('text', '')
        if text:
            lines = text.split('\n')
            for li, line in enumerate(lines):
                if not line:
                    continue
                p = doc.add_paragraph()
                apply_style(p, 'Body Text')
                add_text_mixed_fonts(p, line, cn_font='楷体', en_font='Times New Roman')
                text_para = p
        for img_name in material.get('images', []):
            img_path = os.path.join(images_dir, img_name)
            if os.path.exists(img_path):
                add_image_centered(doc, img_path, logger)
                quality_data['images_inserted'] += 1
            else:
                quality_data['missing_images'].append(img_name)
                logger.warning(f'    图片缺失: {img_name}')
        tables_data = group.get('tables', [])
        for table_ref in material.get('tables', []):
            for td in tables_data:
                if td.get('id') == table_ref:
                    add_table(doc, td, logger)
                    quality_data['tables_inserted'] += 1
                    break
        return text_para
    
    # 预扫描：收集所有 type:"image" 的图片名，用于后续去重
    # 避免图片同时在 text 段（【图片：xxx】占位符）和 image 段各插入一次
    image_seg_names = set()
    for seg in segments:
        if seg.get('type') == 'image':
            name = seg.get('name', '')
            if name:
                image_seg_names.add(name)

    # 按 segments 顺序处理
    last_was_text = False
    for seg in segments:
        seg_type = seg.get('type', '')
        
        if seg_type == 'text':
            content = seg.get('content', '')
            if content:
                # 检查是否包含图片占位符（兼容【图片：xxx】格式）
                img_placeholders = CHINESE_IMAGE_PLACEHOLDER.findall(content)
                if img_placeholders:
                    # 有图片占位符：先输出非占位符文本，再插入图片
                    text_parts = CHINESE_IMAGE_PLACEHOLDER.split(content)
                    has_real_text = any(t.strip() for t in text_parts if t.strip())
                    if has_real_text:
                        p = doc.add_paragraph()
                        apply_style(p, 'Body Text')
                        # 输出非图片文本部分
                        for idx, part in enumerate(text_parts):
                            if part.strip():
                                add_text_mixed_fonts(p, part, cn_font='楷体', en_font='Times New Roman')
                        text_para = p
                        last_was_text = True
                    # 插入占位符对应的图片（跳过已有独立 image segment 的图片，避免重复）
                    for img_name in img_placeholders:
                        if img_name in image_seg_names:
                            logger.debug(f'    跳过重复图片占位符: {img_name}（已有独立 image segment）')
                            continue
                        img_path = os.path.join(images_dir, img_name) if images_dir else None
                        if img_path and os.path.exists(img_path):
                            add_image_centered(doc, img_path, logger)
                            quality_data['images_inserted'] += 1
                            logger.info(f'    占位符图片插入: {img_name}')
                        else:
                            quality_data['missing_images'].append(img_name)
                            logger.warning(f'    占位符图片缺失: {img_name}')
                            # 插入占位符文本作为备用
                            p = doc.add_paragraph()
                            set_paragraph_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)
                            run = p.add_run(f'[图片缺失: {img_name}]')
                            _set_run_font(run, cn_font='宋体', en_font='Times New Roman')
                            run.font.color.rgb = RGBColor(255, 0, 0)
                    last_was_text = False
                else:
                    # 普通文本，无占位符
                    p = doc.add_paragraph()
                    apply_style(p, 'Body Text')
                    add_text_mixed_fonts(p, content, cn_font='楷体', en_font='Times New Roman')
                    text_para = p
                    last_was_text = True
                    logger.debug(f'    材料文字: {content[:50]}...')
        
        elif seg_type == 'image':
            img_name = seg.get('name', '')
            width_cm = seg.get('width_cm')
            height_cm = seg.get('height_cm')
            img_path = os.path.join(images_dir, img_name)
            if os.path.exists(img_path):
                # 使用记录的尺寸或自动计算
                if width_cm and width_cm > 0:
                    # 使用记录的尺寸，限制最大宽度12cm
                    display_width = min(width_cm, 12.0)
                    width = Cm(display_width)
                    p = doc.add_paragraph()
                    set_paragraph_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)
                    apply_style(p, 'Normal')
                    add_image_to_paragraph(p, img_path, width=width)
                    logger.debug(f'    材料图片: {img_name} (原始{width_cm}x{height_cm}cm, 显示{display_width}cm)')
                else:
                    # 无尺寸信息，自动缩放
                    add_image_centered(doc, img_path, logger)
                quality_data['images_inserted'] += 1
            else:
                quality_data['missing_images'].append(img_name)
                logger.warning(f'    图片缺失: {img_name}')
            last_was_text = False
        
        elif seg_type == 'table':
            table_id = seg.get('table_id', '')
            # 从 group 的 tables 字段中查找表格数据
            tables_list = group.get('tables', [])
            table_data = None
            for td in tables_list:
                if td.get('id') == table_id:
                    table_data = td
                    break
            if table_data:
                add_table(doc, table_data, logger, role=table_data.get('role', 'material'))
                quality_data['tables_inserted'] += 1
            last_was_text = False
    
    # 引导语追加到最后一个文字段落
    if instruction and text_para:
        run = text_para.add_run(instruction)
        _set_run_font(run, cn_font='宋体', en_font='Times New Roman')
        logger.debug(f'    引导语(紧跟材料): {instruction}')
    
    return text_para


# 填空题空位模式：3个及以上连续下划线
FILL_IN_BLANK_PATTERN = re.compile(r'_{3,}')


def _format_question_stem(doc, question, logger, quality_data):
    """输出题干段落，并统计题数。返回题干段落。"""
    q_num = question.get('question_number', '?')
    q_type = question.get('question_type', '')
    stem = question.get('stem', '')

    quality_data['total_questions'] += 1
    if q_type == '选择题':
        quality_data['choice_questions'] += 1
    else:
        quality_data['non_choice_questions'] += 1

    logger.info(f'    题{q_num} ({q_type}): {stem[:40]}...')

    # 填空题空位处理：将不一致的下划线统一为6个下划线
    formatted_stem = FILL_IN_BLANK_PATTERN.sub('______', stem)
    if formatted_stem != stem:
        fill_count = len(FILL_IN_BLANK_PATTERN.findall(stem))
        quality_data['fill_in_blank_count'] += fill_count
        logger.debug(f'    填空题: {fill_count} 处空位统一')

    # 选择题括号处理：将空括号（）替换为含两个全角空格的括号（　　）
    # 全角空格 U+3000，用于留出答题书写空间
    formatted_stem = formatted_stem.replace('（）', '（\u3000\u3000）')
    # 兼容半角空括号
    formatted_stem = formatted_stem.replace('()', '（\u3000\u3000）')

    # 题干（混合字体：中文/引号宋体，数字/英文新罗马）
    p = doc.add_paragraph()
    apply_style(p, 'Normal')
    add_text_mixed_fonts(p, f'{q_num}. {formatted_stem}', cn_font='宋体', en_font='Times New Roman')
    return p


def _format_sub_questions(doc, question, logger, quality_data):
    """输出非选择题的子问题。"""
    sub_questions = question.get('sub_questions', [])
    for sq in sub_questions:
        sub_id = sq.get('sub_id', '')
        sub_text = sq.get('text', '')
        # 填空题空位处理
        formatted_sub_text = FILL_IN_BLANK_PATTERN.sub('______', sub_text)
        if formatted_sub_text != sub_text:
            quality_data['fill_in_blank_count'] += 1
        p = doc.add_paragraph()
        apply_style(p, 'Normal')
        if sub_id:
            add_text_mixed_fonts(p, f'{sub_id} {formatted_sub_text}', cn_font='宋体', en_font='Times New Roman')
        else:
            add_text_mixed_fonts(p, formatted_sub_text, cn_font='宋体', en_font='Times New Roman')


def _format_question(doc, question, section_type, images_dir, logger, quality_data):
    """排版选择题：题干 + [stem_images] + [sub_options] + 选项。"""
    _format_question_stem(doc, question, logger, quality_data)

    q_type = question.get('question_type', section_type)
    # 选择题选项
    if q_type == '选择题':
        # 插入题干图片（如有）——位于题干下方、选项上方
        stem_images = question.get('stem_images')
        if stem_images:
            for img_info in stem_images:
                img_name = img_info.get('name', str(img_info))
                img_path = os.path.join(images_dir, img_name) if images_dir else img_name
                if os.path.exists(img_path):
                    add_image_centered(doc, img_path, logger)
                    quality_data['images_inserted'] += 1
                    logger.debug(f'  题干图片: {img_name}')
                else:
                    quality_data['missing_images'].append(img_name)
                    logger.warning(f'  题干图片缺失: {img_name}')

        options = question.get('options')
        sub_options = question.get('sub_options')
        if options:
            # 确定使用的排版规则（与 format_options 保持一致）
            has_images = any(IMAGE_PLACEHOLDER_PATTERN.search(v) for v in options.values())
            if has_images:
                rule_used = 5
            else:
                rule_used = select_option_rule(options)
            quality_data['option_rules'][f'规则{rule_used}'] = \
                quality_data['option_rules'].get(f'规则{rule_used}', 0) + 1

            format_options(doc, options, sub_options, images_dir, logger)


def _generate_quality_report(quality_data, output_path, logger):
    """生成 HTML 质检报告。"""
    report_path = os.path.join(os.path.dirname(output_path), 'quality_report.html')

    has_issues = bool(quality_data['missing_images'] or quality_data['warnings'])
    status_class = 'success' if not has_issues else 'warning'
    status_text = '质检通过' if not has_issues else '发现问题'
    status_desc = '所有检查项通过，未发现异常。' if not has_issues else f'发现 {len(quality_data["missing_images"]) + len(quality_data["warnings"])} 个问题，请检查下方警告信息。'

    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>排版质检报告 - {quality_data["exam_name"]}</title>
<style>
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{
    font-family: "Noto Sans CJK SC", "WenQuanYi Micro Hei", "Microsoft YaHei", sans-serif;
    background: #f0f2f5;
    color: #333;
    line-height: 1.7;
    padding: 20px;
  }}
  .container {{ max-width: 920px; margin: 0 auto; }}
  .report-header {{
    background: linear-gradient(135deg, #1a5276, #2980b9);
    color: #fff;
    padding: 30px 40px;
    border-radius: 12px 12px 0 0;
  }}
  .report-header h1 {{ font-size: 24px; font-weight: 700; margin-bottom: 6px; }}
  .report-header .subtitle {{ font-size: 14px; opacity: 0.85; }}
  .status-badge {{
    display: inline-block;
    padding: 4px 14px;
    border-radius: 20px;
    font-size: 13px;
    font-weight: 600;
    margin-top: 10px;
  }}
  .status-success {{ background: #27ae60; color: #fff; }}
  .status-warning {{ background: #e67e22; color: #fff; }}
  .card {{
    background: #fff;
    border-radius: 0 0 12px 12px;
    padding: 30px 40px;
    box-shadow: 0 2px 8px rgba(0,0,0,0.08);
    margin-bottom: 20px;
  }}
  .card-section {{ margin-bottom: 28px; }}
  .card-section:last-child {{ margin-bottom: 0; }}
  .card-section h2 {{
    font-size: 17px;
    color: #1a5276;
    border-left: 4px solid #2980b9;
    padding-left: 10px;
    margin-bottom: 14px;
    font-weight: 700;
  }}
  .summary-grid {{
    display: grid;
    grid-template-columns: repeat(3, 1fr);
    gap: 12px;
  }}
  .summary-item {{
    background: #f8f9fa;
    border-radius: 8px;
    padding: 14px 16px;
    text-align: center;
  }}
  .summary-item .label {{ font-size: 12px; color: #7f8c8d; margin-bottom: 4px; }}
  .summary-item .value {{ font-size: 22px; font-weight: 700; color: #2c3e50; }}
  .summary-item .unit {{ font-size: 13px; color: #95a5a6; }}
  table {{
    width: 100%;
    border-collapse: collapse;
    font-size: 14px;
  }}
  table th {{
    background: #eaf2f8;
    color: #1a5276;
    font-weight: 600;
    padding: 10px 14px;
    text-align: left;
    border-bottom: 2px solid #d5dbdb;
  }}
  table td {{
    padding: 9px 14px;
    border-bottom: 1px solid #ecf0f1;
  }}
  table tr:last-child td {{ border-bottom: none; }}
  .tag {{
    display: inline-block;
    padding: 2px 10px;
    border-radius: 4px;
    font-size: 12px;
    font-weight: 600;
  }}
  .tag-green {{ background: #e8f8e8; color: #27ae60; }}
  .tag-orange {{ background: #fef5e7; color: #e67e22; }}
  .tag-red {{ background: #fdedec; color: #e74c3c; }}
  .warning-box {{
    background: #fff3cd;
    border: 1px solid #ffeaa7;
    border-left: 4px solid #f39c12;
    border-radius: 8px;
    padding: 14px 18px;
    margin-bottom: 10px;
  }}
  .warning-box .warning-title {{
    font-weight: 700;
    color: #d68910;
    font-size: 14px;
    margin-bottom: 4px;
  }}
  .warning-box .warning-detail {{ font-size: 13px; color: #7d6608; }}
  .warning-box.success {{
    background: #e8f8e8;
    border-color: #abebc6;
    border-left-color: #27ae60;
  }}
  .warning-box.success .warning-title {{ color: #1e8449; }}
  .warning-box.success .warning-detail {{ color: #196f3d; }}
  .footer {{
    text-align: center;
    font-size: 12px;
    color: #bdc3c7;
    padding: 20px 0;
  }}
</style>
</head>
<body>
<div class="container">

  <div class="report-header">
    <h1>排版质检报告</h1>
    <div class="subtitle">{quality_data["exam_name"]}</div>
    <div class="status-badge status-{status_class}">{status_text}</div>
  </div>

  <div class="card">
    <div class="card-section">
      <h2>试卷结构统计</h2>
      <div class="summary-grid">
        <div class="summary-item">
          <div class="label">分区数</div>
          <div class="value">{quality_data["sections"]}</div>
        </div>
        <div class="summary-item">
          <div class="label">题组数</div>
          <div class="value">{quality_data["question_groups"]}<span class="unit"> 组</span></div>
        </div>
        <div class="summary-item">
          <div class="label">总题数</div>
          <div class="value">{quality_data["total_questions"]}</div>
        </div>
        <div class="summary-item">
          <div class="label">选择题</div>
          <div class="value">{quality_data["choice_questions"]}<span class="unit"> 题</span></div>
        </div>
        <div class="summary-item">
          <div class="label">非选择题</div>
          <div class="value">{quality_data["non_choice_questions"]}<span class="unit"> 题</span></div>
        </div>
        <div class="summary-item">
          <div class="label">插入图片</div>
          <div class="value">{quality_data["images_inserted"]}<span class="unit"> 张</span></div>
        </div>
      </div>
    </div>

    <div class="card-section">
      <h2>排版规则使用统计</h2>
      <table>
        <thead>
          <tr><th>规则</th><th>使用次数</th><th>说明</th></tr>
        </thead>
        <tbody>
'''

    rule_descriptions = {
        '规则1': '短选项，4个一行',
        '规则2': '中等选项，AB一行/CD一行',
        '规则3': '长选项，每行1个',
        '规则4': '子选项①②③④',
        '规则5': '图片选项，2×2排列',
    }

    for rule, count in sorted(quality_data['option_rules'].items()):
        desc = rule_descriptions.get(rule, '')
        html += f'          <tr><td>{rule}</td><td>{count} 题</td><td>{desc}</td></tr>\n'

    if not quality_data['option_rules']:
        html += '          <tr><td colspan="3" style="color:#95a5a6;">无选择题或无需排版选项</td></tr>\n'

    html += f'''        </tbody>
      </table>
    </div>

    <div class="card-section">
      <h2>媒体资源统计</h2>
      <table>
        <thead>
          <tr><th>检查项</th><th>结果</th><th>详情</th></tr>
        </thead>
        <tbody>
          <tr>
            <td>插入图片</td>
            <td><span class="tag tag-green">通过</span></td>
            <td>{quality_data["images_inserted"]} 张</td>
          </tr>
          <tr>
            <td>插入表格</td>
            <td><span class="tag tag-green">通过</span></td>
            <td>{quality_data["tables_inserted"]} 个</td>
          </tr>
          <tr>
            <td>填空空位统一</td>
            <td><span class="tag tag-green">通过</span></td>
            <td>{quality_data["fill_in_blank_count"]} 处</td>
          </tr>
        </tbody>
      </table>
    </div>
'''

    if quality_data['missing_images']:
        html += '''
    <div class="card-section">
      <h2>缺失图片</h2>
'''
        for img in quality_data['missing_images']:
            html += f'      <div class="warning-box"><div class="warning-title">图片缺失</div><div class="warning-detail">{img}</div></div>\n'
        html += '    </div>\n'

    if quality_data['warnings']:
        html += '''
    <div class="card-section">
      <h2>警告信息</h2>
'''
        for w in quality_data['warnings']:
            html += f'      <div class="warning-box"><div class="warning-title">警告</div><div class="warning-detail">{w}</div></div>\n'
        html += '    </div>\n'

    html += f'''
    <div class="card-section">
      <h2>质检结论</h2>
      <div class="warning-box {status_class}">
        <div class="warning-title">{status_text}</div>
        <div class="warning-detail">{status_desc}</div>
      </div>
    </div>
  </div>

  <div class="footer">
    地理试卷排版系统 | 生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}
  </div>

</div>
</body>
</html>
'''

    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(html)

    logger.info(f'HTML 质检报告已生成: {report_path}')


def main():
    parser = argparse.ArgumentParser(
        description='地理试卷排版脚本 - 基于 JSON+模板生成排版文档',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例:
  python format_docx.py --json tagged.json --template assets/template.dotx --output formatted.docx
        '''
    )
    parser.add_argument('--json', '-j', required=True, help='打标 JSON 文件路径')
    parser.add_argument('--template', '-t', required=True, help='模板 dotx 文件路径')
    parser.add_argument('--output', '-o', required=True, help='输出 docx 文件路径')
    parser.add_argument('--images-dir', help='图片目录（默认与JSON同目录的images/）')
    parser.add_argument('--log', '-l', help='日志文件路径')

    args = parser.parse_args()

    success = format_docx(args.json, args.template, args.output, args.images_dir, args.log)

    if not success:
        sys.exit(1)


if __name__ == '__main__':
    main()
