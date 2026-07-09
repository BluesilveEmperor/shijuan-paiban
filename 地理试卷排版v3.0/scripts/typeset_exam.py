# -*- coding: utf-8 -*-
"""
地理试卷排版脚本 v3.0 (typeset_exam)

基于 v3.0 统一 Schema (final_exam.json) 和样式模板 (template.dotx)，生成排版好的 Word 文档。

设计原则：
  - 直接消费 v3.0 Schema，不使用任何 v2.0 适配层
  - 图片通过 image_mapping (placeholder_id→image_id→file_name) 三层解析
  - 所有排版逻辑为"纯样式应用"，不理解题目语义
  - 每步输出日志，失败时精确定位

用法:
    python typeset_exam.py --json {工作目录}/试卷数据/final_exam.json \
                           --template assets/template.dotx \
                           --images {工作目录}/清洗产物/images/ \
                           --output {工作目录}/排版文档/final_exam.docx \
                           [--log {工作目录}/排版文档/typeset_log.txt]

输出:
    final_exam.docx       排版后的 Word 文档
    quality_report.html   质检报告 (HTML)
    typeset_log.txt       排版日志
"""

import argparse
import json
import logging
import os
import re
import sys
import zipfile
from datetime import datetime

from docx import Document
from docx.shared import Cm, Emu, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn as docx_qn
from lxml import etree

try:
    from PIL import Image as PILImage
except ImportError:
    PILImage = None

# ============================================================================
# 页面常量
# ============================================================================
PAGE_CONTENT_WIDTH_CM = 21.0 - 1.9 - 1.9       # 版心宽度 17.2cm
PAGE_CONTENT_HEIGHT_CM = 29.7 - 2.54 - 2.54     # 版心高度 24.62cm

# 四号字 (14pt) 字符宽度
_CN_CHAR_WIDTH_CM = 14 * 0.0353                 # 中文 ~0.494cm
_EN_CHAR_WIDTH_CM = 14 * 0.0353 / 2             # 英文 ~0.247cm

# 选项制表位（来自 template.dotx 样式定义）
TAB_STOPS_4 = [Cm(4.54), Cm(8.98), Cm(13.43)]
TAB_STOPS_2 = [Cm(8.98), Cm(9.0)]

# 占位符匹配模式
PLACEHOLDER_TOKEN_PATTERN = re.compile(r'\{\{image:ph_\d{3}\}\}')

# 填空题空位
FILL_IN_BLANK_PATTERN = re.compile(r'_{3,}')

# 中文标点（需强制中文字体）
_CN_PUNCT_CHARS = set(
    '""''<>《》「」『』【】〔〕〖〗〘〙〚〛…—·、。，；：？！'
    '（）()（）【】<>《》「」『』《》〈〉'
    '—–－~～·．・、，；：！？""''""‘’“”'
    '·・×÷±＝≠≤≥≈≡∝∞∑∏∫∮√∛∜∝∠⊥∥'
    '℃°′″‰％‱¤¥€£$₽₩₿'
    '※★☆◆◇▲△▼▽●○◎★☆'
    'ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ'
    '①②③④⑤⑥⑦⑧⑨⑩'
)


# ============================================================================
# 日志
# ============================================================================

def setup_logger(log_path):
    logger = logging.getLogger('geo_exam_typeset')
    logger.setLevel(logging.DEBUG)
    logger.handlers = []

    fh = logging.FileHandler(log_path, mode='w', encoding='utf-8')
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(logging.Formatter('%(asctime)s [%(levelname)s] %(message)s'))
    logger.addHandler(fh)

    sh = logging.StreamHandler()
    sh.setLevel(logging.INFO)
    sh.setFormatter(logging.Formatter('%(message)s'))
    logger.addHandler(sh)

    return logger


# ============================================================================
# 模板加载
# ============================================================================

def load_template(template_path):
    """加载 dotx 模板文件，返回空的 Document（仅保留样式和 sectPr）。"""
    temp_docx = template_path + '.tmp.docx'

    with zipfile.ZipFile(template_path, 'r') as zin:
        with zipfile.ZipFile(temp_docx, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                data = zin.read(item)
                if item == '[Content_Types].xml':
                    data = data.replace(b'template.main+xml', b'document.main+xml')
                zout.writestr(item, data)

    try:
        doc = Document(temp_docx)
        body = doc.element.body
        for child in list(body):
            if child.tag != docx_qn('w:sectPr'):
                body.remove(child)
        return doc
    finally:
        if os.path.exists(temp_docx):
            os.remove(temp_docx)


# ============================================================================
# 样式工具
# ============================================================================

def apply_style(paragraph, style_name, logger=None):
    try:
        paragraph.style = style_name
    except KeyError:
        if logger:
            logger.warning(f'样式不存在: {style_name}, 回退 Normal')
        try:
            paragraph.style = 'Normal'
        except KeyError:
            pass


def set_alignment(paragraph, alignment):
    paragraph.alignment = alignment


# ============================================================================
# 混合字体引擎
# ============================================================================

def _classify_char(ch):
    """分类字符: 'en' 或 'cn'"""
    if ch.isascii() and (ch.isdigit() or ch.isalpha()):
        return 'en'
    if ch == "'":
        return 'en'
    code = ord(ch)
    if 0x2070 <= code <= 0x208F:
        return 'en'
    return 'cn'


def _is_cn_punct(ch):
    return ch in _CN_PUNCT_CHARS


def _set_run_font(run, cn_font=None, en_font=None, size_pt=None, force_cn_all=False):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(docx_qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, docx_qn('w:rFonts'))
    if force_cn_all and cn_font:
        rFonts.set(docx_qn('w:ascii'), cn_font)
        rFonts.set(docx_qn('w:hAnsi'), cn_font)
        rFonts.set(docx_qn('w:eastAsia'), cn_font)
    else:
        if en_font:
            rFonts.set(docx_qn('w:ascii'), en_font)
            rFonts.set(docx_qn('w:hAnsi'), en_font)
        if cn_font:
            rFonts.set(docx_qn('w:eastAsia'), cn_font)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def _parse_subscript_text(text):
    """解析包含 <sub> 标签的文本，返回分段列表 [(text, is_subscript)]。"""
    import re
    parts = []
    pattern = re.compile(r'<sub>(.*?)</sub>', re.DOTALL)
    last_end = 0
    
    for match in pattern.finditer(text):
        start, end = match.span()
        if start > last_end:
            parts.append((text[last_end:start], False))
        parts.append((match.group(1), True))
        last_end = end
    
    if last_end < len(text):
        parts.append((text[last_end:], False))
    
    return parts


def add_mixed_text(paragraph, text, cn_font=None, en_font='Times New Roman', size_pt=None):
    """按字符类型分段，中/西文使用不同字体，支持下标标签。"""
    if not text:
        return
    
    sub_parts = _parse_subscript_text(text)
    
    for sub_text, is_subscript in sub_parts:
        if not sub_text:
            continue
        
        segments = []
        current_type = _classify_char(sub_text[0])
        current_chars = [sub_text[0]]
        for ch in sub_text[1:]:
            t = _classify_char(ch)
            if t == current_type:
                current_chars.append(ch)
            else:
                segments.append((''.join(current_chars), current_type))
                current_chars = [ch]
                current_type = t
        segments.append((''.join(current_chars), current_type))

        for seg_text, seg_type in segments:
            run = paragraph.add_run(seg_text)
            if is_subscript:
                run.font.subscript = True
                if size_pt:
                    run.font.size = Pt(size_pt * 0.7)
            if seg_type == 'en':
                _set_run_font(run, cn_font=cn_font, en_font=en_font, size_pt=size_pt)
            else:
                force = any(_is_cn_punct(c) for c in seg_text)
                _set_run_font(run, cn_font=cn_font, en_font=en_font, size_pt=size_pt, force_cn_all=force)


def clear_run_fonts(paragraph):
    """清除段落中所有 run 的字体覆盖，让样式控制字体。"""
    for run in paragraph.runs:
        rPr = run._element.find(docx_qn('w:rPr'))
        if rPr is not None:
            for tag in ['w:rFonts', 'w:sz', 'w:szCs']:
                for el in rPr.findall(docx_qn(tag)):
                    rPr.remove(el)


# ============================================================================
# 文本宽度估算 (用于选项排版规则选择)
# ============================================================================

def estimate_text_width_cm(text):
    width = 0.0
    for ch in text:
        if ch.isascii() and (ch.isdigit() or ch.isalpha()):
            width += _EN_CHAR_WIDTH_CM
        elif ch == "'":
            width += _EN_CHAR_WIDTH_CM
        else:
            width += _CN_CHAR_WIDTH_CM
    return round(width, 2)


# ============================================================================
# 图片处理
# ============================================================================

def get_image_size_cm(image_path):
    if PILImage is None:
        return None
    try:
        with PILImage.open(image_path) as im:
            px_w, px_h = im.size
            dpi = im.info.get('dpi', (96, 96))
            dpi_x = dpi[0] if dpi[0] else 96
            dpi_y = dpi[1] if dpi[1] else 96
            return px_w / dpi_x * 2.54, px_h / dpi_y * 2.54
    except Exception:
        return None


def compute_display_width(image_path, max_w_cm=12.0, max_h_cm=8.0):
    """等比缩放计算合适的显示宽度（只缩小不放大）。"""
    size = get_image_size_cm(image_path)
    if size is None:
        return Cm(max_w_cm)
    orig_w, orig_h = size
    if orig_w <= 0 or orig_h <= 0:
        return Cm(max_w_cm)
    scale = 1.0
    if orig_w > max_w_cm:
        scale = min(scale, max_w_cm / orig_w)
    if orig_h > max_h_cm:
        scale = min(scale, max_h_cm / orig_h)
    return Cm(orig_w * scale)


def add_picture(paragraph, image_path, width=None):
    if not os.path.exists(image_path):
        run = paragraph.add_run(f'[图片缺失: {os.path.basename(image_path)}]')
        return False
    if width is None:
        width = compute_display_width(image_path)
    run = paragraph.add_run()
    run.add_picture(image_path, width=width)
    return True


def add_centered_picture(doc, image_path, logger):
    p = doc.add_paragraph()
    set_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)
    apply_style(p, 'Normal')
    width = compute_display_width(image_path)
    ok = add_picture(p, image_path, width=width)
    if ok:
        logger.debug(f'  插入图片: {os.path.basename(image_path)}')
    return p, ok


# ============================================================================
# 图片映射解析
# ============================================================================

class ImageResolver:
    """将 placeholder_id 解析为实际图片文件路径。

    解析链路: placeholder_id → image_mapping → image_id → images[].file_name → file_path
    """

    def __init__(self, images_data, image_mapping_data, images_dir):
        self._ph_to_file = {}
        self._unmapped = set()

        # 构建 image_id → file_name 映射
        id_to_file = {}
        for img in images_data:
            img_id = img.get('image_id')
            file_name = img.get('file_name')
            if img_id and file_name:
                id_to_file[img_id] = file_name

        # 构建 placeholder_id → file_path 映射
        for mapping in image_mapping_data:
            ph_id = mapping.get('placeholder_id')
            img_id = mapping.get('image_id')
            if ph_id and img_id:
                file_name = id_to_file.get(img_id)
                if file_name:
                    full_path = os.path.normpath(os.path.join(images_dir, file_name))
                    self._ph_to_file[ph_id] = full_path
                else:
                    self._unmapped.add(ph_id)

    def resolve(self, placeholder_id):
        """返回 (file_path, exists)，不存在时返回 (None, False)"""
        path = self._ph_to_file.get(placeholder_id)
        if path and os.path.exists(path):
            return path, True
        return path, False

    @property
    def unmapped_placeholders(self):
        return self._unmapped


# ============================================================================
# 制表位
# ============================================================================

def set_tab_stops(paragraph, tab_positions):
    pPr = paragraph._element.get_or_add_pPr()
    for existing in pPr.findall(docx_qn('w:tabs')):
        pPr.remove(existing)
    tabs = etree.SubElement(pPr, docx_qn('w:tabs'))
    for pos in tab_positions:
        tab = etree.SubElement(tabs, docx_qn('w:tab'))
        tab.set(docx_qn('w:val'), 'left')
        twips = str(int(pos / 635))
        tab.set(docx_qn('w:pos'), twips)


# ============================================================================
# 表格处理
# ============================================================================

def add_table(doc, table_data, logger, role='material'):
    rows = table_data.get('rows', 0)
    cols = table_data.get('cols', 0)
    data = table_data.get('data', [])
    if rows == 0 or cols == 0 or not data:
        return None

    table = doc.add_table(rows=rows, cols=cols)
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass
    table.autofit = False
    table.allow_autofit = False
    table.width = Cm(PAGE_CONTENT_WIDTH_CM)

    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            if j < cols:
                cell = table.cell(i, j)
                cell.text = ''
                paragraph = cell.paragraphs[0]
                set_alignment(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
                paragraph.paragraph_format.line_spacing = 1.0
                cell_str = str(cell_text)
                if i == 0:
                    add_mixed_text(paragraph, cell_str, cn_font='黑体')
                    for run in paragraph.runs:
                        run.bold = True
                else:
                    add_mixed_text(paragraph, cell_str, cn_font='宋体')

    if rows > 1:
        tbl = table._tbl
        tblPr = tbl.find(docx_qn('w:tblPr'))
        if tblPr is None:
            tblPr = etree.SubElement(tbl, docx_qn('w:tblPr'))
        tblHeader = etree.SubElement(tblPr, docx_qn('w:tblHeader'))
        tblHeader.set(docx_qn('w:val'), 'true')

    logger.debug(f'  表格: {rows}行 x {cols}列')
    return table


# ============================================================================
# 选项排版引擎
# ============================================================================

def _select_option_rule(options):
    """根据选项文本宽度选择排版规则。

    规则 1 (1x4): 四个选项一行，Tab 分隔
    规则 2 (2x2): 两行，每行两个
    规则 3 (4x1): 每行一个
    """
    page_w = PAGE_CONTENT_WIDTH_CM
    tab_gap = 0.5
    opt_widths = {}
    for label, text in options.items():
        clean = PLACEHOLDER_TOKEN_PATTERN.sub('', f'{label}. {text}')
        opt_widths[label] = estimate_text_width_cm(clean)

    max_w = max(opt_widths.values()) if opt_widths else 0
    total_w = sum(opt_widths.values()) + tab_gap * 3

    if max_w <= page_w / 4 and total_w <= page_w:
        return 1
    if max_w <= 8.0:
        return 2
    return 3


def _add_option_label_text(paragraph, label, text, image_resolver, logger, img_width=None):
    """添加选项标签和文本，处理占位符中的图片。"""
    paragraph.add_run(f'{label}. ')

    if img_width is None:
        img_width = Cm(6)

    parts = PLACEHOLDER_TOKEN_PATTERN.split(text)
    for i, part in enumerate(parts):
        if i % 2 == 0:
            if part:
                paragraph.add_run(part)
        else:
            # part 是 placeholder_id (如 ph_001)
            # 但这里需要匹配 token {{image:ph_001}}
            token = f'{{{{image:{part}}}}}'
            # 查找原始占位符
            ph_file_path, exists = image_resolver.resolve(part) if part.startswith('ph_') else (None, False)
            if exists:
                run = paragraph.add_run()
                run.add_picture(ph_file_path, width=img_width)
            else:
                paragraph.add_run(f'[图片:{part}]')


def format_options(doc, options, image_resolver, logger, sub_options=None):
    """选择题选项排版入口。"""
    if not options:
        return

    # v3.0 的 options 是 [{label, text}] 数组，转为 {label: text} 字典
    if isinstance(options, list):
        opt_dict = {}
        for opt in options:
            label = opt.get('label', '')
            text = opt.get('text', '')
            if label:
                opt_dict[label] = text
        options = opt_dict

    if not options:
        return

    # 子选项
    if sub_options:
        _format_sub_options(doc, sub_options, logger)

    # 检查是否含图片占位符
    has_images = any(PLACEHOLDER_TOKEN_PATTERN.search(v) for v in options.values())

    if has_images:
        rule = 5
    else:
        rule = _select_option_rule(options)

    logger.debug(f'  选项排版: 规则{rule}')

    if rule == 1:
        _format_rule_1x4(doc, options, image_resolver, logger)
    elif rule == 2:
        _format_rule_2x2(doc, options, image_resolver, logger)
    elif rule == 3:
        _format_rule_4x1(doc, options, image_resolver, logger)
    elif rule == 5:
        _format_rule_image(doc, options, image_resolver, logger)

    return rule


def _format_sub_options(doc, sub_options, logger):
    """排版子选项（①②③④等）。"""
    current_line = []
    current_width = 0.0
    gap_cm = 0.5
    max_width = PAGE_CONTENT_WIDTH_CM

    for sub in sub_options:
        sub_w = estimate_text_width_cm(sub)
        if current_line and current_width + gap_cm + sub_w > max_width:
            p = doc.add_paragraph()
            apply_style(p, '选项')
            p.add_run('  '.join(current_line))
            clear_run_fonts(p)
            current_line = [sub]
            current_width = sub_w
        else:
            if current_line:
                current_width += gap_cm
            current_line.append(sub)
            current_width += sub_w

    if current_line:
        p = doc.add_paragraph()
        apply_style(p, '选项')
        p.add_run('  '.join(current_line))
        clear_run_fonts(p)


def _format_rule_1x4(doc, options, image_resolver, logger):
    p = doc.add_paragraph()
    apply_style(p, '选项')
    set_tab_stops(p, TAB_STOPS_4)
    letters = sorted(options.keys())
    for i, letter in enumerate(letters):
        if i > 0:
            p.add_run('\t')
        _add_option_label_text(p, letter, options[letter], image_resolver, logger)
    clear_run_fonts(p)


def _format_rule_2x2(doc, options, image_resolver, logger):
    letters = sorted(options.keys())

    p1 = doc.add_paragraph()
    apply_style(p1, '选项')
    set_tab_stops(p1, TAB_STOPS_2)
    if 'A' in options:
        _add_option_label_text(p1, 'A', options['A'], image_resolver, logger)
    if 'B' in options:
        p1.add_run('\t\t')
        _add_option_label_text(p1, 'B', options['B'], image_resolver, logger)
    clear_run_fonts(p1)

    p2 = doc.add_paragraph()
    apply_style(p2, '选项')
    set_tab_stops(p2, TAB_STOPS_2)
    if 'C' in options:
        _add_option_label_text(p2, 'C', options['C'], image_resolver, logger)
    if 'D' in options:
        p2.add_run('\t\t')
        _add_option_label_text(p2, 'D', options['D'], image_resolver, logger)
    clear_run_fonts(p2)


def _format_rule_4x1(doc, options, image_resolver, logger):
    for letter in sorted(options.keys()):
        p = doc.add_paragraph()
        apply_style(p, '选项')
        _add_option_label_text(p, letter, options[letter], image_resolver, logger)
        clear_run_fonts(p)


def _format_rule_image(doc, options, image_resolver, logger):
    letters = sorted(options.keys())
    img_width = Cm(6.9)

    p1 = doc.add_paragraph()
    apply_style(p1, '选项')
    set_tab_stops(p1, TAB_STOPS_2)
    if 'A' in options:
        _add_option_label_text(p1, 'A', options['A'], image_resolver, logger, img_width)
    if 'B' in options:
        p1.add_run('\t')
        _add_option_label_text(p1, 'B', options['B'], image_resolver, logger, img_width)

    p2 = doc.add_paragraph()
    apply_style(p2, '选项')
    set_tab_stops(p2, TAB_STOPS_2)
    if 'C' in options:
        _add_option_label_text(p2, 'C', options['C'], image_resolver, logger, img_width)
    if 'D' in options:
        p2.add_run('\t')
        _add_option_label_text(p2, 'D', options['D'], image_resolver, logger, img_width)

    logger.debug(f'  图片选项: 2x2, 宽度={img_width.cm:.1f}cm')


# ============================================================================
# 题目排版
# ============================================================================

def _format_question_stem(doc, question, image_resolver, logger, quality):
    """排版题干段落，处理填空题空位、括号、图片占位符。"""
    q_num = question.get('number', '?')
    q_type = question.get('question_type', '')
    stem = question.get('stem', '')

    quality['total_questions'] += 1
    if q_type == '选择题':
        quality['choice_questions'] += 1
    else:
        quality['non_choice_questions'] += 1

    logger.info(f'  题{q_num} ({q_type}): {stem[:50]}...')

    # 填空题空位统一
    formatted = FILL_IN_BLANK_PATTERN.sub('______', stem)
    if formatted != stem:
        quality['fill_in_blank_count'] += len(FILL_IN_BLANK_PATTERN.findall(stem))

    # 选择题括号替换
    formatted = formatted.replace('（）', '（\u3000\u3000）')
    formatted = formatted.replace('()', '（\u3000\u3000）')

    # 按占位符分段处理
    parts = PLACEHOLDER_TOKEN_PATTERN.split(formatted)
    # 使用 findall 获取 token
    tokens = PLACEHOLDER_TOKEN_PATTERN.findall(formatted)

    if not tokens:
        p = doc.add_paragraph()
        apply_style(p, 'Normal')
        add_mixed_text(p, f'{q_num}. {formatted}', cn_font='宋体')
        return

    # 有图片占位符：分段处理
    first_text = True
    for i, part in enumerate(parts):
        part = part.strip()
        if part:
            p = doc.add_paragraph()
            apply_style(p, 'Normal')
            prefix = f'{q_num}. ' if first_text else ''
            add_mixed_text(p, f'{prefix}{part}', cn_font='宋体')
            first_text = False

        if i < len(tokens):
            # 解析 token: {{image:ph_xxx}}
            token = tokens[i]
            ph_id = token.replace('{{image:', '').replace('}}', '')
            file_path, exists = image_resolver.resolve(ph_id)
            if exists:
                add_centered_picture(doc, file_path, logger)
                quality['images_inserted'] += 1
                logger.info(f'    题干图片: {os.path.basename(file_path)}')
            else:
                quality['missing_images'].append(ph_id)
                logger.warning(f'    图片缺失: placeholder_id={ph_id}')


def _format_subquestions(doc, question, logger, quality):
    """排版非选择题子问题。"""
    subquestions = question.get('subquestions', [])
    for sq in subquestions:
        label = sq.get('label', '')
        stem = sq.get('stem', '')
        formatted = FILL_IN_BLANK_PATTERN.sub('______', stem)
        if formatted != stem:
            quality['fill_in_blank_count'] += 1
        p = doc.add_paragraph()
        apply_style(p, 'Normal')
        add_mixed_text(p, f'{label} {formatted}', cn_font='宋体')


# ============================================================================
# 材料排版
# ============================================================================

def _format_materials(doc, question, image_resolver, logger, quality):
    """排版题目材料（含 segments: text/image/table）。"""
    materials = question.get('materials', [])
    if not materials:
        return

    for material in materials:
        content = material.get('content', '')
        guide_sentence = material.get('guide_sentence', '')
        segments = material.get('segments', [])

        if not segments and content:
            text_parts = PLACEHOLDER_TOKEN_PATTERN.split(content)
            text_tokens = PLACEHOLDER_TOKEN_PATTERN.findall(content)
            
            if text_tokens:
                for k, tp in enumerate(text_parts):
                    tp = tp.strip()
                    if tp:
                        p = doc.add_paragraph()
                        apply_style(p, 'Body Text')
                        add_mixed_text(p, tp, cn_font='楷体')
                    if k < len(text_tokens):
                        ph_id = text_tokens[k].replace('{{image:', '').replace('}}', '')
                        file_path, exists = image_resolver.resolve(ph_id)
                        if exists:
                            add_centered_picture(doc, file_path, logger)
                            quality['images_inserted'] += 1
                            logger.info(f'    材料图片: {os.path.basename(file_path)}')
                        else:
                            quality['missing_images'].append(ph_id)
                            logger.warning(f'    图片缺失: placeholder_id={ph_id}')
                
                if guide_sentence:
                    last_p = doc.paragraphs[-1] if doc.paragraphs else None
                    if last_p:
                        add_mixed_text(last_p, guide_sentence, cn_font='宋体')
                    else:
                        p = doc.add_paragraph()
                        apply_style(p, 'Body Text')
                        add_mixed_text(p, guide_sentence, cn_font='宋体')
            else:
                p = doc.add_paragraph()
                apply_style(p, 'Body Text')
                add_mixed_text(p, content, cn_font='楷体')
                if guide_sentence:
                    add_mixed_text(p, guide_sentence, cn_font='宋体')
                logger.debug(f'  材料: {content[:50]}...')
            continue

        if not segments:
            continue

        # 按 segments 顺序处理
        for seg in segments:
            seg_type = seg.get('type', '')

            if seg_type == 'text':
                text = seg.get('content', '')
                if text:
                    # 处理占位符
                    text_parts = PLACEHOLDER_TOKEN_PATTERN.split(text)
                    text_tokens = PLACEHOLDER_TOKEN_PATTERN.findall(text)
                    if text_tokens:
                        p = doc.add_paragraph()
                        apply_style(p, 'Body Text')
                        for j, tp in enumerate(text_parts):
                            if tp.strip():
                                add_mixed_text(p, tp.strip(), cn_font='楷体')
                            if j < len(text_tokens):
                                ph_id = text_tokens[j].replace('{{image:', '').replace('}}', '')
                                file_path, exists = image_resolver.resolve(ph_id)
                                if exists:
                                    add_centered_picture(doc, file_path, logger)
                                    quality['images_inserted'] += 1
                    else:
                        p = doc.add_paragraph()
                        apply_style(p, 'Body Text')
                        add_mixed_text(p, text, cn_font='楷体')

            elif seg_type == 'image':
                img_name = seg.get('name', '')
                if img_name:
                    # 直接通过文件名查找
                    img_path = os.path.join(image_resolver._ph_to_file.get('', '') or '', img_name)
                    # fallback: 通过 image_mapping 中的 file_name 查找
                    from pathlib import Path
                    resolved = False
                    for ph_id, path in image_resolver._ph_to_file.items():
                        if os.path.basename(path) == img_name:
                            img_path = path
                            resolved = True
                            break
                    if not resolved:
                        # 直接拼接 images_dir
                        for key, val in image_resolver._ph_to_file.items():
                            img_path = os.path.join(os.path.dirname(val), img_name)
                            if os.path.exists(img_path):
                                break
                    if os.path.exists(img_path):
                        add_centered_picture(doc, img_path, logger)
                        quality['images_inserted'] += 1

            elif seg_type == 'table':
                table_data = seg.get('data', seg.get('table_data'))
                if table_data:
                    add_table(doc, table_data, logger)
                    quality['tables_inserted'] += 1


# ============================================================================
# 分区排版
# ============================================================================

def _format_section(doc, section, image_resolver, logger, quality):
    """排版一个分区（选择题/非选择题等）。"""
    sec_type = section.get('type', '')
    sec_id = section.get('id', '')
    sec_title = section.get('title', '')
    instructions = section.get('instructions', [])
    questions = section.get('questions', [])

    logger.info(f'')
    logger.info(f'--- 排版 {sec_id}: {sec_title[:40] if sec_title else sec_type} ---')

    # 分区标题
    if sec_title:
        p = doc.add_paragraph()
        apply_style(p, '题型标题')
        p.add_run(sec_title)

    # 分区说明（跳过与标题重复的内容）
    for instr in instructions:
        instr_stripped = instr.strip()
        if sec_title and instr_stripped in sec_title:
            logger.debug(f'  跳过重复说明: {instr_stripped[:30]}')
            continue
        p = doc.add_paragraph()
        apply_style(p, 'Body Text')
        add_mixed_text(p, instr, cn_font='宋体')

    quality['sections'] += 1

    for question in questions:
        q_type = question.get('question_type', sec_type)

        if q_type == '非选择题':
            # 非选择题：题干 → 材料 → 子问题
            _format_question_stem(doc, question, image_resolver, logger, quality)
            _format_materials(doc, question, image_resolver, logger, quality)
            _format_subquestions(doc, question, logger, quality)
        else:
            # 选择题：材料(如有) → 题干 → 选项
            _format_materials(doc, question, image_resolver, logger, quality)
            _format_question_stem(doc, question, image_resolver, logger, quality)

            if q_type == '选择题':
                options = question.get('options', [])
                if options:
                    rule = format_options(doc, options, image_resolver, logger)
                    if rule:
                        quality['option_rules'][f'规则{rule}'] = \
                            quality['option_rules'].get(f'规则{rule}', 0) + 1


# ============================================================================
# 考前内容排版
# ============================================================================

def _format_meta_notes(doc, meta, logger):
    """排版 meta.notes（注意事项等）。优先使用 notes_items 结构化数据，回退到 notes 纯文本。"""
    notes_items = meta.get('notes_items', [])
    
    if notes_items:
        for item in notes_items:
            item_type = item.get('type', '')
            content = item.get('content', '')
            number = item.get('number', '')
            
            if not content:
                continue
            
            p = doc.add_paragraph()
            if item_type == 'title':
                apply_style(p, '题型标题')
                add_mixed_text(p, content, cn_font='黑体')
                logger.debug(f'  注意事项标题: {content[:40]}')
            elif item_type == 'item':
                apply_style(p, '注意事项内容')
                if number:
                    add_mixed_text(p, f'{number}. {content}', cn_font='楷体')
                else:
                    add_mixed_text(p, content, cn_font='楷体')
                logger.debug(f'  注意事项内容: {content[:40]}')
        return
    
    notes = meta.get('notes', '')
    if not notes:
        return

    lines = notes.strip().split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('注意事项') or line.startswith('注意：'):
            p = doc.add_paragraph()
            apply_style(p, '题型标题')
            add_mixed_text(p, line, cn_font='黑体')
            logger.debug(f'  注意事项标题: {line[:40]}')
        else:
            p = doc.add_paragraph()
            apply_style(p, '注意事项内容')
            add_mixed_text(p, line, cn_font='楷体')
            logger.debug(f'  注意事项内容: {line[:40]}')


# ============================================================================
# 考试信息排版
# ============================================================================

def _format_exam_header(doc, meta, logger):
    """排版考试名称和科目。"""
    title = meta.get('title', '')
    subtitle = meta.get('subtitle', '')
    subject = meta.get('subject', '')

    if title:
        p = doc.add_paragraph()
        apply_style(p, '考试名称')
        add_mixed_text(p, title, cn_font='黑体')
        logger.info(f'  考试名称: {title}')

    if subtitle:
        p = doc.add_paragraph()
        apply_style(p, '考试名称')
        add_mixed_text(p, subtitle, cn_font='黑体')
        logger.info(f'  副标题: {subtitle}')

    if subject:
        subject_display = '地  理' if subject == '地理' else subject
        p = doc.add_paragraph()
        apply_style(p, '科目名称')
        p.add_run(subject_display)
        logger.info(f'  科目: {subject_display}')


# ============================================================================
# 质检报告
# ============================================================================

def _generate_report(quality, output_path, logger):
    report_path = os.path.join(os.path.dirname(output_path), 'quality_report.html')
    has_issues = bool(quality['missing_images'] or quality['warnings'])
    status_class = 'success' if not has_issues else 'warning'
    status_text = '质检通过' if not has_issues else '发现问题'
    issue_count = len(quality['missing_images']) + len(quality['warnings'])
    status_desc = '所有检查项通过。' if not has_issues else f'发现 {issue_count} 个问题。'

    rule_descriptions = {
        '规则1': '短选项 1x4',
        '规则2': '中等选项 2x2',
        '规则3': '长选项 4x1',
        '规则5': '图片选项 2x2',
    }

    option_rules_html = ''
    for rule, count in sorted(quality['option_rules'].items()):
        desc = rule_descriptions.get(rule, '')
        option_rules_html += f'<tr><td>{rule}</td><td>{count} 题</td><td>{desc}</td></tr>\n'
    if not quality['option_rules']:
        option_rules_html += '<tr><td colspan="3" style="color:#999;">无</td></tr>\n'

    missing_html = ''
    for img in quality['missing_images']:
        missing_html += f'<div class="warn"><div class="warn-title">图片缺失</div><div class="warn-detail">{img}</div></div>\n'

    warn_html = ''
    for w in quality['warnings']:
        warn_html += f'<div class="warn"><div class="warn-title">警告</div><div class="warn-detail">{w}</div></div>\n'

    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>排版质检报告 - {quality['exam_name']}</title>
<style>
*{{margin:0;padding:0;box-sizing:border-box}}
body{{font-family:"Microsoft YaHei",sans-serif;background:#f0f2f5;color:#333;line-height:1.7;padding:20px}}
.container{{max-width:920px;margin:0 auto}}
.header{{background:linear-gradient(135deg,#1a5276,#2980b9);color:#fff;padding:30px 40px;border-radius:12px 12px 0 0}}
.header h1{{font-size:24px;margin-bottom:6px}}
.badge{{display:inline-block;padding:4px 14px;border-radius:20px;font-size:13px;font-weight:600;margin-top:10px}}
.badge-ok{{background:#27ae60}}
.badge-warn{{background:#e67e22}}
.card{{background:#fff;border-radius:0 0 12px 12px;padding:30px 40px;box-shadow:0 2px 8px rgba(0,0,0,.08);margin-bottom:20px}}
.card h2{{font-size:17px;color:#1a5276;border-left:4px solid #2980b9;padding-left:10px;margin:0 0 14px 0}}
.grid{{display:grid;grid-template-columns:repeat(3,1fr);gap:12px}}
.grid-item{{background:#f8f9fa;border-radius:8px;padding:14px 16px;text-align:center}}
.grid-item .lbl{{font-size:12px;color:#7f8c8d}}
.grid-item .val{{font-size:22px;font-weight:700;color:#2c3e50}}
table{{width:100%;border-collapse:collapse;font-size:14px}}
th{{background:#eaf2f8;color:#1a5276;font-weight:600;padding:10px 14px;text-align:left;border-bottom:2px solid #d5dbdb}}
td{{padding:9px 14px;border-bottom:1px solid #ecf0f1}}
.tag{{display:inline-block;padding:2px 10px;border-radius:4px;font-size:12px;font-weight:600}}
.tag-ok{{background:#e8f8e8;color:#27ae60}}
.warn{{background:#fff3cd;border:1px solid #ffeaa7;border-left:4px solid #f39c12;border-radius:8px;padding:14px 18px;margin-bottom:10px}}
.warn-title{{font-weight:700;color:#d68910;font-size:14px;margin-bottom:4px}}
.warn-detail{{font-size:13px;color:#7d6608}}
.warn.ok{{background:#e8f8e8;border-color:#abebc6;border-left-color:#27ae60}}
.warn.ok .warn-title{{color:#1e8449}}
.warn.ok .warn-detail{{color:#196f3d}}
.footer{{text-align:center;font-size:12px;color:#bdc3c7;padding:20px 0}}
</style>
</head>
<body>
<div class="container">
<div class="header">
<h1>排版质检报告</h1>
<div>{quality['exam_name']}</div>
<div class="badge {'badge-ok' if status_class == 'success' else 'badge-warn'}">{status_text}</div>
</div>
<div class="card">
<h2>结构统计</h2>
<div class="grid">
<div class="grid-item"><div class="lbl">分区</div><div class="val">{quality['sections']}</div></div>
<div class="grid-item"><div class="lbl">总题数</div><div class="val">{quality['total_questions']}</div></div>
<div class="grid-item"><div class="lbl">选择题</div><div class="val">{quality['choice_questions']}</div></div>
<div class="grid-item"><div class="lbl">非选择题</div><div class="val">{quality['non_choice_questions']}</div></div>
<div class="grid-item"><div class="lbl">图片</div><div class="val">{quality['images_inserted']}</div></div>
<div class="grid-item"><div class="lbl">表格</div><div class="val">{quality['tables_inserted']}</div></div>
</div>
</div>
<div class="card-section">
<h2>选项排版规则</h2>
<table><thead><tr><th>规则</th><th>次数</th><th>说明</th></tr></thead><tbody>{option_rules_html}</tbody></table>
</div>
<div class="card-section">
<h2>资源统计</h2>
<table>
<tr><td>插入图片</td><td><span class="tag tag-ok">OK</span></td><td>{quality['images_inserted']} 张</td></tr>
<tr><td>插入表格</td><td><span class="tag tag-ok">OK</span></td><td>{quality['tables_inserted']} 个</td></tr>
<tr><td>填空空位</td><td><span class="tag tag-ok">OK</span></td><td>{quality['fill_in_blank_count']} 处</td></tr>
</table>
</div>
{missing_html}
{warn_html}
<div class="card-section">
<h2>结论</h2>
<div class="warn {'ok' if status_class == 'success' else ''}"><div class="warn-title">{status_text}</div><div class="warn-detail">{status_desc}</div></div>
</div>
</div>
<div class="footer">地理试卷排版 v3.0 | {datetime.now().strftime('%Y-%m-%d %H:%M')}</div>
</div>
</body>
</html>'''

    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(html)
    logger.info(f'质检报告: {report_path}')


# ============================================================================
# 主排版流程
# ============================================================================

def typeset_exam(json_path, template_path, output_path, images_dir=None, log_path=None):
    """主排版入口函数。

    Args:
        json_path:   final_exam.json (v3.0 Schema)
        template_path: template.dotx
        output_path:  输出 docx 路径
        images_dir:   图片目录
        log_path:     日志路径

    Returns:
        bool: 排版是否成功
    """
    json_path = os.path.abspath(json_path)
    template_path = os.path.abspath(template_path)
    output_path = os.path.abspath(output_path)

    # 输入校验
    errors = []
    if not os.path.exists(json_path):
        errors.append(f'JSON 文件不存在: {json_path}')
    if not os.path.exists(template_path):
        errors.append(f'模板文件不存在: {template_path}')
    if errors:
        for e in errors:
            print(f'错误: {e}', file=sys.stderr)
        return False

    if images_dir is None:
        images_dir = os.path.join(os.path.dirname(json_path), '..', '清洗产物', 'images')
    images_dir = os.path.abspath(images_dir)

    if log_path is None:
        log_path = os.path.join(os.path.dirname(output_path), 'typeset_log.txt')

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    logger = setup_logger(log_path)

    logger.info('=' * 60)
    logger.info('地理试卷排版 v3.0 (typeset_exam)')
    logger.info(f'JSON:   {json_path}')
    logger.info(f'模板:   {template_path}')
    logger.info(f'图片:   {images_dir}')
    logger.info(f'输出:   {output_path}')
    logger.info('=' * 60)

    try:
        # 1. 加载 JSON
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        meta = data.get('meta', {})
        document = data.get('document', {})
        images_data = data.get('images', [])
        image_mapping = data.get('image_mapping', [])
        validation = data.get('validation', {})

        # 2. 构建图片映射解析器
        image_resolver = ImageResolver(images_data, image_mapping, images_dir)

        # 3. 初始质检数据
        quality = {
            'exam_name': meta.get('title', ''),
            'sections': 0,
            'total_questions': 0,
            'choice_questions': 0,
            'non_choice_questions': 0,
            'images_inserted': 0,
            'tables_inserted': 0,
            'option_rules': {},
            'missing_images': [],
            'fill_in_blank_count': 0,
            'warnings': [],
        }

        # 注入校验警告
        if validation.get('has_unmapped_placeholders'):
            quality['warnings'].append('存在未映射的占位符：' + ', '.join(
                validation.get('unmapped_placeholders', [])))
        if validation.get('has_unused_images'):
            quality['warnings'].append('存在未使用的图片：' + ', '.join(
                validation.get('unused_images', [])))
        for w in validation.get('warnings', []):
            if w not in quality['warnings']:
                quality['warnings'].append(w)

        # 4. 加载模板
        logger.info('加载模板...')
        doc = load_template(template_path)
        logger.info('模板加载成功')

        # 5. 排版：考试信息
        logger.info('')
        logger.info('--- 考试信息 ---')
        _format_exam_header(doc, meta, logger)

        # 6. 排版：注意事项
        _format_meta_notes(doc, meta, logger)

        # 7. 排版：各分区
        sections = document.get('sections', [])
        for section in sections:
            _format_section(doc, section, image_resolver, logger, quality)

        # 8. 未归类块（如有）
        unclassified = document.get('unclassified_blocks', [])
        if unclassified:
            logger.warning(f'存在 {len(unclassified)} 个未归类块，未排版')
            quality['warnings'].append(f'存在 {len(unclassified)} 个未归类文本块')

        # 9. 保存
        logger.info('')
        logger.info('--- 保存文档 ---')
        doc.save(output_path)
        logger.info(f'文档已保存: {output_path}')

        # 10. 质检报告
        _generate_report(quality, output_path, logger)

        # 打印摘要
        logger.info('')
        logger.info('=' * 60)
        logger.info(f'排版完成!')
        logger.info(f'  分区: {quality["sections"]}  题目: {quality["total_questions"]}')
        logger.info(f'  选择题: {quality["choice_questions"]}  非选择题: {quality["non_choice_questions"]}')
        logger.info(f'  图片: {quality["images_inserted"]}  表格: {quality["tables_inserted"]}')
        logger.info(f'  缺失: {len(quality["missing_images"])}  警告: {len(quality["warnings"])}')
        logger.info('=' * 60)

        return True

    except json.JSONDecodeError as e:
        logger.error(f'JSON 解析失败: {e}')
        return False
    except Exception as e:
        logger.error(f'排版失败: {e}', exc_info=True)
        return False


# ============================================================================
# CLI
# ============================================================================

def main():
    parser = argparse.ArgumentParser(
        description='地理试卷排版 v3.0 - 基于 final_exam.json + template.dotx 生成排版文档',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例:
  python typeset_exam.py --json output/final_exam.json \\
                         --template assets/template.dotx \\
                         --images {工作目录}/清洗产物/images/ \
                         --output dist/final_exam.docx
        '''
    )
    parser.add_argument('--json', '-j', required=True, help='final_exam.json 路径 (v3.0 Schema)')
    parser.add_argument('--template', '-t', required=True, help='template.dotx 路径')
    parser.add_argument('--output', '-o', required=True, help='输出 docx 路径')
    parser.add_argument('--images', '-i', help='图片目录 (默认 {工作目录}/清洗产物/images/)')
    parser.add_argument('--log', '-l', help='日志文件路径')

    args = parser.parse_args()

    success = typeset_exam(
        json_path=args.json,
        template_path=args.template,
        output_path=args.output,
        images_dir=args.images,
        log_path=args.log,
    )

    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
