# -*- coding: utf-8 -*-
"""诊断测试：生成含三种规则的选项排版，检查 XML 制表位"""
import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'scripts'))

from docx import Document
from typeset_exam import (
    format_options, set_tab_stops, TAB_STOPS_4, TAB_STOPS_2,
    apply_style, add_mixed_text, clear_run_fonts,
    _select_option_rule, measure_text_width_cm,
)
from docx.oxml.ns import qn as docx_qn

# 创建文档
doc = Document()

# 模拟 image_resolver (空实现)
class FakeResolver:
    def resolve(self, ph_id):
        return None, False

import logging
logger = logging.getLogger('test')
logger.setLevel(logging.DEBUG)
handler = logging.StreamHandler()
handler.setFormatter(logging.Formatter('%(message)s'))
logger.addHandler(handler)

print('=== 测试1: 短选项(1x4) ===')
opts1 = {'A': '北京', 'B': '上海', 'C': '广州', 'D': '深圳'}
rule1 = format_options(doc, opts1, FakeResolver(), logger)
print(f'  规则: {rule1}')

print()
print('=== 测试2: A/C长B/D短(2x2) ===')
opts2 = {'A': '这是一个中等长度的选项文字', 'B': '短选项B', 'C': '另一个中等长度选项C', 'D': '短选项D'}
rule2 = format_options(doc, opts2, FakeResolver(), logger)
print(f'  规则: {rule2}')

print()
print('=== 测试3: 超长选项(4x1) ===')
opts3 = {'A': '这是一个非常非常长的选项文字内容用于测试降级到一行一个的情况', 'B': '另一个也很长的选项文字内容', 'C': '第三个长选项文字内容', 'D': '第四个长选项文字'}
rule3 = format_options(doc, opts3, FakeResolver(), logger)
print(f'  规则: {rule3}')

# 保存
out_path = os.path.join(os.path.dirname(__file__), '_test_options.docx')
doc.save(out_path)
print(f'\n已保存: {out_path}')

# 检查 XML
print('\n=== XML 制表位检查 ===')
doc2 = Document(out_path)
for i, p in enumerate(doc2.paragraphs):
    pPr = p._element.find(docx_qn('w:pPr'))
    if pPr is not None:
        tabs_elem = pPr.find(docx_qn('w:tabs'))
        text_preview = p.text[:40] if p.text else '(空)'
        if tabs_elem is not None:
            tab_positions = []
            for tab in tabs_elem.findall(docx_qn('w:tab')):
                pos = tab.get(docx_qn('w:pos'))
                val = tab.get(docx_qn('w:val'))
                pos_cm = int(pos) * 635 / 360000  # twips → cm
                tab_positions.append(f'{pos_cm:.2f}cm({val})')
            print(f'  段落{i}: 制表位=[{", ".join(tab_positions)}] | {text_preview}')
        else:
            # 检查是否有 tab 字符
            has_tab = '\t' in p.text
            tab_info = ' [含Tab字符!]' if has_tab else ''
            print(f'  段落{i}: 无制表位{tab_info} | {text_preview}')
