# -*- coding: utf-8 -*-
"""
地理试卷排版脚本 v3.0 (typeset_exam)

基于 v3.0 统一 Schema (final_exam.json) 和样式模板 (template.dotx)，生成排版好的 Word 文档。

设计原则：
  - 直接消费 v3.0 Schema，不使用任何 v2.0 适配层
  - 图片通过 image_mapping (placeholder_id→image_id→file_name) 三层解析
  - 所有排版逻辑为"纯样式应用"，不理解题目语义
  - 每步输出日志，失败时精确定位

用法:
    python typeset_exam.py --json {工作目录}/试卷数据/final_exam.json \
                           --template assets/template.dotx \
                           --images {工作目录}/清洗产物/images/ \
                           --output {工作目录}/{试卷名称}-排版后.docx \
                           [--format v1|v2] \
                           [--log {工作目录}/排版文档/typeset_log.txt]

    --format v1 (默认): 标准版式，含考试名称/科目、正文、末尾选择题答题卡
    --format v2: 封面版式，含独立封面（考试名称+答题卡）、注意事项、正文（无末尾答题卡）

输出:
    {试卷名称}-排版后.docx  排版后的 Word 文档
    quality_report.html     质检报告 (HTML，位于排版文档/)
    typeset_log.txt         排版日志（位于排版文档/）
"""

import argparse
import json
import logging
import os
import platform
import re
import sys
import zipfile
from datetime import datetime

from docx import Document
from docx.shared import Cm, Emu, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn as docx_qn
from docx.oxml import OxmlElement
from lxml import etree

try:
    from PIL import Image as PILImage
except ImportError:
    PILImage = None

# ============================================================================
# 页面常量
# ============================================================================
PAGE_CONTENT_WIDTH_CM = 21.0 - 1.9 - 1.9       # 版心宽度 17.2cm
PAGE_CONTENT_HEIGHT_CM = 29.7 - 2.54 - 2.54     # 版心高度 24.62cm

# 四号字 (14pt) 字符宽度
_CN_CHAR_WIDTH_CM = 14 * 0.0353                        # 中文全角 ~0.494cm
_EN_CHAR_WIDTH_CM = 14 * 0.0353 * 0.43                # 西文比例字体平均 ~0.212cm（Times New Roman）

# 选项制表位（取自 template.dotx「选项」样式，确保与模板一致）
# 模板原始制表位: [0.46, 4.54, 8.98, 13.43] cm（T1 为标签→文本跳距，排版不使用）
# 1x4: A\tB\tC\tD → B@T2(4.54), C@T3(8.98), D@T4(13.43)
#       各槽宽: A=4.54, B=4.44, C=4.45, D=3.77 cm
TAB_STOPS_4 = [Cm(4.54), Cm(8.98), Cm(13.43)]
# 2x2: 三档对齐，B/D 落在 4.54/8.98/13.43 之一（与 1x4 的 B/C/D 列对齐）
#       优先选最大可行位置（B/D 尽量靠右最美观），动态 1/2/3 tab 精确命中目标
#       情况2: A/C 短(<4.54)，B/D 长(≤12.66)  → B/D @ 4.54cm
#       情况3: A/C 中(<8.98)，B/D 中(≤8.22)   → B/D @ 8.98cm
#       情况4: A/C 长(<13.43)，B/D 短(≤3.77)  → B/D @ 13.43cm
TAB_STOPS_2 = [Cm(4.54), Cm(8.98), Cm(13.43)]

# 占位符匹配模式（支持 ph_xxx / img_xxx / ph_anchor_xxx，捕获组用于 split 分隔）
PLACEHOLDER_TOKEN_PATTERN = re.compile(r'\{\{image:(ph(?:_anchor)?_\d{3}|img_\d{3})\}\}')

# HTML 格式标签（sub/sup），宽度估算前需剔除
FORMAT_TAG_PATTERN = re.compile(r'</?(?:sub|sup)>')

# 填空题空位
FILL_IN_BLANK_PATTERN = re.compile(r'_{3,}')

# 中文标点（需强制中文字体）
_CN_PUNCT_CHARS = set(
    '""''<>《》「」『』【】〔〕〖〗〘〙〚〛…—·、。，；：？！'
    '（）()（）【】<>《》「」『』《》〈〉'
    '—–－~～·．・、，；：！？""''""‘’“”'
    '·・×÷±＝≠≤≥≈≡∝∞∑∏∫∮√∛∜∝∠⊥∥'
    '℃°′″‰％‱¤¥€£$₽₩₿'
    '※★☆◆◇▲△▼▽●○◎★☆'
    'ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ'
    '①②③④⑤⑥⑦⑧⑨⑩'
)

# 中文圈码数字（纯子条目标识，如 ①②③）
_CIRCLED_NUMBERS = frozenset('①②③④⑤⑥⑦⑧⑨⑩')


# ============================================================================
# 页面空间追踪器（用于自适应图片尺寸）
# ============================================================================

class PageSpaceTracker:
    """页面垂直空间追踪器。

    在排版过程中累加估算每段内容的高度，与版心高度比较，
    计算当前位置的剩余空间。当图片无法放入剩余空间时，
    自动缩小图片尺寸以适配当前页。

    原理：python-docx 不提供实时分页信息，因此通过
    累计估算逼近 Word 渲染引擎的实际分页行为。
    追踪器对文本高度做保守估算，实际 Word 渲染可能有细微偏差。
    """
    LINE_HEIGHT_CM = 0.74         # 四号字（14pt）单行高度
    PARAGRAPH_SPACING_CM = 0.15   # 段间距估算值
    CHARS_PER_LINE = 34           # 版心宽度约容纳的四号中文/全角字符数

    def __init__(self, content_height_cm=24.62):
        self.content_height = content_height_cm
        self.current_offset = 0.0
        self._page_count = 1

    @property
    def page_count(self):
        return self._page_count

    def remaining_space(self):
        """返回当前页面剩余可用空间（cm）。"""
        return max(0, self.content_height - self.current_offset)

    def consume_paragraph(self, estimated_lines=1):
        """记录一个段落消耗的垂直空间。"""
        self.current_offset += estimated_lines * self.LINE_HEIGHT_CM + self.PARAGRAPH_SPACING_CM

    def consume_text_lines(self, lines):
        """记录指定行数的文本消耗的垂直空间。"""
        self.current_offset += lines * self.LINE_HEIGHT_CM + self.PARAGRAPH_SPACING_CM

    def consume_image(self, height_cm):
        """记录一张图片消耗的垂直空间。"""
        self.current_offset += height_cm + self.PARAGRAPH_SPACING_CM

    def consume_table(self, height_cm):
        """记录一个表格消耗的预估垂直空间。"""
        self.current_offset += height_cm + self.PARAGRAPH_SPACING_CM

    def new_page(self):
        """重置为下一页（手动或 Word 自动换页后调用）。"""
        self.current_offset = 0.0
        self._page_count += 1

    def estimate_text_height_cm(self, text):
        """估算一段文本的垂直高度（含段间距）。

        Args:
            text: 文本内容

        Returns:
            估算高度（cm），含段间距
        """
        if not text:
            return 0
        cn_chars = 0
        en_chars = 0
        for ch in text:
            if ch.isascii():
                en_chars += 1
            else:
                cn_chars += 1
        effective = cn_chars + en_chars * 0.5
        lines = max(1, effective / self.CHARS_PER_LINE)
        return lines * self.LINE_HEIGHT_CM + self.PARAGRAPH_SPACING_CM


# 自适应图片尺寸阈值
IMAGE_MIN_WIDTH_CM = 6.0         # 图片最小宽度（低于此值不再缩小，允许跨页）
IMAGE_MIN_HEIGHT_CM = 4.0        # 图片最小高度
IMAGE_MAX_HEIGHT_CM = 16.0       # 绝对最大高度（不超过页面的 2/3）


# ============================================================================
# 日志
# ============================================================================

def setup_logger(log_path):
    logger = logging.getLogger('geo_exam_typeset')
    logger.setLevel(logging.DEBUG)
    logger.handlers = []

    fh = logging.FileHandler(log_path, mode='w', encoding='utf-8')
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(logging.Formatter('%(asctime)s [%(levelname)s] %(message)s'))
    logger.addHandler(fh)

    sh = logging.StreamHandler()
    sh.setLevel(logging.INFO)
    sh.setFormatter(logging.Formatter('%(message)s'))
    logger.addHandler(sh)

    return logger


# ============================================================================
# 模板加载
# ============================================================================

def load_template(template_path):
    """加载 dotx 模板文件，返回空的 Document（仅保留样式和 sectPr）。"""
    temp_docx = template_path + '.tmp.docx'

    with zipfile.ZipFile(template_path, 'r') as zin:
        with zipfile.ZipFile(temp_docx, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                data = zin.read(item)
                if item == '[Content_Types].xml':
                    data = data.replace(b'template.main+xml', b'document.main+xml')
                zout.writestr(item, data)

    try:
        doc = Document(temp_docx)
        body = doc.element.body
        for child in list(body):
            if child.tag != docx_qn('w:sectPr'):
                body.remove(child)
        return doc
    finally:
        if os.path.exists(temp_docx):
            os.remove(temp_docx)


# ============================================================================
# 页脚页码
# ============================================================================

def _add_page_number_footer(doc):
    """为文档所有节添加页脚：第 X 页 共 X 页，小四(12pt)，数字 Times New Roman，下方居中。

    使用 Word PAGE 和 NUMPAGES 域实现动态页码，打开文档后自动更新。
    """
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        # 清除页脚现有段落中的内容
        for p in footer.paragraphs:
            for child in list(p._element):
                if child.tag != docx_qn('w:pPr'):
                    p._element.remove(child)

        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 确保段落有 pPr
        pPr = p._element.find(docx_qn('w:pPr'))
        if pPr is None:
            pPr = etree.SubElement(p._element, docx_qn('w:pPr'))
        jc = pPr.find(docx_qn('w:jc'))
        if jc is None:
            jc = etree.SubElement(pPr, docx_qn('w:jc'))
        jc.set(docx_qn('w:val'), 'center')

        # "第 PAGE 页 共 NUMPAGES 页"
        _add_footer_text(p, '\u7b2c ')             # "第 "
        _add_field_code(p, ' PAGE ')               # PAGE 域
        _add_footer_text(p, ' \u9875 \u5171 ')     # " 页 共 "
        _add_field_code(p, ' NUMPAGES ')           # NUMPAGES 域
        _add_footer_text(p, ' \u9875')             # " 页"


def _add_footer_text(paragraph, text):
    """页脚普通文本 run：小四号 宋体"""
    run = paragraph.add_run(text)
    _set_run_font(run, cn_font='宋体', en_font='Times New Roman', size_pt=12)


def _add_field_code(paragraph, instr):
    """页脚 Word 域 run (PAGE / NUMPAGES)：小四号 Times New Roman"""
    # fldChar begin
    r1 = paragraph.add_run()
    _set_run_font(r1, cn_font='宋体', en_font='Times New Roman', size_pt=12)
    fld_begin = etree.SubElement(r1._element, docx_qn('w:fldChar'))
    fld_begin.set(docx_qn('w:fldCharType'), 'begin')

    # instrText
    r2 = paragraph.add_run()
    _set_run_font(r2, cn_font='宋体', en_font='Times New Roman', size_pt=12)
    instr_text = etree.SubElement(r2._element, docx_qn('w:instrText'))
    instr_text.set(docx_qn('xml:space'), 'preserve')
    instr_text.text = instr

    # fldChar separate
    r3 = paragraph.add_run()
    _set_run_font(r3, cn_font='宋体', en_font='Times New Roman', size_pt=12)
    fld_sep = etree.SubElement(r3._element, docx_qn('w:fldChar'))
    fld_sep.set(docx_qn('w:fldCharType'), 'separate')

    # field display value (默认值，打开 Word 后自动更新)
    r4 = paragraph.add_run('1')
    _set_run_font(r4, cn_font='宋体', en_font='Times New Roman', size_pt=12)

    # fldChar end
    r5 = paragraph.add_run()
    _set_run_font(r5, cn_font='宋体', en_font='Times New Roman', size_pt=12)
    fld_end = etree.SubElement(r5._element, docx_qn('w:fldChar'))
    fld_end.set(docx_qn('w:fldCharType'), 'end')


# ============================================================================
# 样式工具
# ============================================================================

def apply_style(paragraph, style_name, logger=None):
    try:
        paragraph.style = style_name
    except KeyError:
        if logger:
            logger.warning(f'样式不存在: {style_name}, 回退 Normal')
        try:
            paragraph.style = 'Normal'
        except KeyError:
            pass


def set_alignment(paragraph, alignment):
    paragraph.alignment = alignment


# ============================================================================
# 混合字体引擎
# ============================================================================

def _classify_char(ch):
    """分类字符: 'en' 或 'cn'。

    所有 ASCII 字符（含标点符号 ~ - ( ) % 等）归为 'en'，使用西文字体。
    非 ASCII 字符（含中文弯引号、CJK 字符等）归为 'cn'。
    """
    if ch.isascii():
        return 'en'
    if ch == '\u201c' or ch == '\u201d':  # 中文弯引号
        return 'cn'
    code = ord(ch)
    if 0x2070 <= code <= 0x208F:
        return 'en'
    return 'cn'


def _is_cn_punct(ch):
    return ch in _CN_PUNCT_CHARS


def _set_run_font(run, cn_font=None, en_font=None, size_pt=None, force_cn_all=False):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(docx_qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, docx_qn('w:rFonts'))
    if force_cn_all and cn_font:
        rFonts.set(docx_qn('w:ascii'), cn_font)
        rFonts.set(docx_qn('w:hAnsi'), cn_font)
        rFonts.set(docx_qn('w:eastAsia'), cn_font)
    else:
        if en_font:
            rFonts.set(docx_qn('w:ascii'), en_font)
            rFonts.set(docx_qn('w:hAnsi'), en_font)
        if cn_font:
            rFonts.set(docx_qn('w:eastAsia'), cn_font)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def _parse_format_tags(text):
    """解析包含 <sub> 和 <sup> 标签的文本，返回 [(text, tag_type)]。tag_type 为 None / 'sub' / 'sup'。"""
    import re
    parts = []
    pattern = re.compile(r'<(sub|sup)>(.*?)</\1>', re.DOTALL)
    last_end = 0
    
    for match in pattern.finditer(text):
        start, end = match.span()
        tag_type = match.group(1)  # 'sub' 或 'sup'
        if start > last_end:
            parts.append((text[last_end:start], None))
        parts.append((match.group(2), tag_type))
        last_end = end
    
    if last_end < len(text):
        parts.append((text[last_end:], None))
    
    return parts


def add_mixed_text(paragraph, text, cn_font=None, en_font='Times New Roman', size_pt=None):
    """按字符类型分段，中/西文使用不同字体，支持下标标签。"""
    if not text:
        return
    
    format_parts = _parse_format_tags(text)
    
    for sub_text, tag_type in format_parts:
        if not sub_text:
            continue
        
        segments = []
        current_type = _classify_char(sub_text[0])
        current_chars = [sub_text[0]]
        for ch in sub_text[1:]:
            t = _classify_char(ch)
            if t == current_type:
                current_chars.append(ch)
            else:
                segments.append((''.join(current_chars), current_type))
                current_chars = [ch]
                current_type = t
        segments.append((''.join(current_chars), current_type))

        for seg_text, seg_type in segments:
            run = paragraph.add_run(seg_text)
            if tag_type == 'sub':
                run.font.subscript = True
            elif tag_type == 'sup':
                run.font.superscript = True
            if tag_type in ('sub', 'sup') and size_pt is not None:
                actual_size = size_pt * 0.7
            else:
                actual_size = size_pt
            if seg_type == 'en':
                _set_run_font(run, cn_font=cn_font, en_font=en_font, size_pt=actual_size)
            else:
                force = any(_is_cn_punct(c) for c in seg_text)
                _set_run_font(run, cn_font=cn_font, en_font=en_font, size_pt=actual_size, force_cn_all=force)


def add_multi_paragraph_material(doc, text, cn_font='楷体', first_line_prefix=None, page_tracker=None):
    """将多段材料文本按 \\n 拆分为多个独立段落，每段应用 Body Text 样式（含首行缩进）。
    
    Args:
        doc: Word 文档对象
        text: 可能包含 \\n 的原始文本
        cn_font: 中文字体名称
        first_line_prefix: 可选，追加到首段文本前的内联标题（如 '材料一'），使用黑体渲染（不加粗）
        page_tracker: 可选，页面空间追踪器，用于记录段落消耗的垂直空间
    Returns:
        最后一个文本段落的 Word paragraph 对象
    """
    if not text:
        return None
    lines = text.split('\n')
    last_p = None
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        apply_style(p, 'Body Text')
        if i == 0 and first_line_prefix:
            add_mixed_text(p, first_line_prefix, cn_font='黑体')
            add_mixed_text(p, line, cn_font=cn_font)
            rendered = first_line_prefix + line
        else:
            add_mixed_text(p, line, cn_font=cn_font)
            rendered = line
        if page_tracker:
            page_tracker.current_offset += page_tracker.estimate_text_height_cm(rendered)
        last_p = p
    return last_p


def clear_run_fonts(paragraph):
    """清除段落中所有 run 的字体覆盖，让样式控制字体。"""
    for run in paragraph.runs:
        rPr = run._element.find(docx_qn('w:rPr'))
        if rPr is not None:
            for tag in ['w:rFonts', 'w:sz', 'w:szCs']:
                for el in rPr.findall(docx_qn(tag)):
                    rPr.remove(el)


# ============================================================================
# 文本宽度测量 (Pillow 精确测量优先，估算兜底)
# ============================================================================

# --- 估算用精细化字符宽度表（Times New Roman 14pt 实测近似值，cm） ---
# 窄字符: i l j 1 | . , ; : ' " ( ) [ ] t f I r
_EN_NARROW_CHARS = set('ilj1|.,;:\'"()[]tfIr')
_EN_NARROW_WIDTH_CM = 14 * 0.0353 * 0.25       # ~0.124cm
# 宽字符: W M m w @
_EN_WIDE_CHARS = set('WMmw@')
_EN_WIDE_WIDTH_CM = 14 * 0.0353 * 0.78          # ~0.384cm
# 数字
_EN_DIGIT_WIDTH_CM = 14 * 0.0353 * 0.50          # ~0.247cm


def estimate_text_width_cm(text, size_pt=14):
    """精细化文本宽度估算（cm）。

    修复 BUG-09：ASCII 标点不再使用统一西文宽度，按字符类别细分。
    保留旧签名兼容（text-only），新增 size_pt 参数支持缩放。
    """
    scale = size_pt / 14.0 if size_pt != 14 else 1.0
    width = 0.0
    for ch in text:
        if ch.isascii():
            if ch in _EN_NARROW_CHARS:
                width += _EN_NARROW_WIDTH_CM
            elif ch in _EN_WIDE_CHARS:
                width += _EN_WIDE_WIDTH_CM
            elif ch.isdigit():
                width += _EN_DIGIT_WIDTH_CM
            elif ch == ' ':
                width += _EN_NARROW_WIDTH_CM  # 空格与窄标点相近
            else:
                width += _EN_CHAR_WIDTH_CM    # 默认西文（多数字母）
        else:
            width += _CN_CHAR_WIDTH_CM
    return round(width * scale, 2)


# --- 字体文件路径探测（跨平台） ---

_font_path_cache = {}
# 已加载的 ImageFont 缓存: {(font_path, size_px): ImageFont}
_image_font_cache = {}

_FONT_CANDIDATES = {
    '宋体': [
        r'C:\Windows\Fonts\simsun.ttc',
        '/System/Library/Fonts/Supplemental/Songti.ttc',
        '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
        '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
    ],
    'Times New Roman': [
        r'C:\Windows\Fonts\times.ttf',
        '/Library/Fonts/Times New Roman.ttf',
        '/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf',
    ],
}


def _find_font_path(font_name):
    """跨平台查找字体文件路径，结果缓存。找不到返回 None。"""
    if font_name in _font_path_cache:
        return _font_path_cache[font_name]

    # 1. 尝试已知候选路径
    for path in _FONT_CANDIDATES.get(font_name, []):
        if os.path.exists(path):
            _font_path_cache[font_name] = path
            return path

    # 2. Windows 注册表查找（兜底，覆盖字体重命名/非标准安装路径）
    if platform.system() == 'Windows':
        try:
            import winreg
            key = winreg.OpenKey(
                winreg.HKEY_LOCAL_MACHINE,
                r'SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts'
            )
            i = 0
            while True:
                try:
                    name, value, _ = winreg.EnumValue(key, i)
                    if font_name in name and value.lower().endswith(('.ttf', '.ttc', '.otf')):
                        full_path = value if os.path.isabs(value) else os.path.join(
                            os.environ.get('WINDIR', r'C:\Windows'), 'Fonts', value)
                        if os.path.exists(full_path):
                            _font_path_cache[font_name] = full_path
                            return full_path
                except OSError:
                    break
                i += 1
            winreg.CloseKey(key)
        except Exception:
            pass

    _font_path_cache[font_name] = None
    return None


def _get_image_font(font_path, size_pt):
    """获取缓存的 PIL ImageFont 对象。"""
    size_px = max(1, int(round(size_pt * 96 / 72)))  # pt → px @96dpi
    key = (font_path, size_px)
    if key not in _image_font_cache:
        from PIL import ImageFont
        # .ttc 文件: index=0 通常为主字体（simsun.ttc[0]=宋体）
        if font_path.lower().endswith('.ttc'):
            _image_font_cache[key] = ImageFont.truetype(font_path, size_px, index=0)
        else:
            _image_font_cache[key] = ImageFont.truetype(font_path, size_px)
    return _image_font_cache[key]


def _measure_with_pillow(text, size_pt=14):
    """使用 Pillow + 真实字体文件精确测量文本宽度（cm）。

    按字符类型分别使用对应字体：
    - 中文字符（含中文标点）→ 宋体
    - 西文字符（含 ASCII 标点）→ Times New Roman

    Returns:
        宽度（cm），或 None（字体文件不可用时）
    """
    cn_font_path = _find_font_path('宋体')
    en_font_path = _find_font_path('Times New Roman')
    if not cn_font_path or not en_font_path:
        return None

    cn_font = _get_image_font(cn_font_path, size_pt)
    en_font = _get_image_font(en_font_path, size_pt)

    dpi = 96
    total_px = 0.0
    for ch in text:
        char_type = _classify_char(ch)
        font = cn_font if char_type == 'cn' else en_font
        total_px += font.getlength(ch)

    # px → cm
    return round(total_px / dpi * 2.54, 2)


def measure_text_width_cm(text, size_pt=14):
    """统一文本宽度测量：Pillow 精确测量优先，估算兜底。

    Args:
        text: 纯文本（调用方应先剔除占位符和格式标签）
        size_pt: 字号（pt），默认四号字 14pt

    Returns:
        文本宽度（cm）
    """
    # 主方案：Pillow 精确测量
    try:
        result = _measure_with_pillow(text, size_pt)
        if result is not None:
            return result
    except Exception:
        pass

    # 兜底：精细化估算
    return estimate_text_width_cm(text, size_pt)


# ============================================================================
# 图片处理
# ============================================================================

def get_image_size_cm(image_path):
    if PILImage is None:
        return None
    try:
        with PILImage.open(image_path) as im:
            px_w, px_h = im.size
            dpi = im.info.get('dpi', (96, 96))
            dpi_x = dpi[0] if dpi[0] else 96
            dpi_y = dpi[1] if dpi[1] else 96
            return px_w / dpi_x * 2.54, px_h / dpi_y * 2.54
    except Exception:
        return None


def compute_display_width(image_path, max_w_cm=12.0, max_h_cm=8.0):
    """等比缩放计算合适的显示宽度（只缩小不放大）。"""
    size = get_image_size_cm(image_path)
    if size is None:
        return Cm(max_w_cm)
    orig_w, orig_h = size
    if orig_w <= 0 or orig_h <= 0:
        return Cm(max_w_cm)
    scale = 1.0
    if orig_w > max_w_cm:
        scale = min(scale, max_w_cm / orig_w)
    if orig_h > max_h_cm:
        scale = min(scale, max_h_cm / orig_h)
    return Cm(orig_w * scale)


def add_picture(paragraph, image_path, width=None):
    if not os.path.exists(image_path):
        run = paragraph.add_run(f'[图片缺失: {os.path.basename(image_path)}]')
        return False
    if width is None:
        width = compute_display_width(image_path)
    run = paragraph.add_run()
    run.add_picture(image_path, width=width)
    return True


def add_centered_picture(doc, image_path, logger):
    p = doc.add_paragraph()
    set_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)
    apply_style(p, 'Normal')
    width = compute_display_width(image_path)
    ok = add_picture(p, image_path, width=width)
    if ok:
        logger.debug(f'  插入图片: {os.path.basename(image_path)}')
    return p, ok


def compute_adaptive_display_size(image_path, page_tracker):
    """根据页面剩余空间自适应计算图片最佳显示尺寸。

    策略（按优先级）：
    1. 获取图片原始尺寸
    2. 以版心宽度（17.2cm）为宽度上限计算候选尺寸
    3. 若候选高度 > 剩余空间，以剩余空间为高度约束等比缩小
    4. 缩小后不低于最小阈值（低于阈值保持最小尺寸，允许跨页）
    5. 不放大原始图片

    Args:
        image_path: 图片文件路径
        page_tracker: PageSpaceTracker 实例

    Returns:
        (display_width_cm, display_height_cm, was_scaled)
        - display_width_cm: 显示宽度（cm）
        - display_height_cm: 显示高度（cm）
        - was_scaled: 是否进行了自适应缩放
    """
    original_size = get_image_size_cm(image_path)
    if original_size is None:
        # 无法获取尺寸：保守使用中等尺寸
        safe_h = min(8.0, page_tracker.remaining_space() * 0.55)
        return 12.0, safe_h, True

    orig_w, orig_h = original_size
    if orig_w <= 0 or orig_h <= 0:
        return PAGE_CONTENT_WIDTH_CM, 8.0, True

    remaining = page_tracker.remaining_space()
    aspect_ratio = orig_w / orig_h

    # Step 1: 先按版心宽度约束计算候选尺寸
    target_w = min(orig_w, PAGE_CONTENT_WIDTH_CM)
    target_h = target_w / aspect_ratio

    # Step 2: 检查高度是否超出剩余空间
    if target_h > remaining - 0.3:
        # 以剩余空间为高度约束（留0.3cm缓冲）
        target_h = max(0.5, remaining - 0.3)
        target_w = target_h * aspect_ratio
        # 如果宽度也因此超出，再以宽度约束
        if target_w > PAGE_CONTENT_WIDTH_CM:
            target_w = PAGE_CONTENT_WIDTH_CM
            target_h = target_w / aspect_ratio

    # Step 3: 高度不超过绝对上限
    if target_h > IMAGE_MAX_HEIGHT_CM:
        target_h = IMAGE_MAX_HEIGHT_CM
        target_w = target_h * aspect_ratio

    # Step 4: 检查下限阈值
    was_scaled = False
    if target_h < IMAGE_MIN_HEIGHT_CM or target_w < IMAGE_MIN_WIDTH_CM:
        # 缩小后会太小：保持最小尺寸，允许图片跨页
        target_h = max(target_h, IMAGE_MIN_HEIGHT_CM)
        target_w = max(target_w, IMAGE_MIN_WIDTH_CM)
        # 按最小高度重新等比计算
        target_h = IMAGE_MIN_HEIGHT_CM
        target_w = target_h * aspect_ratio
        was_scaled = True
    elif abs(target_w - orig_w) > 0.05 or abs(target_h - orig_h) > 0.05:
        was_scaled = True

    # Step 5: 不放大原始图片
    target_w = min(target_w, orig_w)
    target_h = min(target_h, orig_h)

    return round(target_w, 2), round(target_h, 2), was_scaled


def add_centered_picture_adaptive(doc, image_path, page_tracker, logger):
    """居中插入图片（自适应版），根据页面剩余空间自动调整尺寸。

    Args:
        doc: Word Document 对象
        image_path: 图片文件路径
        page_tracker: PageSpaceTracker 实例
        logger: 日志记录器

    Returns:
        (paragraph, success)
    """
    p = doc.add_paragraph()
    set_alignment(p, WD_ALIGN_PARAGRAPH.CENTER)
    apply_style(p, 'Normal')

    if not os.path.exists(image_path):
        run = p.add_run(f'[图片缺失: {os.path.basename(image_path)}]')
        page_tracker.consume_paragraph(estimated_lines=1)
        return p, False

    width_cm, height_cm, was_scaled = compute_adaptive_display_size(image_path, page_tracker)

    width_emu = Cm(width_cm)
    ok = add_picture(p, image_path, width=width_emu)

    if ok:
        page_tracker.consume_image(height_cm)
        if was_scaled:
            logger.info(f'  图片自适应: {os.path.basename(image_path)} → {width_cm:.1f}×{height_cm:.1f}cm '
                        f'(剩余空间: {page_tracker.remaining_space():.1f}cm)')
        else:
            logger.debug(f'  插入图片: {os.path.basename(image_path)} ({width_cm:.1f}×{height_cm:.1f}cm)')
    else:
        page_tracker.consume_paragraph(estimated_lines=1)

    return p, ok


# ============================================================================
# 图片映射解析
# ============================================================================

class ImageResolver:
    """将 placeholder_id 解析为实际图片文件路径。

    解析链路: placeholder_id → image_mapping → image_id → images[].file_name → file_path

    v3.6.7 新增：_images_data 存储原始图片元数据（含 clues），供非选择题兜底匹配使用；
             _resolved_phs 追踪已渲染占位符，避免 clues 匹配时重复渲染。
    """

    def __init__(self, images_data, image_mapping_data, images_dir):
        self._ph_to_file = {}
        self._unmapped = set()
        self._images_data = images_data
        self._image_mapping_data = image_mapping_data
        self._resolved_phs = set()

        # 构建 image_id → file_name 映射
        id_to_file = {}
        for img in images_data:
            img_id = img.get('image_id')
            file_name = img.get('file_name')
            if img_id and file_name:
                id_to_file[img_id] = file_name

        # 构建 placeholder_id → file_path 映射
        for mapping in image_mapping_data:
            ph_id = mapping.get('placeholder_id')
            img_id = mapping.get('image_id')
            if ph_id and img_id:
                file_name = id_to_file.get(img_id)
                if file_name:
                    full_path = os.path.normpath(os.path.join(images_dir, file_name))
                    self._ph_to_file[ph_id] = full_path
                else:
                    self._unmapped.add(ph_id)

    def resolve(self, placeholder_id):
        """返回 (file_path, exists)，不存在时返回 (None, False)。渲染成功后自动追踪。"""
        path = self._ph_to_file.get(placeholder_id)
        if path and os.path.exists(path):
            self._resolved_phs.add(placeholder_id)
            return path, True
        return path, False

    def is_rendered(self, placeholder_id):
        """检查某占位符是否已被 resolve 并渲染过。"""
        return placeholder_id in self._resolved_phs

    def find_images_by_clues(self, question):
        """通过 images[].clues 匹配题目子问题，返回未被渲染的图片列表。

        匹配规则：clues 中去除 {{image:xxx}} 格式的条目后，与 subquestions[].stem
        做精确文本匹配。每个匹配结果包含 file_path 和对应的 subquestion 索引。

        Returns:
            list of dict: [{'file_path': str, 'placeholder_id': str, 'subq_index': int}, ...]
        """
        subq_stems = [sq.get('stem', '').strip() for sq in question.get('subquestions', [])]
        if not subq_stems:
            return []

        matched = []

        for mapping in self._image_mapping_data:
            ph_id = mapping.get('placeholder_id', '')
            img_id = mapping.get('image_id', '')
            if not ph_id or not img_id:
                continue
            if ph_id in self._resolved_phs:
                continue

            # 在 images_data 中找对应的图片元数据
            img_meta = next((i for i in self._images_data if i.get('image_id') == img_id), None)
            if not img_meta:
                continue

            clues = img_meta.get('clues', [])
            file_path = self._ph_to_file.get(ph_id)
            if not file_path or not os.path.exists(file_path):
                continue

            # 遍历 clues 中的非占位符条目，匹配子问题 stem
            for clue in clues:
                if clue.startswith('{{image:'):
                    continue
                clue_clean = clue.strip()
                for idx, stem in enumerate(subq_stems):
                    if clue_clean == stem:
                        matched.append({
                            'file_path': file_path,
                            'placeholder_id': ph_id,
                            'subq_index': idx,
                        })
                        break
                else:
                    continue
                break  # 找到匹配后跳出 clues 循环

        return matched

    @property
    def unmapped_placeholders(self):
        return self._unmapped


# ============================================================================
# 制表位
# ============================================================================

def set_tab_stops(paragraph, tab_positions):
    pPr = paragraph._element.get_or_add_pPr()
    for existing in pPr.findall(docx_qn('w:tabs')):
        pPr.remove(existing)
    tabs = etree.SubElement(pPr, docx_qn('w:tabs'))
    for pos in tab_positions:
        tab = etree.SubElement(tabs, docx_qn('w:tab'))
        tab.set(docx_qn('w:val'), 'left')
        twips = str(int(pos / 635))
        tab.set(docx_qn('w:pos'), twips)


# ============================================================================
# 表格处理
# ============================================================================

def _compute_table_column_widths(data, rows, cols, spans, covered_cells, page_width_cm):
    """根据单元格内容估算最优列宽分配，最小化表格总行数。

    算法：
    1. 计算每列中最宽单元格的预估宽度（取 max）
    2. 按比例分配版心宽度，确保各列宽度≥最小宽度
    3. 若总需求超出，等比压缩

    Args:
        data: 表格数据二维数组
        rows, cols: 行列数
        spans: 合并信息
        covered_cells: 被合并覆盖的单元格集合
        page_width_cm: 版心总宽度(cm)

    Returns:
        list[float]: 每列宽度(cm)，总和 == page_width_cm
    """
    MIN_COL_WIDTH_CM = 1.5       # 单列最小宽度
    HEADER_PADDING_CM = 0.3      # 表头额外宽度（黑体加粗需要更多空间）
    PADDING_CM = 0.2             # 单元格内边距余量

    # 计算每列最大预估宽度
    col_max_widths = [0.0] * cols
    for i in range(rows):
        for j in range(cols):
            if (i, j) in covered_cells:
                continue
            row_data = data[i] if i < len(data) else []
            cell_text = str(row_data[j]) if j < len(row_data) else ''
            if not cell_text:
                continue
            # 估算多行文本的宽度（取最长一行的宽度）
            text_lines = cell_text.split('\n')
            max_line_w = max((measure_text_width_cm(line) for line in text_lines), default=0.0)
            col_max_widths[j] = max(col_max_widths[j], max_line_w)

    # 加上内边距
    col_max_widths = [w + PADDING_CM for w in col_max_widths]

    # 应用最小宽度
    col_max_widths = [max(w, MIN_COL_WIDTH_CM) for w in col_max_widths]

    total_required = sum(col_max_widths)

    if total_required <= page_width_cm:
        # 总需求不超过版心：按比例放大到版心宽度
        scale = page_width_cm / total_required if total_required > 0 else 1.0
        col_widths = [w * scale for w in col_max_widths]
    else:
        # 总需求超出：等比压缩
        scale = page_width_cm / total_required
        col_widths = [max(w * scale, MIN_COL_WIDTH_CM) for w in col_max_widths]
        # 重新归一化到版心宽度
        adjusted_total = sum(col_widths)
        if adjusted_total > 0:
            scale2 = page_width_cm / adjusted_total
            col_widths = [w * scale2 for w in col_widths]

    return col_widths


def _set_table_column_widths(table, col_widths_cm, total_width_cm):
    """为表格所有行设置列宽。

    Word 表格列宽需要在每行的第一个单元格（或每行每列）上设置 tcW。
    这里对每行的每个单元格设置宽度，确保合并单元格也正确处理。
    """
    rows = len(table.rows)
    cols = len(table.columns)
    for i in range(rows):
        row = table.rows[i]
        for j in range(min(cols, len(col_widths_cm))):
            cell = row.cells[j]
            tc = cell._tc
            tcPr = tc.find(docx_qn('w:tcPr'))
            if tcPr is None:
                tcPr = etree.SubElement(tc, docx_qn('w:tcPr'))
            tcW = tcPr.find(docx_qn('w:tcW'))
            if tcW is None:
                tcW = etree.SubElement(tcPr, docx_qn('w:tcW'))
            tcW.set(docx_qn('w:w'), str(int(Cm(col_widths_cm[j]))))
            tcW.set(docx_qn('w:type'), 'dxa')


def add_table(doc, table_data, logger, image_resolver=None, role='material'):
    rows = table_data.get('rows', 0)
    cols = table_data.get('cols', 0)
    data = table_data.get('data', [])
    spans = table_data.get('spans', [])
    if rows == 0 or cols == 0 or not data:
        return None

    # ---- 计算 header_rows（兼容旧 has_header 字段） ----
    header_rows = table_data.get('header_rows')
    if header_rows is None:
        # 兼容旧数据：has_header true → 1, false/缺失 → 0
        header_rows = 1 if table_data.get('has_header', False) else 0
    header_rows = max(0, int(header_rows or 0))

    # ---- table_style ----
    table_style_name = table_data.get('table_style', 'grid')
    STYLE_MAP = {
        'grid': 'Table Grid',
        'plain': 'Table Normal',
        'preserve': 'Table Grid',  # 暂退化
    }
    applied_style = STYLE_MAP.get(table_style_name, 'Table Grid')

    # ---- 先创建空表格并设置样式 ----
    table = doc.add_table(rows=rows, cols=cols)
    try:
        table.style = applied_style
    except KeyError:
        try:
            table.style = 'Table Grid'
        except KeyError:
            pass
    table.autofit = False
    table.allow_autofit = False
    table.width = Cm(PAGE_CONTENT_WIDTH_CM)

    # ---- 先执行合并，构建 covered_cells 集合 ----
    covered_cells = set()
    if spans:
        for span in spans:
            try:
                r = span.get('row', 0)
                c = span.get('col', 0)
                rs = span.get('rowspan', 1)
                cs = span.get('colspan', 1)
                if rs > 1 or cs > 1:
                    start_cell = table.cell(r, c)
                    end_cell = table.cell(r + rs - 1, c + cs - 1)
                    start_cell.merge(end_cell)
                    # 标记被覆盖的单元格
                    for rr in range(r, r + rs):
                        for cc in range(c, c + cs):
                            if rr == r and cc == c:
                                continue
                            covered_cells.add((rr, cc))
            except Exception as e:
                logger.warning(f'  表格合并失败 (row={r}, col={c}, rowspan={rs}, colspan={cs}): {e}')

    # ---- 再填充内容（跳过 covered_cells） ----
    for i in range(rows):
        for j in range(cols):
            if (i, j) in covered_cells:
                continue
            row_data = data[i] if i < len(data) else []
            cell_text = row_data[j] if j < len(row_data) else ''
            cell = table.cell(i, j)
            cell.text = ''
            paragraph = cell.paragraphs[0]
            # 数据表格表头居中对齐，非表头根据角色决定对齐方式
            if i < header_rows:
                set_alignment(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
                # 表头垂直居中
                tcPr = cell._tc.get_or_add_tcPr()
                vAlign = etree.SubElement(tcPr, docx_qn('w:vAlign'))
                vAlign.set(docx_qn('w:val'), 'center')
            else:
                set_alignment(paragraph, WD_ALIGN_PARAGRAPH.LEFT)
            paragraph.paragraph_format.line_spacing = 1.0

            cell_str = str(cell_text)

            # 根据 header_rows 决定字体
            if i < header_rows:
                cell_font = '黑体'
                is_header = True
            else:
                # 材料表格用楷体，与材料正文一致；其他表格用宋体
                cell_font = '楷体' if role == 'material' else '宋体'
                is_header = False

            # 渲染单元格内容（支持占位符）
            _render_cell_content(paragraph, cell_str, cell_font, is_header,
                                 image_resolver, logger)

    # ---- 计算并设置列宽（根据内容智能分配） ----
    col_widths = _compute_table_column_widths(data, rows, cols, spans, covered_cells, PAGE_CONTENT_WIDTH_CM)
    _set_table_column_widths(table, col_widths, PAGE_CONTENT_WIDTH_CM)

    # ---- 设置跨页表头（仅 header_rows > 0 时） ----
    if header_rows > 0 and rows > 1:
        tbl = table._tbl
        tblPr = tbl.find(docx_qn('w:tblPr'))
        if tblPr is None:
            tblPr = etree.SubElement(tbl, docx_qn('w:tblPr'))
        tblHeader = etree.SubElement(tblPr, docx_qn('w:tblHeader'))
        tblHeader.set(docx_qn('w:val'), 'true')

    merge_info = f', 合并{len(spans)}处' if spans else ''
    style_info = f', {table_style_name}' if table_style_name != 'grid' else ''
    header_info = f', 表头{header_rows}行' if header_rows > 0 else ''
    col_widths_str = ', '.join(f'{w:.1f}' for w in col_widths)
    col_info = f', 列宽=[{col_widths_str}]cm'
    logger.debug(f'  表格: {rows}行 x {cols}列{merge_info}{style_info}{header_info}{col_info}')
    return table


def _render_cell_content(paragraph, cell_str, cn_font, is_header, image_resolver, logger):
    """渲染单元格内容，支持 {{image:xxx}} 和 {{symbol:xxx}} 占位符。

    将单元格文本按占位符拆分，对文本添加 run，对图片插入实际图片。

    Args:
        paragraph: python-docx Paragraph 对象
        cell_str: 单元格文本（可能含占位符）
        cn_font: 中文字体名
        is_header: 是否为表头行
        image_resolver: ImageResolver 实例或 None
        logger: logger 实例
    """
    if not cell_str:
        return

    parts = PLACEHOLDER_TOKEN_PATTERN.split(cell_str)
    for k, part in enumerate(parts):
        if not part:
            continue
        if k % 2 == 0:
            # 纯文本
            add_mixed_text(paragraph, part, cn_font=cn_font)
        else:
            # 占位符 ID（如 img_001, ph_001, ph_anchor_001）
            ph_id = part
            if image_resolver is not None:
                file_path, exists = image_resolver.resolve(ph_id)
                if exists:
                    # 在已有段落中插入图片
                    run = paragraph.add_run()
                    try:
                        width = compute_display_width(file_path)
                        run.add_picture(file_path, width=width)
                    except Exception as e:
                        logger.warning(f'    表格单元格图片插入失败 ({ph_id}): {e}')
                        run.text = f'[图片:{ph_id}]'
                else:
                    # 图片找不到，保留占位文本
                    run = paragraph.add_run(f'[图片:{ph_id}]')
            else:
                # 无 image_resolver，保留占位文本
                run = paragraph.add_run(f'[图片:{ph_id}]')


# ============================================================================
# 选项排版引擎
# ============================================================================

def _select_option_rule(options, logger=None):
    """根据选项文本宽度选择排版规则。

    规则 1 (1x4): 四个选项一行，Tab 分隔
    规则 2 (2x2): 两行，每行两个，B/D 对齐 1x4 的 B/C/D 之一（三档）
    规则 3 (4x1): 每行一个

    1x4 制表位 TAB_STOPS_4 = [4.54, 8.98, 13.43]，版心宽 17.2cm
    各槽位实际宽度: A=4.54, B=4.44, C=4.45, D=3.77

    2x2 三档对齐（制表位与 1x4 相同 [4.54, 8.98, 13.43]）：
      情况2: B/D @ 4.54cm  — A/C 短(<4.54)，B/D 长(≤12.66)，单 tab
      情况3: B/D @ 8.98cm  — A/C 中等(<8.98)，B/D 中等(≤8.22)，单/双 tab
      情况4: B/D @ 13.43cm — A/C 长(<13.43)，B/D 短(≤3.77)，单/双/三 tab
    优先选最大可行目标位置（B/D 尽量靠右，视觉美观度最优先）。

    测量使用 measure_text_width_cm（Pillow 精确测量优先，估算兜底）。
    所有槽位/列宽检查均预留 SAFETY_MARGIN_CM 安全余量，防止 Word 渲染时
    因微小差异导致换行。
    """
    SAFETY_MARGIN_CM = 0.15  # 1x4 安全余量，补偿 tab 间距与渲染偏差
    SAFETY_MARGIN_2X2_CM = 0.05  # 2x2 安全余量较小（无 tab 间距，仅纯文本列宽）

    opt_widths = {}
    for label, text in options.items():
        clean = PLACEHOLDER_TOKEN_PATTERN.sub('', f'{label}. {text}')
        clean = FORMAT_TAG_PATTERN.sub('', clean)  # 剔除 <sub>/<sup> 标签
        opt_widths[label] = measure_text_width_cm(clean)

    letters = sorted(options.keys())
    n = len(letters)

    # 规则1: 检查每个选项是否都能放入 1x4 的对应槽位（含安全余量）
    rule1_failed = []
    if n <= 4:
        slot_widths = _compute_1x4_slot_widths(n)
        all_fit = True
        for i in range(n):
            label = letters[i]
            w = opt_widths[label]
            sw = slot_widths[i] - SAFETY_MARGIN_CM
            if w > sw:
                all_fit = False
                rule1_failed.append(f'{label}={w:.2f}>{sw:.2f}')
        if all_fit:
            return 1
        if logger:
            logger.debug(f'  1x4不适用: {rule1_failed}')

    # 规则2: 2x2 — 三档对齐（B/D 落在 4.54/8.98/13.43 之一，与 1x4 的 B/C/D 对齐）
    first_col_labels = ['A', 'C']
    second_col_labels = ['B', 'D']
    first_col_max = max(
        (opt_widths.get(l, 0) for l in first_col_labels), default=0
    )
    second_col_max = max(
        (opt_widths.get(l, 0) for l in second_col_labels), default=0
    )
    # 依次尝试三个目标位置，选最大可行位置（B/D 尽量靠右，与 1x4 的 D 对齐最美观）
    targets_2x2 = [13.43, 8.98, 4.54]
    rule2_ok = False
    chosen_target = None
    for t in targets_2x2:
        if (first_col_max + SAFETY_MARGIN_2X2_CM <= t and
                second_col_max + SAFETY_MARGIN_2X2_CM <= PAGE_CONTENT_WIDTH_CM - t):
            rule2_ok = True
            chosen_target = t
            break

    if logger:
        width_detail = ', '.join(f'{l}={opt_widths[l]:.2f}cm' for l in letters)
        target_str = f'{chosen_target:.2f}cm' if chosen_target else '无'
        logger.debug(
            f'  2x2检查: 左列最大={first_col_max:.2f}cm, 右列最大={second_col_max:.2f}cm, '
            f'目标位置={target_str}, 通过={rule2_ok} | 选项宽: {width_detail}'
        )

    if rule2_ok:
        return 2
    return 3


def _compute_1x4_slot_widths(n):
    """计算 n 个选项在 1x4 布局中各槽位的实际宽度(cm)。

    制表位: TAB_STOPS_4 = [4.54, 8.98, 13.43] (取自模板「选项」样式)
    最后一个选项使用版心剩余宽度，其余选项使用相邻制表位之差。
    """
    tab_stops = [4.54, 8.98, 13.43]
    page_w = PAGE_CONTENT_WIDTH_CM
    if n == 1:
        return [page_w]
    widths = [tab_stops[0]]  # 第一个选项: 0 → 第一个制表位
    for i in range(1, n - 1):
        widths.append(tab_stops[i] - tab_stops[i - 1])
    # 最后一个选项: 最后一个使用的制表位 → 版心边缘
    widths.append(page_w - tab_stops[n - 2])
    return widths


def _add_option_label_text(paragraph, label, text, image_resolver, logger, img_width=None, page_tracker=None):
    """添加选项标签和文本，处理占位符中的图片。"""
    paragraph.add_run(f'{label}. ')

    if img_width is None:
        img_width = Cm(6)

    parts = PLACEHOLDER_TOKEN_PATTERN.split(text)
    for i, part in enumerate(parts):
        if i % 2 == 0:
            if part:
                add_mixed_text(paragraph, part, cn_font='宋体')
        else:
            # part 是 placeholder_id (如 ph_001)
            # 但这里需要匹配 token {{image:ph_001}}
            token = f'{{{{image:{part}}}}}'
            # 查找原始占位符
            ph_file_path, exists = image_resolver.resolve(part) if (part.startswith('ph_') or part.startswith('img_')) else (None, False)
            if exists:
                run = paragraph.add_run()
                # 选项内图片：如有 tracker，超过剩余空间时自适应缩小
                if page_tracker:
                    width_cm, height_cm, was_scaled = compute_adaptive_display_size(ph_file_path, page_tracker)
                    run.add_picture(ph_file_path, width=Cm(width_cm))
                    page_tracker.consume_image(height_cm + 0.2)
                    if was_scaled:
                        logger.info(f'  选项图片自适应: {os.path.basename(ph_file_path)} → {width_cm:.1f}×{height_cm:.1f}cm')
                else:
                    run.add_picture(ph_file_path, width=img_width)
            elif not part.startswith('img_'):
                paragraph.add_run(f'[图片:{part}]')
            # img_ 开头但未映射的占位符：可能是已被清洗替换的小图/符号图，静默跳过


def format_options(doc, options, image_resolver, logger, page_tracker=None, sub_options=None):
    """选择题选项排版入口。"""
    if not options:
        return

    # v3.0 的 options 是 [{label, text}] 数组，转为 {label: text} 字典
    if isinstance(options, list):
        opt_dict = {}
        for opt in options:
            label = opt.get('label', '')
            text = opt.get('text', '')
            if label:
                opt_dict[label] = text
        options = opt_dict

    if not options:
        return

    # 子选项
    if sub_options:
        logger.debug(f'  [format_options] 收到 sub_options: {len(sub_options)}项')
        _format_sub_options(doc, sub_options, logger, page_tracker)
    else:
        logger.debug(f'  [format_options] sub_options 为空，跳过子选项排版')

    # 检查是否含图片占位符
    has_images = any(PLACEHOLDER_TOKEN_PATTERN.search(v) for v in options.values())

    if has_images:
        rule = 5
    else:
        rule = _select_option_rule(options, logger)

    logger.debug(f'  选项排版: 规则{rule}')

    if rule == 1:
        _format_rule_1x4(doc, options, image_resolver, logger, page_tracker)
    elif rule == 2:
        _format_rule_2x2(doc, options, image_resolver, logger, page_tracker)
    elif rule == 3:
        _format_rule_4x1(doc, options, image_resolver, logger, page_tracker)
    elif rule == 5:
        _format_rule_image(doc, options, image_resolver, logger, page_tracker)

    return rule


def _format_sub_options(doc, sub_options, logger, page_tracker=None):
    """排版子选项（①②③④等）。按精确宽度估算分行，避免单个子选项换行。

    每两个子选项之间用2个空格分隔。
    如果一行能放下则放一行，放不下则自动分行（一行两个、三个或四个）。
    """
    if not sub_options:
        return

    # 按内容宽度自动分行
    current_line = []
    current_width = 0.0
    gap_cm = 0.5  # 2个空格在四号字下的宽度约0.5cm
    max_width = PAGE_CONTENT_WIDTH_CM
    line_count = 0

    logger.debug(f'    [子选项贪心] 共{len(sub_options)}项, max_width={max_width:.2f}cm')

    for sub in sub_options:
        label = sub.get('label', '')
        text = sub.get('text', '')
        full = f'{label} {text}'
        sub_w = measure_text_width_cm(full)

        if current_line and current_width + gap_cm + sub_w > max_width:
            # 当前行放不下 → 输出当前行，开新行
            logger.debug(f'    [子选项贪心]   "{label} {text[:20]}..." w={sub_w:.2f} → '
                         f'当前行={current_width:.2f}+{gap_cm}+{sub_w:.2f}={current_width+gap_cm+sub_w:.2f} > {max_width:.2f} ⚡换行')
            p = doc.add_paragraph()
            apply_style(p, '选项')
            add_mixed_text(p, '  '.join(current_line), cn_font='宋体')
            clear_run_fonts(p)
            line_count += 1
            current_line = [full]
            current_width = sub_w
        else:
            if current_line:
                current_width += gap_cm
            current_line.append(full)
            current_width += sub_w
            logger.debug(f'    [子选项贪心]   "{label} {text[:20]}..." w={sub_w:.2f} → '
                         f'当前行={current_width:.2f} ✅同行({len(current_line)}项)')

    # 输出最后一行
    if current_line:
        p = doc.add_paragraph()
        apply_style(p, '选项')
        add_mixed_text(p, '  '.join(current_line), cn_font='宋体')
        clear_run_fonts(p)
        line_count += 1

    if page_tracker and line_count > 0:
        page_tracker.consume_text_lines(line_count)


def _format_rule_1x4(doc, options, image_resolver, logger, page_tracker=None):
    p = doc.add_paragraph()
    apply_style(p, '选项')
    set_tab_stops(p, TAB_STOPS_4)
    letters = sorted(options.keys())
    for i, letter in enumerate(letters):
        if i > 0:
            p.add_run('\t')
        _add_option_label_text(p, letter, options[letter], image_resolver, logger, page_tracker=page_tracker)
    clear_run_fonts(p)
    if page_tracker:
        page_tracker.consume_paragraph(estimated_lines=1)


def _format_rule_2x2(doc, options, image_resolver, logger, page_tracker=None):
    """2x2 布局：三档对齐，B/D 落在 4.54/8.98/13.43 之一。

    根据 A/C 和 B/D 的宽度，选择能让两者都放下的最大目标位置（B/D 尽量靠右最美观）：
      情况2: B/D @ 4.54cm  — A/C 短，B/D 长
      情况3: B/D @ 8.98cm  — A/C 中等，B/D 中等
      情况4: B/D @ 13.43cm — A/C 长，B/D 短
    A/C 在 0cm，B/D 对齐 1x4 的 B/C/D 之一，跨题垂直对齐。
    每行根据左选项宽度动态选择 tab 数量（1/2/3），使右选项精确落在目标位置。
    """
    SAFETY_MARGIN_2X2_CM = 0.05  # 2x2 安全余量较小（无 tab 间距，仅纯文本列宽）
    tab_stops_cm = [4.54, 8.98, 13.43]

    # 测量各选项宽度
    opt_widths = {}
    for label, text in options.items():
        clean = PLACEHOLDER_TOKEN_PATTERN.sub('', f'{label}. {text}')
        clean = FORMAT_TAG_PATTERN.sub('', clean)
        opt_widths[label] = measure_text_width_cm(clean)

    first_col_max = max(opt_widths.get('A', 0), opt_widths.get('C', 0))
    second_col_max = max(opt_widths.get('B', 0), opt_widths.get('D', 0))

    # 选择目标位置（最大可行位置优先，B/D 尽量靠右最美观）
    # 注意：tab_stops_cm 保持正序 [4.54, 8.98, 13.43] 用于 _tab_count 的 index 计算，
    #       此处仅选择循环使用逆序，使 B/D 尽量对齐到 D 列(13.43)
    target = None
    for t in [13.43, 8.98, 4.54]:
        if (first_col_max + SAFETY_MARGIN_2X2_CM <= t and
                second_col_max + SAFETY_MARGIN_2X2_CM <= PAGE_CONTENT_WIDTH_CM - t):
            target = t
            break

    if target is None:
        # 无可行位置，降级 4x1（理论上 _select_option_rule 已过滤）
        if logger:
            logger.warning('  2x2 无可行目标位置，降级 4x1')
        _format_rule_4x1(doc, options, image_resolver, logger, page_tracker)
        return

    target_index = tab_stops_cm.index(target)

    def _tab_count(left_label):
        """根据左选项宽度和目标位置计算 tab 数量，使右选项精确落在目标位置。"""
        if left_label not in opt_widths:
            return 1
        w = opt_widths[left_label]
        # 找到左选项宽度对应的起始制表位索引（第一个 >= w 的制表位）
        start_index = 0
        for i, pos in enumerate(tab_stops_cm):
            if w < pos:
                start_index = i
                break
        else:
            start_index = len(tab_stops_cm)
        # tab 数量 = 目标索引 - 起始索引 + 1
        return max(1, target_index - start_index + 1)

    p1 = doc.add_paragraph()
    apply_style(p1, '选项')
    set_tab_stops(p1, TAB_STOPS_2)
    if 'A' in options:
        _add_option_label_text(p1, 'A', options['A'], image_resolver, logger, page_tracker=page_tracker)
    if 'B' in options:
        p1.add_run('\t' * _tab_count('A'))
        _add_option_label_text(p1, 'B', options['B'], image_resolver, logger, page_tracker=page_tracker)
    clear_run_fonts(p1)

    p2 = doc.add_paragraph()
    apply_style(p2, '选项')
    set_tab_stops(p2, TAB_STOPS_2)
    if 'C' in options:
        _add_option_label_text(p2, 'C', options['C'], image_resolver, logger, page_tracker=page_tracker)
    if 'D' in options:
        p2.add_run('\t' * _tab_count('C'))
        _add_option_label_text(p2, 'D', options['D'], image_resolver, logger, page_tracker=page_tracker)
    clear_run_fonts(p2)
    if page_tracker:
        page_tracker.consume_text_lines(2)


def _format_rule_4x1(doc, options, image_resolver, logger, page_tracker=None):
    for letter in sorted(options.keys()):
        p = doc.add_paragraph()
        apply_style(p, '选项')
        _add_option_label_text(p, letter, options[letter], image_resolver, logger, page_tracker=page_tracker)
        clear_run_fonts(p)
        if page_tracker:
            page_tracker.consume_paragraph(estimated_lines=1)


def _format_rule_image(doc, options, image_resolver, logger, page_tracker=None):
    letters = sorted(options.keys())
    img_width = Cm(6.9)

    for row_start in range(0, len(letters), 2):
        p = doc.add_paragraph()
        apply_style(p, '选项')
        set_tab_stops(p, TAB_STOPS_2)
        for col in range(2):
            idx = row_start + col
            if idx >= len(letters):
                break
            letter = letters[idx]
            if col > 0:
                p.add_run('\t')
            _add_option_label_text(p, letter, options[letter], image_resolver, logger, img_width, page_tracker=page_tracker)
        if page_tracker:
            page_tracker.consume_paragraph(estimated_lines=1)

    logger.debug(f'  图片选项: 2x{len(letters)}, 宽度={img_width.cm:.1f}cm')


# ============================================================================
# 题目排版
# ============================================================================

def _format_question_stem(doc, question, image_resolver, logger, quality, page_tracker=None):
    """排版题干段落，处理填空题空位、括号、图片占位符。"""
    q_num = question.get('number', '?')
    q_type = question.get('question_type', '')
    stem = question.get('stem', '')

    quality['total_questions'] += 1
    if q_type == '选择题':
        quality['choice_questions'] += 1
    else:
        quality['non_choice_questions'] += 1

    logger.info(f'  题{q_num} ({q_type}): {stem[:50]}...')

    # 填空题空位统一
    formatted = FILL_IN_BLANK_PATTERN.sub('______', stem)
    if formatted != stem:
        quality['fill_in_blank_count'] += len(FILL_IN_BLANK_PATTERN.findall(stem))

    # 选择题：题干末尾加/替换为标准答题括号（　　）
    if q_type == '选择题':
        # 先剥离末尾的图片占位符，避免占位符在尾部时 trailing_bracket 误判
        # 场景：stem = "xxx（）{{image:img_009}}"，此时括号不在字符串末尾
        trailing_ph_match = re.search(r'(\s*\{\{image:[^}]+\}\}\s*)+$', formatted)
        trailing_ph_tail = trailing_ph_match.group(0) if trailing_ph_match else ''
        check_str = formatted[:len(formatted) - len(trailing_ph_tail)] if trailing_ph_tail else formatted

        trailing_bracket = re.compile(r'[（(]\s*[A-Da-d]?\s*[）)]\s*$')
        if trailing_bracket.search(check_str):
            replaced = trailing_bracket.sub('（\u3000\u3000）', check_str)
            formatted = replaced + trailing_ph_tail
        else:
            formatted = check_str.rstrip() + '（\u3000\u3000）' + trailing_ph_tail

    # 填空题题干：含 ______ 时使用下划线渲染
    if q_type == '填空题' and '______' in formatted:
        _format_fill_in_blank_paragraph(doc, f'{q_num}.', formatted, logger)
        return

    # 按占位符分段处理（有捕获组时，偶数索引=文本，奇数索引=占位符ID）
    parts = PLACEHOLDER_TOKEN_PATTERN.split(formatted)

    if not any(i % 2 == 1 for i in range(len(parts))):
        p = doc.add_paragraph()
        apply_style(p, 'Normal')
        add_mixed_text(p, f'{q_num}. {formatted}', cn_font='宋体')
        if page_tracker:
            page_tracker.consume_paragraph(estimated_lines=1)
        return

    # 有图片占位符：分段处理
    first_text = True
    for i, part in enumerate(parts):
        if i % 2 == 0:
            # 纯文本
            part = part.strip()
            if part:
                p = doc.add_paragraph()
                apply_style(p, 'Normal')
                prefix = f'{q_num}. ' if first_text else ''
                add_mixed_text(p, f'{prefix}{part}', cn_font='宋体')
                if page_tracker:
                    page_tracker.consume_paragraph(estimated_lines=1)
                first_text = False
        else:
            # 占位符 ID
            ph_id = part
            file_path, exists = image_resolver.resolve(ph_id)
            if exists:
                if page_tracker:
                    add_centered_picture_adaptive(doc, file_path, page_tracker, logger)
                else:
                    add_centered_picture(doc, file_path, logger)
                quality['images_inserted'] += 1
                logger.info(f'    题干图片: {os.path.basename(file_path)}')
            else:
                quality['missing_images'].append(ph_id)
                logger.warning(f'    图片缺失: placeholder_id={ph_id}')


def _format_subquestions(doc, question, image_resolver, logger, quality, page_tracker=None):
    """排版非选择题子问题。"""
    subquestions = question.get('subquestions', [])
    for sq in subquestions:
        label = sq.get('label', '')
        stem = sq.get('stem', '')
        formatted = FILL_IN_BLANK_PATTERN.sub('______', stem)
        if formatted != stem:
            quality['fill_in_blank_count'] += 1

        parts = PLACEHOLDER_TOKEN_PATTERN.split(formatted)

        if not any(i % 2 == 1 for i in range(len(parts))):
            p = doc.add_paragraph()
            apply_style(p, 'Normal')
            if label:
                add_mixed_text(p, f'{label} {formatted}', cn_font='宋体')
            else:
                add_mixed_text(p, formatted, cn_font='宋体')
            if page_tracker:
                page_tracker.consume_paragraph(estimated_lines=1)
            continue

        first_text = True
        for i, part in enumerate(parts):
            if i % 2 == 0:
                part = part.strip()
                if part:
                    p = doc.add_paragraph()
                    apply_style(p, 'Normal')
                    prefix = f'{label} ' if first_text and label else ''
                    add_mixed_text(p, f'{prefix}{part}', cn_font='宋体')
                    if page_tracker:
                        page_tracker.consume_paragraph(estimated_lines=1)
                    first_text = False
            else:
                ph_id = part
                file_path, exists = image_resolver.resolve(ph_id)
                if exists:
                    if page_tracker:
                        add_centered_picture_adaptive(doc, file_path, page_tracker, logger)
                    else:
                        add_centered_picture(doc, file_path, logger)
                    quality['images_inserted'] += 1
                    logger.info(f'    子问题图片: {os.path.basename(file_path)}')
                else:
                    quality['missing_images'].append(ph_id)
                    logger.warning(f'    子问题图片缺失: placeholder_id={ph_id}')


# ============================================================================
# 材料排版
# ============================================================================

def _format_choice_material_with_images(doc, text_parts, text_tokens, guide_sentence, image_resolver, logger, quality, page_tracker=None):
    """选择题材料含图片时的特殊处理：图片放在引导语之后。
    
    正确排版顺序：材料正文 + 引导语 → 图片
    而非：材料正文 → 图片 → 引导语
    
    注意：PLACEHOLDER_TOKEN_PATTERN 有捕获组，split 后奇数索引是占位符 ID，
    偶数索引才是纯文本。
    """
    all_text = ''.join(text_parts[i] for i in range(0, len(text_parts), 2)).strip()
    
    p = doc.add_paragraph()
    apply_style(p, 'Body Text')
    add_mixed_text(p, all_text, cn_font='楷体')
    
    if guide_sentence:
        add_mixed_text(p, guide_sentence, cn_font='宋体')
    
    if page_tracker:
        tracking_text = all_text + (guide_sentence or '')
        page_tracker.current_offset += page_tracker.estimate_text_height_cm(tracking_text)
    
    for k, token in enumerate(text_tokens):
        ph_id = token.replace('{{image:', '').replace('}}', '')
        file_path, exists = image_resolver.resolve(ph_id)
        if exists:
            if page_tracker:
                add_centered_picture_adaptive(doc, file_path, page_tracker, logger)
            else:
                add_centered_picture(doc, file_path, logger)
            quality['images_inserted'] += 1
            logger.info(f'    材料图片: {os.path.basename(file_path)}')
        else:
            quality['missing_images'].append(ph_id)
            logger.warning(f'    图片缺失: placeholder_id={ph_id}')


def _is_image_only_materials(question):
    """判断题目的 materials 是否仅含图片（无文本/表格/引导语/标题）。

    用于选择题渲染顺序决策：当 materials 仅含图片时，该图片是题干配图
    （如含①②③④序号的选项配图），应放在题干之后而非之前。
    """
    materials = question.get('materials', [])
    if not materials:
        return False
    for material in materials:
        # 含正文文本 → 不是纯图片
        if (material.get('content') or '').strip():
            return False
        # 含引导语 → 不是纯图片
        if (material.get('guide_sentence') or '').strip():
            return False
        # 含标题 → 不是纯图片
        if (material.get('title') or '').strip():
            return False
        segments = material.get('segments', [])
        if not segments:
            return False
        # 任一段非 image 类型 → 不是纯图片
        for seg in segments:
            if seg.get('type', '') != 'image':
                return False
    return True


def _is_image_only_material(material):
    """判断单个 material 是否仅含图片（无文本/表格/引导语/标题）。
    
    返回 True 表示该 material 是纯图片材料（题干配图），应在题干之后、选项之前渲染。
    """
    # 含正文文本 → 不是纯图片
    if (material.get('content') or '').strip():
        return False
    # 含引导语 → 不是纯图片
    if (material.get('guide_sentence') or '').strip():
        return False
    # 含标题 → 不是纯图片
    if (material.get('title') or '').strip():
        return False
    segments = material.get('segments', [])
    if not segments:
        return False
    # 任一段非 image 类型 → 不是纯图片
    for seg in segments:
        if seg.get('type', '') != 'image':
            return False
    return True


def _split_materials(question):
    """将题目的 materials 分离为 text_materials 和 image_materials。
    
    - text_materials: 包含文本/表格的普通材料，渲染在题干之前
    - image_materials: 纯图片材料（题干配图），渲染在题干之后、选项之前
    """
    materials = question.get('materials', [])
    text_materials = []
    image_materials = []
    for mat in materials:
        if _is_image_only_material(mat):
            image_materials.append(mat)
        else:
            text_materials.append(mat)
    return text_materials, image_materials


def _format_materials(doc, question, image_resolver, logger, quality, page_tracker=None):
    """排版题目材料（含 segments: text/image/table）。"""
    materials = question.get('materials', [])
    if not materials:
        return

    q_type = question.get('question_type', '')

    for material in materials:
        _format_materials_inner(doc, material, image_resolver, logger, quality, q_type, page_tracker)


def _format_one_material(doc, material, image_resolver, logger, quality, q_type='', page_tracker=None):
    """排版单个 material 对象（用于 order 排序后的逐一渲染）。"""
    _format_materials_inner(doc, material, image_resolver, logger, quality, q_type, page_tracker)


def _format_materials_inner(doc, material, image_resolver, logger, quality, q_type='', page_tracker=None):
    """单个 material 的核心渲染逻辑。"""
    content = material.get('content', '')
    guide_sentence = material.get('guide_sentence', '')
    segments = list(material.get('segments', []))
    title = material.get('title', '')
    title_style = material.get('title_style', '')

    # Block 标题：独立段落，黑体（非标准标题如【研学背景】）
    if title and title_style == 'block':
        p_title = doc.add_paragraph()
        apply_style(p_title, 'Body Text')
        add_mixed_text(p_title, title, cn_font='黑体')

    # Inline 标题：同行排版，【材料一】黑体 + 正文楷体
    inline_title = None
    if title and title_style == 'inline':
        inline_title = f'【{title}】'

    if not segments and content:
        if inline_title:
            # Inline 标题 + content：同行渲染
            text_parts = PLACEHOLDER_TOKEN_PATTERN.split(content)
            has_placeholders = any(i % 2 == 1 for i in range(len(text_parts)))
            if has_placeholders:
                p = None
                inline_title_used = False
                for k, tp in enumerate(text_parts):
                    if k % 2 == 0:
                        tp = tp.strip()
                        if tp:
                            if not inline_title_used:
                                p = add_multi_paragraph_material(doc, tp, cn_font='楷体', first_line_prefix=inline_title, page_tracker=page_tracker)
                                inline_title_used = True
                            else:
                                p = add_multi_paragraph_material(doc, tp, cn_font='楷体', page_tracker=page_tracker)
                    else:
                        ph_id = tp
                        file_path, exists = image_resolver.resolve(ph_id)
                        if exists:
                            if page_tracker:
                                add_centered_picture_adaptive(doc, file_path, page_tracker, logger)
                            else:
                                add_centered_picture(doc, file_path, logger)
                            quality['images_inserted'] += 1
                        else:
                            quality['missing_images'].append(ph_id)
                            logger.warning(f'    图片缺失: placeholder_id={ph_id}')
                if guide_sentence:
                    if p is not None:
                        add_mixed_text(p, guide_sentence, cn_font='楷体')
                    else:
                        p = doc.add_paragraph()
                        apply_style(p, 'Body Text')
                        add_mixed_text(p, guide_sentence, cn_font='楷体')
            else:
                p = add_multi_paragraph_material(doc, content, cn_font='楷体', first_line_prefix=inline_title, page_tracker=page_tracker)
                if guide_sentence:
                    if p is not None:
                        add_mixed_text(p, guide_sentence, cn_font='宋体')
                    else:
                        p = doc.add_paragraph()
                        apply_style(p, 'Body Text')
                        add_mixed_text(p, guide_sentence, cn_font='宋体')
            return

        # 无 inline 标题：原有逻辑
        text_parts = PLACEHOLDER_TOKEN_PATTERN.split(content)
        has_placeholders = any(i % 2 == 1 for i in range(len(text_parts)))
        if has_placeholders:
            if q_type == '选择题':
                text_tokens = PLACEHOLDER_TOKEN_PATTERN.findall(content)
                _format_choice_material_with_images(doc, text_parts, text_tokens, guide_sentence, image_resolver, logger, quality, page_tracker)
            else:
                p = None
                for k, tp in enumerate(text_parts):
                    if k % 2 == 0:
                        tp = tp.strip()
                        if tp:
                            p = add_multi_paragraph_material(doc, tp, cn_font='楷体', page_tracker=page_tracker)
                    else:
                        ph_id = tp
                        file_path, exists = image_resolver.resolve(ph_id)
                        if exists:
                            if page_tracker:
                                add_centered_picture_adaptive(doc, file_path, page_tracker, logger)
                            else:
                                add_centered_picture(doc, file_path, logger)
                            quality['images_inserted'] += 1
                        else:
                            quality['missing_images'].append(ph_id)
                            logger.warning(f'    图片缺失: placeholder_id={ph_id}')
                if p is not None and guide_sentence:
                    add_mixed_text(p, guide_sentence, cn_font='楷体')
                elif guide_sentence:
                    p = doc.add_paragraph()
                    apply_style(p, 'Body Text')
                    add_mixed_text(p, guide_sentence, cn_font='楷体')
        else:
            p = add_multi_paragraph_material(doc, content, cn_font='楷体', page_tracker=page_tracker)
            if guide_sentence:
                if p is not None:
                    add_mixed_text(p, guide_sentence, cn_font='宋体')
                else:
                    p = doc.add_paragraph()
                    apply_style(p, 'Body Text')
                    add_mixed_text(p, guide_sentence, cn_font='宋体')
        return

    if not segments:
        return

    # 当 segments 无 text 段但 content 有内容时，插入虚拟 text segment
    if content:
        has_text_segment = any(seg.get('type') == 'text' for seg in segments)
        if not has_text_segment:
            segments = [{'type': 'text', 'content': content}] + list(segments)

    if q_type == '选择题':
        _format_choice_material_segments(doc, segments, guide_sentence, image_resolver, logger, quality, inline_title=inline_title, page_tracker=page_tracker)
    else:
        _format_non_choice_material_segments(doc, segments, guide_sentence, image_resolver, logger, quality, inline_title=inline_title, page_tracker=page_tracker)


def _format_choice_material_segments(doc, segments, guide_sentence, image_resolver, logger, quality, inline_title=None, page_tracker=None):
    """选择题材料含 segments 时的处理。

    排版规则（修复注释位置错误）：
    - **表格/图片之前的 text 段** 视为材料正文，合并为一个段落渲染（保持"文本在前"的选择题版式习惯）
    - **表格/图片之后的 text 段** 视为图表注释，紧跟在该表格/图片之后渲染（按 \\n 拆分多行）
    - 引导语附在材料正文段落末尾
    - 文本段中的内嵌图片占位符按出现顺序延后渲染
    """
    # 找到第一个 image/table 段的位置，作为"前置文本"和"后置内容"的分界点
    first_media_idx = None
    for i, seg in enumerate(segments):
        if seg.get('type', '') in ('image', 'table'):
            first_media_idx = i
            break

    # ---- 收集前置文本段（材料正文 + 内嵌图片占位符） ----
    pre_text_parts = []
    image_phs_from_pre_text = []

    if first_media_idx is None:
        # 没有 image/table 段：所有 text 段都是前置文本
        pre_segments = segments
        post_segments = []
    else:
        pre_segments = segments[:first_media_idx]
        post_segments = segments[first_media_idx:]

    for seg in pre_segments:
        if seg.get('type', '') != 'text':
            continue
        text = seg.get('content', '')
        if not text:
            continue
        text_parts = PLACEHOLDER_TOKEN_PATTERN.split(text)
        for j, tp in enumerate(text_parts):
            if j % 2 == 0:
                tp_stripped = tp.strip()
                if tp_stripped:
                    pre_text_parts.append(tp_stripped)
            else:
                image_phs_from_pre_text.append(tp)

    # ---- 渲染前置文本 + 引导语（同一段落） ----
    full_text = ''.join(pre_text_parts)
    if inline_title or full_text or guide_sentence:
        p = doc.add_paragraph()
        apply_style(p, 'Body Text')
        if inline_title:
            add_mixed_text(p, inline_title, cn_font='黑体')
        if full_text:
            add_mixed_text(p, full_text, cn_font='楷体')
        if guide_sentence:
            add_mixed_text(p, guide_sentence, cn_font='宋体')
        if page_tracker:
            tracking_text = (inline_title or '') + (full_text or '') + (guide_sentence or '')
            page_tracker.current_offset += page_tracker.estimate_text_height_cm(tracking_text)

    # ---- 渲染前置文本中的内嵌图片占位符 ----
    for ph_id in image_phs_from_pre_text:
        file_path, exists = image_resolver.resolve(ph_id)
        if exists:
            if page_tracker:
                add_centered_picture_adaptive(doc, file_path, page_tracker, logger)
            else:
                add_centered_picture(doc, file_path, logger)
            quality['images_inserted'] += 1
            logger.info(f'    材料图片: {os.path.basename(file_path)}')
        else:
            quality['missing_images'].append(ph_id)
            logger.warning(f'    图片缺失: placeholder_id={ph_id}')

    # ---- 按原始顺序渲染 image/table 段 + 紧邻其后的 text 注释段 ----
    for seg in post_segments:
        seg_type = seg.get('type', '')

        if seg_type == 'image':
            img_name = seg.get('name', '')
            img_path = None
            if img_name:
                resolved = False
                for ph_id, path in image_resolver._ph_to_file.items():
                    if os.path.basename(path) == img_name:
                        img_path = path
                        resolved = True
                        break
                if not resolved:
                    for key, val in image_resolver._ph_to_file.items():
                        img_path = os.path.join(os.path.dirname(val), img_name)
                        if os.path.exists(img_path):
                            break
            if img_path and os.path.exists(img_path):
                if page_tracker:
                    add_centered_picture_adaptive(doc, img_path, page_tracker, logger)
                else:
                    add_centered_picture(doc, img_path, logger)
                quality['images_inserted'] += 1
            elif img_name:
                logger.warning(f'    图片未找到: {img_name}')

        elif seg_type == 'table':
            table_data = seg.get('data', seg.get('table_data'))
            if table_data:
                add_table(doc, table_data, logger, image_resolver)
                quality['tables_inserted'] += 1

        elif seg_type == 'text':
            # 后置文本段（一般是表格/图片下方的注释）
            # 按 \n 拆分多行渲染；保留每个 ① ② 序号独占一行
            # 同时处理文本中可能存在的内嵌图片占位符
            text = seg.get('content', '')
            if not text:
                continue
            text_parts = PLACEHOLDER_TOKEN_PATTERN.split(text)
            for j, tp in enumerate(text_parts):
                if j % 2 == 0:
                    # 纯文本部分：按 \n 拆分为多段，每段独立成段
                    if tp.strip():
                        add_multi_paragraph_material(doc, tp, cn_font='楷体', page_tracker=page_tracker)
                else:
                    # 内嵌图片占位符
                    ph_id = tp
                    file_path, exists = image_resolver.resolve(ph_id)
                    if exists:
                        if page_tracker:
                            add_centered_picture_adaptive(doc, file_path, page_tracker, logger)
                        else:
                            add_centered_picture(doc, file_path, logger)
                        quality['images_inserted'] += 1
                        logger.info(f'    材料图片: {os.path.basename(file_path)}')
                    else:
                        quality['missing_images'].append(ph_id)
                        logger.warning(f'    图片缺失: placeholder_id={ph_id}')


def _format_non_choice_material_segments(doc, segments, guide_sentence, image_resolver, logger, quality, inline_title=None, page_tracker=None):
    """非选择题材料含 segments 时的处理：按 segments 顺序渲染，引导语附在末尾。"""
    last_text_para = None

    for seg in segments:
        seg_type = seg.get('type', '')

        if seg_type == 'text':
            text = seg.get('content', '')
            if text or inline_title:
                text_parts = PLACEHOLDER_TOKEN_PATTERN.split(text) if text else []
                has_placeholders = any(i % 2 == 1 for i in range(len(text_parts)))
                if has_placeholders:
                    p = None
                    inline_title_used = False
                    for j, tp in enumerate(text_parts):
                        if j % 2 == 0:
                            if tp.strip():
                                if not inline_title_used and inline_title:
                                    p = add_multi_paragraph_material(doc, tp.strip(), cn_font='楷体', first_line_prefix=inline_title, page_tracker=page_tracker)
                                    inline_title_used = True
                                else:
                                    p = add_multi_paragraph_material(doc, tp.strip(), cn_font='楷体', page_tracker=page_tracker)
                        else:
                            ph_id = tp
                            file_path, exists = image_resolver.resolve(ph_id)
                            if exists:
                                if page_tracker:
                                    add_centered_picture_adaptive(doc, file_path, page_tracker, logger)
                                else:
                                    add_centered_picture(doc, file_path, logger)
                                quality['images_inserted'] += 1
                    last_text_para = p
                else:
                    p = add_multi_paragraph_material(doc, text, cn_font='楷体', first_line_prefix=inline_title, page_tracker=page_tracker)
                    if inline_title:
                        inline_title = None  # 只渲染一次
                    last_text_para = p

        elif seg_type == 'image':
            img_name = seg.get('name', '')
            img_path = None
            if img_name:
                resolved = False
                for ph_id, path in image_resolver._ph_to_file.items():
                    if os.path.basename(path) == img_name:
                        img_path = path
                        resolved = True
                        break
                if not resolved:
                    for key, val in image_resolver._ph_to_file.items():
                        img_path = os.path.join(os.path.dirname(val), img_name)
                        if os.path.exists(img_path):
                            break
            if img_path and os.path.exists(img_path):
                if page_tracker:
                    add_centered_picture_adaptive(doc, img_path, page_tracker, logger)
                else:
                    add_centered_picture(doc, img_path, logger)
                quality['images_inserted'] += 1
            elif img_name:
                logger.warning(f'    图片未找到: {img_name}')

        elif seg_type == 'table':
            table_data = seg.get('data', seg.get('table_data'))
            if table_data:
                add_table(doc, table_data, logger, image_resolver)
                quality['tables_inserted'] += 1

    # 引导语：附加到最后一个文本段落（同段），否则新建段落
    if guide_sentence:
        if last_text_para is not None:
            add_mixed_text(last_text_para, guide_sentence, cn_font='楷体')
        else:
            p = doc.add_paragraph()
            apply_style(p, 'Body Text')
            add_mixed_text(p, guide_sentence, cn_font='楷体')


def _has_essay_content(stem):
    """判断 mixed 类型的 stem 是否包含填空白之外的论述内容。

    如果 stem 含 ______ 填空标记，同时又有论述关键词（请分析、为什么等）
    或问号，则说明除了填空还需要考生展开论述，需要额外留白。
    如果仅有填空标记而无任何论述提示，则为纯填空题，不需要留白。

    Returns:
        bool: True 表示有论述内容（需要留白），False 表示纯填空（不需要留白）
    """
    has_blanks = FILL_IN_BLANK_PATTERN.search(stem)

    # 检查含论述关键词
    essay_keywords = re.compile(
        r'请分析|请说明|试分析|试说明|简述|概括|阐述|说明原因|分析原因|'
        r'请回答|请描述|试述|评价|论证|探讨|为什么|如何|怎样|怎么办|'
        r'有哪些|是什么原因|请归纳|请总结|试归纳|试概括'
    )
    if essay_keywords.search(stem):
        return True

    # 检查是否有问号（疑问句）—— 纯填空通常以句号结尾
    if '？' in stem or '?' in stem:
        return True

    # 有填空但无任何论述标识 → 纯填空题
    if has_blanks:
        return False

    # 无填空 → 纯论述（essay 类型走不到这里，但作为兜底）
    return True


def _add_blank_writing_lines(doc, count=3, line_spacing=1.5, page_tracker=None):
    """为考生添加书写空白行。

    每行一个空段落，1.5 倍行距，使用 Normal 样式。

    Args:
        doc: Word Document 对象
        count: 空白行数（默认 3 行）
        line_spacing: 行距倍数（默认 1.5）
        page_tracker: 页面空间追踪器（可选）
    """
    for _ in range(count):
        p = doc.add_paragraph()
        apply_style(p, 'Normal')
        p.paragraph_format.line_spacing = line_spacing
    if page_tracker:
        page_tracker.consume_text_lines(count * line_spacing)


def _is_sub_item_label(label):
    """检查 label 是否为纯圈码数字（①②③…），属于子条目而非独立子问题。"""
    return label in _CIRCLED_NUMBERS


def _format_one_subquestion(doc, subq, image_resolver, logger, quality, page_tracker=None,
                            is_last_in_group=True):
    """排版单个子问题（含子问题专属 materials 和填空类型渲染）。

    Args:
        is_last_in_group: 是否为其所在问题组的最后一个条目。
            子条目（①②③）只在是组末时才添加书写空白行；
            独立子问题（如 (1)、(2)）按原有逻辑处理。
    """
    label = subq.get('label', '')
    stem = subq.get('stem', '')
    sub_question_type = subq.get('sub_question_type', 'essay')
    subq_materials = subq.get('materials', [])

    formatted = FILL_IN_BLANK_PATTERN.sub('______', stem)
    if formatted != stem:
        quality['fill_in_blank_count'] += 1

    # 填空式子问题：在 ______ 处添加下划线
    is_fill = sub_question_type in ('fill_in_blank', 'mixed')

    parts = PLACEHOLDER_TOKEN_PATTERN.split(formatted)
    has_placeholders = any(i % 2 == 1 for i in range(len(parts)))

    if not has_placeholders and not is_fill:
        # 无图片、无填空：简单段落
        p = doc.add_paragraph()
        apply_style(p, 'Normal')
        if label:
            add_mixed_text(p, f'{label} {formatted}', cn_font='宋体')
        else:
            add_mixed_text(p, formatted, cn_font='宋体')
        if page_tracker:
            page_tracker.consume_paragraph(estimated_lines=1)
    elif not has_placeholders and is_fill:
        # 有填空无图片：将 ______ 渲染为下划线
        _format_fill_in_blank_paragraph(doc, label, formatted, logger)
        if page_tracker:
            page_tracker.consume_paragraph(estimated_lines=1)
    else:
        # 有图片占位符：分段处理
        first_text = True
        for i, part in enumerate(parts):
            if i % 2 == 0:
                part = part.strip()
                if part:
                    if is_fill and '______' in part:
                        _format_fill_in_blank_paragraph(doc, label if first_text else '', part, logger)
                    else:
                        p = doc.add_paragraph()
                        apply_style(p, 'Normal')
                        prefix = f'{label} ' if first_text and label else ''
                        add_mixed_text(p, f'{prefix}{part}', cn_font='宋体')
                    if page_tracker:
                        page_tracker.consume_paragraph(estimated_lines=1)
                    first_text = False
            else:
                ph_id = part
                file_path, exists = image_resolver.resolve(ph_id)
                if exists:
                    if page_tracker:
                        add_centered_picture_adaptive(doc, file_path, page_tracker, logger)
                    else:
                        add_centered_picture(doc, file_path, logger)
                    quality['images_inserted'] += 1
                    logger.info(f'    子问题图片: {os.path.basename(file_path)}')
                else:
                    quality['missing_images'].append(ph_id)
                    logger.warning(f'    子问题图片缺失: placeholder_id={ph_id}')

    # 子问题专属材料（如访谈调查表）
    if subq_materials:
        for mat in subq_materials:
            _format_materials_inner(doc, mat, image_resolver, logger, quality, page_tracker=page_tracker)

    # 主观题后留白
    # 规则：非组末条目（后面还有子条目 ①②③ 的）不添加留白，留白统一由组末控制
    is_sub_item = _is_sub_item_label(label)

    if not is_last_in_group:
        # 此条目后面紧跟着子条目（①②③等），留白由组末条目决定
        logger.debug(f'    子问题 {label}: 非组末，不添加书写空白')
    elif is_sub_item:
        # 子条目（组末）：按类型处理
        if sub_question_type == 'essay':
            _add_blank_writing_lines(doc, count=3, line_spacing=1.5, page_tracker=page_tracker)
            logger.debug(f'    子条目 {label}: 组末论述题，添加3行书写空白')
        elif sub_question_type == 'mixed' and _has_essay_content(stem):
            _add_blank_writing_lines(doc, count=3, line_spacing=1.5, page_tracker=page_tracker)
            logger.debug(f'    子条目 {label}: 组末混合题型(含论述)，添加3行书写空白')
        else:
            logger.debug(f'    子条目 {label}: 不添加书写空白')
    elif sub_question_type == 'essay':
        _add_blank_writing_lines(doc, count=3, line_spacing=1.5, page_tracker=page_tracker)
        logger.debug(f'    子问题 {label}: 论述题，添加3行书写空白')
    elif sub_question_type == 'mixed':
        if _has_essay_content(stem):
            _add_blank_writing_lines(doc, count=3, line_spacing=1.5, page_tracker=page_tracker)
            logger.debug(f'    子问题 {label}: 混合题型(含论述)，添加3行书写空白')
        else:
            logger.debug(f'    子问题 {label}: 纯填空题，不添加书写空白')


def _format_fill_in_blank_paragraph(doc, label, text, logger):
    """将含 ______ 的文本渲染为带下划线填空区域的段落。"""
    blank_pattern = re.compile(r'(_{3,})')
    parts = blank_pattern.split(text)

    p = doc.add_paragraph()
    apply_style(p, 'Normal')
    # 填空题增加行间距
    p.paragraph_format.line_spacing = 1.8

    if label:
        add_mixed_text(p, f'{label} ', cn_font='宋体')

    for part in parts:
        if blank_pattern.match(part):
            # 填空区域：用空格+下划线渲染
            blank_len = len(part)
            # 使用全角空格模拟填空宽度，每个 ______ (6个_) 渲染为约 4 个全角空格
            space_count = max(4, blank_len // 2)
            run = p.add_run('\u3000' * space_count)
            run.underline = True
            run.font.size = Pt(14)
            _set_run_font(run, cn_font='宋体', en_font='Times New Roman')
        else:
            if part:
                add_mixed_text(p, part, cn_font='宋体')


# ============================================================================
# 分区排版
# ============================================================================

def _get_subq_group_info(subquestions):
    """分析子问题分组，返回 [(subq, is_last_in_group), ...] 列表。

    分组规则：
    - 一个独立子问题（label 如 "(1)"）与其后续的子条目（label 如 "①②③"）属于同一组
    - 只有每组的最后一个条目标记 is_last_in_group=True
    - 若独立子问题后没有子条目，则其自身即为组末
    """
    if not subquestions:
        return []
    result = []
    i = 0
    n = len(subquestions)
    while i < n:
        start = i
        i += 1
        # 收集后续子条目（圈码数字 label）
        while i < n and _is_sub_item_label(subquestions[i].get('label', '')):
            i += 1
        for j in range(start, i):
            result.append((subquestions[j], j == i - 1))
    return result


def _format_non_choice_body(doc, question, image_resolver, logger, quality, page_tracker=None):
    """排版非选择题主体：若 materials/subquestions 含 order 字段则按 order 交错渲染，否则默认顺序。"""
    materials = question.get('materials', [])
    subquestions = question.get('subquestions', [])

    # 检测是否有 order 字段
    has_material_order = any(m.get('order') is not None for m in materials)
    has_subq_order = any(sq.get('order') is not None for sq in subquestions)

    if not has_material_order and not has_subq_order:
        # 默认顺序：所有材料 → 所有子问题（按分组控制留白）
        _format_materials(doc, question, image_resolver, logger, quality, page_tracker)
        for sq, is_last in _get_subq_group_info(subquestions):
            _format_one_subquestion(doc, sq, image_resolver, logger, quality, page_tracker,
                                    is_last_in_group=is_last)
        return

    # 有 order 字段：构建排序列表
    ordered_items = []
    for m in materials:
        order = m.get('order')
        if order is not None:
            ordered_items.append((order, 'material', m))
    for sq in subquestions:
        order = sq.get('order')
        if order is not None:
            ordered_items.append((order, 'subquestion', sq))

    ordered_items.sort(key=lambda x: x[0])

    # 验证是否有遗漏（有 order 的部分按 order 排，无 order 的放在最后）
    no_order_materials = [m for m in materials if m.get('order') is None]
    no_order_subqs = [sq for sq in subquestions if sq.get('order') is None]

    # 为有序子问题构建分组信息（忽略材料项，只对子问题分组）
    ordered_subqs = [sq for _, item_type, sq in ordered_items if item_type == 'subquestion']
    _sq_to_is_last = {id(sq): is_last for sq, is_last in _get_subq_group_info(ordered_subqs)}

    # 按 order 顺序渲染
    for _, item_type, item in ordered_items:
        if item_type == 'material':
            _format_one_material(doc, item, image_resolver, logger, quality, page_tracker=page_tracker)
        elif item_type == 'subquestion':
            is_last = _sq_to_is_last.get(id(item), True)
            _format_one_subquestion(doc, item, image_resolver, logger, quality, page_tracker,
                                    is_last_in_group=is_last)

    # 渲染无 order 的剩余项
    for m in no_order_materials:
        _format_one_material(doc, m, image_resolver, logger, quality, page_tracker=page_tracker)
    for sq, is_last in _get_subq_group_info(no_order_subqs):
        _format_one_subquestion(doc, sq, image_resolver, logger, quality, page_tracker,
                                is_last_in_group=is_last)


def _format_section(doc, section, image_resolver, logger, quality):
    """排版一个分区（选择题/非选择题等）。"""
    sec_type = section.get('type', '')
    sec_id = section.get('id', '')
    sec_title = section.get('title', '')
    header_blocks = section.get('header_blocks', [])
    instructions = section.get('instructions', [])
    questions = section.get('questions', [])

    # 创建页面空间追踪器（每个分区独立追踪）
    page_tracker = PageSpaceTracker()

    logger.info(f'')
    logger.info(f'--- 排版 {sec_id}: {sec_title[:40] if sec_title else sec_type} ---')

    # 分区头部内容块（按原文顺序渲染：卷标题 / 注意事项 / 说明文字等）
    for block in header_blocks:
        block_type = block.get('type', '')
        content = block.get('content', '')
        if not content:
            continue

        if block_type == 'notes':
            # 注意事项：拆分标题行（黑体）+ 内容行（楷体）
            _format_notes_block(doc, content, logger)
            page_tracker.consume_text_lines(len(content.split('\n')))
        elif block_type == 'volume_title':
            # 卷标题：使用"第几卷"样式，不强制覆盖字体
            p = doc.add_paragraph()
            apply_style(p, '第几卷')
            add_mixed_text(p, content)
            page_tracker.consume_paragraph(estimated_lines=1)
            logger.debug(f'  header_block (volume_title): {content[:40]}')
        else:
            # instructions / 其他说明文字：宋体
            p = doc.add_paragraph()
            apply_style(p, 'Body Text')
            add_mixed_text(p, content, cn_font='宋体')
            page_tracker.consume_paragraph(estimated_lines=1)
            logger.debug(f'  header_block ({block_type}): {content[:40]}')

    # 分区题型标题
    if sec_title:
        p = doc.add_paragraph()
        apply_style(p, '题型标题')
        p.add_run(sec_title)
        page_tracker.consume_paragraph(estimated_lines=1)

    # 分区说明（跳过与标题或 header_blocks 重复的内容）
    header_contents = {b.get('content', '') for b in header_blocks}
    for instr in instructions:
        instr_stripped = instr.strip()
        if sec_title and instr_stripped in sec_title:
            logger.debug(f'  跳过重复说明: {instr_stripped[:30]}')
            continue
        if any(instr_stripped in hc for hc in header_contents):
            logger.debug(f'  跳过重复说明（已在header_blocks中）: {instr_stripped[:30]}')
            continue
        p = doc.add_paragraph()
        apply_style(p, 'Body Text')
        add_mixed_text(p, instr, cn_font='宋体')
        page_tracker.consume_paragraph(estimated_lines=1)

    quality['sections'] += 1

    for question in questions:
        q_type = question.get('question_type', sec_type)

        if q_type == '非选择题' or q_type == '综合题':
            # 非选择题/综合题：题干 → 按 order 交错渲染 materials/subquestions（或默认顺序）
            _format_question_stem(doc, question, image_resolver, logger, quality, page_tracker)
            _format_non_choice_body(doc, question, image_resolver, logger, quality, page_tracker)
        elif q_type == '填空题':
            # 填空题：材料(如有) → 题干（含下划线渲染），无选项
            _format_materials(doc, question, image_resolver, logger, quality, page_tracker)
            _format_question_stem(doc, question, image_resolver, logger, quality, page_tracker)
        else:
            # 选择题：材料(如有) → 题干 → 选项
            # 选择题：材料分离为文本材料和纯图片材料
            # 文本材料 → 题干前；纯图片材料 → 题干后、选项前（题干配图）
            text_mats, image_mats = _split_materials(question)
            # 先渲染文本材料
            if text_mats:
                orig_mats = question.get('materials', [])
                question['materials'] = text_mats
                _format_materials(doc, question, image_resolver, logger, quality, page_tracker)
                question['materials'] = orig_mats
            # 渲染题干
            _format_question_stem(doc, question, image_resolver, logger, quality, page_tracker)
            # 再渲染纯图片材料（题干配图）
            if image_mats:
                orig_mats = question.get('materials', [])
                question['materials'] = image_mats
                _format_materials(doc, question, image_resolver, logger, quality, page_tracker)
                question['materials'] = orig_mats

            if q_type == '选择题':
                options = question.get('options', [])
                sub_options = question.get('sub_options', [])
                if options:
                    rule = format_options(doc, options, image_resolver, logger, page_tracker, sub_options)
                    if rule:
                        quality['option_rules'][f'规则{rule}'] = \
                            quality['option_rules'].get(f'规则{rule}', 0) + 1


# ============================================================================
# 考前内容排版
# ============================================================================

def _format_meta_notes(doc, meta, logger):
    """排版 meta.notes（注意事项等）。优先使用 notes_items 结构化数据，回退到 notes 纯文本。"""
    notes_items = meta.get('notes_items', [])
    
    if notes_items:
        for item in notes_items:
            item_type = item.get('type', '')
            content = item.get('content', '')
            number = item.get('number', '')
            
            if not content:
                continue
            
            p = doc.add_paragraph()
            if item_type == 'title':
                apply_style(p, '题型标题')
                add_mixed_text(p, content, cn_font='黑体', en_font='Times New Roman')
                logger.debug(f'  注意事项标题: {content[:40]}')
            elif item_type == 'item':
                apply_style(p, '注意事项内容')
                if number:
                    add_mixed_text(p, f'{number}. {content}', cn_font='楷体', en_font='Times New Roman')
                else:
                    add_mixed_text(p, content, cn_font='楷体', en_font='Times New Roman')
                logger.debug(f'  注意事项内容: {content[:40]}')
        return
    
    notes = meta.get('notes', '')
    if not notes:
        return

    lines = notes.strip().split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('注意事项') or line.startswith('注意：'):
            p = doc.add_paragraph()
            apply_style(p, '题型标题')
            add_mixed_text(p, line, cn_font='黑体', en_font='Times New Roman')
            logger.debug(f'  注意事项标题: {line[:40]}')
        else:
            p = doc.add_paragraph()
            apply_style(p, '注意事项内容')
            add_mixed_text(p, line, cn_font='楷体', en_font='Times New Roman')
            logger.debug(f'  注意事项内容: {line[:40]}')


def _format_notes_block(doc, content, logger):
    """渲染 header_blocks 中的注意事项：标题行→黑体，内容行→楷体。
    
    支持两种输入格式：
    1. 多行（有 \\n）：按行拆分，首行含"注意事项"→黑体，其余→楷体（与 _format_meta_notes 一致）
    2. 单行（无 \\n）：提取"注意事项："/"注意："前缀→黑体，剩余按序号（1. 2. 3. 等）拆分→楷体
    """
    text = content.strip()
    if not text:
        return

    # 提取标题前缀（如有）："注意事项：" / "注意事项" / "注意："
    heading = ''
    body = text
    for prefix in ('注意事项：', '注意事项', '注意：'):
        if text.startswith(prefix):
            heading = prefix
            body = text[len(prefix):].strip()
            break

    # 渲染标题
    if heading:
        p = doc.add_paragraph()
        apply_style(p, '题型标题')
        add_mixed_text(p, heading, cn_font='黑体', en_font='Times New Roman')
        logger.debug(f'  注意事项标题: {heading}')

    if not body:
        return

    # 拆分正文：优先按 \\n 拆分，若无则按序号（1. 2. 3. 或 1、2、3、等）拆分
    if '\n' in body:
        items = [l.strip() for l in body.split('\n') if l.strip()]
    else:
        items = [m.strip() for m in re.split(r'(?<!\d)(?=\d+[\.\、\)）])', body) if m.strip()]
        if not items:
            items = [body]

    for item in items:
        p = doc.add_paragraph()
        apply_style(p, '注意事项内容')
        add_mixed_text(p, item, cn_font='楷体', en_font='Times New Roman')
        logger.debug(f'  注意事项内容: {item[:40]}')


# ============================================================================
# 考试信息排版
# ============================================================================

SYMBOL_TOKEN_RE = re.compile(r'\{\{symbol:[^}]+\}\}')
IMAGE_TOKEN_RE = re.compile(r'\{\{image:[^}]+\}\}')


def _strip_placeholder_tokens(text):
    """移除文本中的占位符 token（{{symbol:xxx}} / {{image:xxx}}），清理首尾空白。"""
    text = SYMBOL_TOKEN_RE.sub('', text)
    text = IMAGE_TOKEN_RE.sub('', text)
    return text.strip()


def _format_exam_header(doc, meta, image_resolver, logger, quality):
    """排版考试名称和科目。"""
    title = meta.get('title', '')
    subtitle = meta.get('subtitle', '')
    subject = meta.get('subject', '')

    # 清理 title 中的占位符
    title = _strip_placeholder_tokens(title)

    if title:
        # 版式一不进行 AI 断句，渲染完整考试名称
        # python-docx 1.2.0 会把真实 \n 转为 <w:br/>，必须清理
        title_clean = title.replace('\n', '').replace('\r', '').replace('\\n', '')
        p = doc.add_paragraph()
        apply_style(p, '考试名称')
        add_mixed_text(p, title_clean)
        logger.info(f'  考试名称: {title_clean[:50]}')
    else:
        logger.warning('  meta.title 为空或仅含占位符 token，已跳过')

    # subtitle：清理占位符，仅在有实质内容时渲染
    subtitle = _strip_placeholder_tokens(subtitle)
    if subtitle:
        p = doc.add_paragraph()
        apply_style(p, '科目名称')
        add_mixed_text(p, subtitle)
        logger.info(f'  副标题: {subtitle}')

    # subject 不再作为独立行渲染。subtitle 已承担科目名称的显示职责，
    # subject 字段是 tag_structure 的硬编码填充，渲染会导致重复。
    if subject:
        logger.info(f'  科目(仅记录，不渲染): {subject}')


# ============================================================================
# 质检报告
# ============================================================================

def _generate_report(quality, output_path, logger, report_dir=None):
    if report_dir:
        os.makedirs(report_dir, exist_ok=True)
        report_path = os.path.join(report_dir, 'quality_report.html')
    else:
        report_path = os.path.join(os.path.dirname(output_path), 'quality_report.html')
    has_issues = bool(quality['missing_images'] or quality['warnings'])
    status_class = 'success' if not has_issues else 'warning'
    status_text = '质检通过' if not has_issues else '发现问题'
    issue_count = len(quality['missing_images']) + len(quality['warnings'])
    status_desc = '所有检查项通过。' if not has_issues else f'发现 {issue_count} 个问题。'

    rule_descriptions = {
        '规则1': '短选项 1x4（每槽 4.54/4.44/4.45/3.77cm）',
        '规则2': '中等选项 2x2（左列≤8.98cm, 右列≤3.97cm，含比例字体容差）',
        '规则3': '长选项 4x1',
        '规则5': '图片选项 2x2',
    }

    option_rules_html = ''
    for rule, count in sorted(quality['option_rules'].items()):
        desc = rule_descriptions.get(rule, '')
        option_rules_html += f'<tr><td>{rule}</td><td>{count} 题</td><td>{desc}</td></tr>\n'
    if not quality['option_rules']:
        option_rules_html += '<tr><td colspan="3" style="color:#999;">无</td></tr>\n'

    missing_html = ''
    for img in quality['missing_images']:
        missing_html += f'<div class="warn"><div class="warn-title">图片缺失</div><div class="warn-detail">{img}</div></div>\n'

    warn_html = ''
    for w in quality['warnings']:
        warn_html += f'<div class="warn"><div class="warn-title">警告</div><div class="warn-detail">{w}</div></div>\n'

    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>排版质检报告 - {quality['exam_name']}</title>
<style>
*{{margin:0;padding:0;box-sizing:border-box}}
body{{font-family:"Microsoft YaHei",sans-serif;background:#f0f2f5;color:#333;line-height:1.7;padding:20px}}
.container{{max-width:920px;margin:0 auto}}
.header{{background:linear-gradient(135deg,#1a5276,#2980b9);color:#fff;padding:30px 40px;border-radius:12px 12px 0 0}}
.header h1{{font-size:24px;margin-bottom:6px}}
.badge{{display:inline-block;padding:4px 14px;border-radius:20px;font-size:13px;font-weight:600;margin-top:10px}}
.badge-ok{{background:#27ae60}}
.badge-warn{{background:#e67e22}}
.card{{background:#fff;border-radius:0 0 12px 12px;padding:30px 40px;box-shadow:0 2px 8px rgba(0,0,0,.08);margin-bottom:20px}}
.card h2{{font-size:17px;color:#1a5276;border-left:4px solid #2980b9;padding-left:10px;margin:0 0 14px 0}}
.grid{{display:grid;grid-template-columns:repeat(3,1fr);gap:12px}}
.grid-item{{background:#f8f9fa;border-radius:8px;padding:14px 16px;text-align:center}}
.grid-item .lbl{{font-size:12px;color:#7f8c8d}}
.grid-item .val{{font-size:22px;font-weight:700;color:#2c3e50}}
table{{width:100%;border-collapse:collapse;font-size:14px}}
th{{background:#eaf2f8;color:#1a5276;font-weight:600;padding:10px 14px;text-align:left;border-bottom:2px solid #d5dbdb}}
td{{padding:9px 14px;border-bottom:1px solid #ecf0f1}}
.tag{{display:inline-block;padding:2px 10px;border-radius:4px;font-size:12px;font-weight:600}}
.tag-ok{{background:#e8f8e8;color:#27ae60}}
.warn{{background:#fff3cd;border:1px solid #ffeaa7;border-left:4px solid #f39c12;border-radius:8px;padding:14px 18px;margin-bottom:10px}}
.warn-title{{font-weight:700;color:#d68910;font-size:14px;margin-bottom:4px}}
.warn-detail{{font-size:13px;color:#7d6608}}
.warn.ok{{background:#e8f8e8;border-color:#abebc6;border-left-color:#27ae60}}
.warn.ok .warn-title{{color:#1e8449}}
.warn.ok .warn-detail{{color:#196f3d}}
.footer{{text-align:center;font-size:12px;color:#bdc3c7;padding:20px 0}}
</style>
</head>
<body>
<div class="container">
<div class="header">
<h1>排版质检报告</h1>
<div>{quality['exam_name']}</div>
<div class="badge {'badge-ok' if status_class == 'success' else 'badge-warn'}">{status_text}</div>
</div>
<div class="card">
<h2>结构统计</h2>
<div class="grid">
<div class="grid-item"><div class="lbl">分区</div><div class="val">{quality['sections']}</div></div>
<div class="grid-item"><div class="lbl">总题数</div><div class="val">{quality['total_questions']}</div></div>
<div class="grid-item"><div class="lbl">选择题</div><div class="val">{quality['choice_questions']}</div></div>
<div class="grid-item"><div class="lbl">非选择题</div><div class="val">{quality['non_choice_questions']}</div></div>
<div class="grid-item"><div class="lbl">图片</div><div class="val">{quality['images_inserted']}</div></div>
<div class="grid-item"><div class="lbl">表格</div><div class="val">{quality['tables_inserted']}</div></div>
</div>
</div>
<div class="card-section">
<h2>选项排版规则</h2>
<table><thead><tr><th>规则</th><th>次数</th><th>说明</th></tr></thead><tbody>{option_rules_html}</tbody></table>
</div>
<div class="card-section">
<h2>资源统计</h2>
<table>
<tr><td>插入图片</td><td><span class="tag tag-ok">OK</span></td><td>{quality['images_inserted']} 张</td></tr>
<tr><td>插入表格</td><td><span class="tag tag-ok">OK</span></td><td>{quality['tables_inserted']} 个</td></tr>
<tr><td>填空空位</td><td><span class="tag tag-ok">OK</span></td><td>{quality['fill_in_blank_count']} 处</td></tr>
</table>
</div>
{missing_html}
{warn_html}
<div class="card-section">
<h2>结论</h2>
<div class="warn {'ok' if status_class == 'success' else ''}"><div class="warn-title">{status_text}</div><div class="warn-detail">{status_desc}</div></div>
</div>
</div>
<div class="footer">地理试卷排版 v3.0 | {datetime.now().strftime('%Y-%m-%d %H:%M')}</div>
</div>
</body>
</html>'''

    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(html)
    logger.info(f'质检报告: {report_path}')


# ============================================================================
# 选择题答题表
# ============================================================================

# 答题表布局参数
ANSWER_TABLE_CELL_WIDTH_CM = 1.5          # 单元格宽度（用于计算最大物理列数）
ANSWER_TABLE_Q_ROW_HEIGHT_CM = 0.7        # 题号行高度
ANSWER_TABLE_A_ROW_HEIGHT_CM = 1.0        # 答案行高度（较高方便书写）
ANSWER_TABLE_GAP_ROW_HEIGHT_CM = 0.3      # 块间间隔行高度

# 引导语中题组范围匹配模式
_GROUP_RANGE_RE = re.compile(r'(\d+)\s*[～~\-—–]+\s*(\d+)\s*题')
# 计数模式：需"下面/下"前缀或"小题"后缀，避免"完成14题"误匹配为14道
_GROUP_COUNT_RE = re.compile(r'(?:下面|下)\s*(\d+)\s*小?\s*题|(\d+)\s*小题')


def _calc_target_lines(total_questions):
    """根据选择题总数计算目标行数。

    规则：
      - < 20 题 → 2 行
      - 20~29 题 → 3 行
      - 30~39 题 → 4 行
      - 以此类推（每 10 题增加一行）

    Args:
        total_questions: 选择题总数

    Returns:
        目标行数（int）
    """
    if total_questions <= 0:
        return 0
    # 1~19 → 2行, 20~29 → 3行, 30~39 → 4行, 以此类推
    return max(2, (total_questions + 10) // 10)


def _auto_calc_target_per_line(groups, target_lines, page_width):
    """计算 target_per_line 和 max_physical_cols，使贪心分行行数 ≤ target_lines。

    先尝试默认单元格宽度（1.5cm），若物理列数限制导致无法达标，
    则逐步缩小单元格宽度（最低 1.0cm）来增大物理列数容量。

    Args:
        groups: 题组列表
        target_lines: 目标行数
        page_width: 可用页面宽度（cm）

    Returns:
        (target_per_line, max_physical_cols) 元组
    """
    total_q = sum(len(g) for g in groups)
    lo_q = max(max(len(g) for g in groups), 1) if groups else 1

    # 候选单元格宽度：从大到小尝试（越大越美观）
    candidate_widths = [1.5, 1.4, 1.3, 1.2, 1.1, 1.0]

    for cell_width in candidate_widths:
        max_cols = int(page_width / cell_width)

        # 先检查：无题数限制时能否达标
        lines = _greedy_layout_lines(groups, target_per_line=total_q,
                                     max_physical_cols=max_cols)
        if len(lines) > target_lines:
            continue  # 物理列数仍不够，尝试更小的单元格宽度

        # 达标，二分搜索最小的 target_per_line
        lo, hi = lo_q, total_q
        result = hi
        while lo <= hi:
            mid = (lo + hi) // 2
            lines = _greedy_layout_lines(groups, target_per_line=mid,
                                         max_physical_cols=max_cols)
            if len(lines) <= target_lines:
                result = mid
                hi = mid - 1
            else:
                lo = mid + 1

        return result, max_cols

    # 所有宽度都试过仍无法达标，使用最窄的
    max_cols = int(page_width / 1.0)
    return total_q, max_cols


def _parse_group_size(guide_sentence):
    """从材料引导语解析题组包含的题目数量。

    匹配模式：
      - "据此完成1～3题" → 3
      - "完成下面3小题" → 3
      - "完成1-3题"     → 3
      - "完成下列要求"   → None（无法确定）

    Args:
        guide_sentence: 材料末尾的引导语句

    Returns:
        题组大小（int），无法确定时返回 None
    """
    if not guide_sentence:
        return None

    m = _GROUP_RANGE_RE.search(guide_sentence)
    if m:
        start = int(m.group(1))
        end = int(m.group(2))
        if end >= start and 1 <= (end - start + 1) <= 20:
            return end - start + 1

    m = _GROUP_COUNT_RE.search(guide_sentence)
    if m:
        count = int(m.group(1) or m.group(2))
        if 2 <= count <= 20:  # "1小题"不作为题组信号，需≥2才有意义
            return count

    return None


def _identify_choice_groups(choice_questions):
    """识别选择题题组。

    规则：
      1. 有 materials 的题是新题组起始（材料题）
      2. 优先从 guide_sentence 解析题组大小，明确收集对应数量的小题
      3. 无法解析大小时，题组延续到下一个有 materials 的题
      4. 无 materials 的独立题各为一组

    Args:
        choice_questions: 选择题列表（已按题号顺序排列）

    Returns:
        题组列表，如 [[1, 2], [3, 4, 5], [6, 7, 8]]
    """
    groups = []
    i = 0
    n = len(choice_questions)

    while i < n:
        q = choice_questions[i]
        has_materials = bool(q.get('materials'))
        number = q.get('number', '')

        if has_materials:
            # 提取 guide_sentence
            guide = ''
            for mat in q.get('materials', []):
                g = mat.get('guide_sentence', '')
                if g:
                    guide = g
                    break

            group_size = _parse_group_size(guide)

            if group_size and group_size >= 1:
                # 明确大小：收集 group_size 个题
                group = [number]
                for j in range(1, group_size):
                    if i + j < n:
                        group.append(choice_questions[i + j].get('number', ''))
                groups.append(group)
                i += group_size
            else:
                # 无法确定大小：延续到下一个有 materials 的题
                group = [number]
                i += 1
                while i < n and not bool(choice_questions[i].get('materials')):
                    group.append(choice_questions[i].get('number', ''))
                    i += 1
                groups.append(group)
        else:
            # 无材料 → 独立成组
            groups.append([number])
            i += 1

    return groups


def _greedy_layout_lines(choice_groups, target_per_line,
                         max_physical_cols=None):
    """贪心分行算法。

    以题组为排版单位，同一题组不拆分。当加入当前行会超过
    target_per_line 或 max_physical_cols 时换行。

    题组超过一行容量时，单独成行（不拆分）。

    Args:
        choice_groups: 题组列表，如 [[1,2], [3,4,5], [6,7,8]]
        target_per_line: 每行目标题数
        max_physical_cols: 最大物理列数（None 时自动计算）

    Returns:
        分行结果，每个元素是一行中的题组列表
        如 [[[1,2], [3,4,5], [6,7,8]], [[9,10], [11,12,13], [14,15]]]
    """
    if max_physical_cols is None:
        max_physical_cols = int(PAGE_CONTENT_WIDTH_CM / ANSWER_TABLE_CELL_WIDTH_CM)

    lines = []
    current_line = []
    current_q_count = 0
    current_p_count = 0

    for group in choice_groups:
        group_size = len(group)

        if current_line:
            new_p_count = current_p_count + 1 + group_size
            new_q_count = current_q_count + group_size
            if new_q_count <= target_per_line and new_p_count <= max_physical_cols:
                current_line.append(group)
                current_q_count = new_q_count
                current_p_count = new_p_count
            else:
                lines.append(current_line)
                current_line = [group]
                current_q_count = group_size
                current_p_count = group_size
        else:
            current_line.append(group)
            current_q_count = group_size
            current_p_count = group_size

    if current_line:
        lines.append(current_line)

    return lines


def _set_answer_cell_border(cell, has_border=True):
    """设置答题表单元格边框。

    Args:
        cell: python-docx Cell 对象
        has_border: True 绘制四边框，False 无边框
    """
    tc = cell._tc
    tcPr = tc.find(docx_qn('w:tcPr'))
    if tcPr is None:
        tcPr = etree.SubElement(tc, docx_qn('w:tcPr'))

    for existing in tcPr.findall(docx_qn('w:tcBorders')):
        tcPr.remove(existing)

    tcBorders = etree.SubElement(tcPr, docx_qn('w:tcBorders'))
    sides = ('top', 'left', 'bottom', 'right')
    if has_border:
        for side in sides:
            b = etree.SubElement(tcBorders, docx_qn(f'w:{side}'))
            b.set(docx_qn('w:val'), 'single')
            b.set(docx_qn('w:sz'), '4')
            b.set(docx_qn('w:space'), '0')
            b.set(docx_qn('w:color'), '000000')
    else:
        for side in sides:
            b = etree.SubElement(tcBorders, docx_qn(f'w:{side}'))
            b.set(docx_qn('w:val'), 'nil')


def _set_answer_cell_vertical_center(cell):
    """设置单元格内容垂直居中。"""
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(docx_qn('w:vAlign')):
        tcPr.remove(existing)
    vAlign = etree.SubElement(tcPr, docx_qn('w:vAlign'))
    vAlign.set(docx_qn('w:val'), 'center')


def _build_choice_answer_table(doc, layout_lines, logger):
    """根据分行结果构建选择题答题表。

    表格结构（每行题组生成 2 行 + 块间 1 行间隔）：
      题号行 | 答案行 | [间隔行] | 题号行 | 答案行 | ...

    单元格类型：
      - question：有边框，题号行填题号/答案行留空
      - gap：无边框，题组间隔
      - padding：无边框，补齐列数

    Args:
        doc: Word Document 对象
        layout_lines: 分行结果
        logger: 日志记录器

    Returns:
        python-docx Table 对象
    """
    if not layout_lines:
        return None

    # 计算最大物理列数
    max_physical_cols = 0
    for line in layout_lines:
        cols = sum(len(g) for g in line) + max(0, len(line) - 1)
        max_physical_cols = max(max_physical_cols, cols)
    if max_physical_cols == 0:
        return None

    total_rows = len(layout_lines) * 2 + max(0, len(layout_lines) - 1)

    table = doc.add_table(rows=total_rows, cols=max_physical_cols)
    try:
        table.style = 'Table Normal'
    except KeyError:
        pass
    table.autofit = False
    table.allow_autofit = False

    cell_width = PAGE_CONTENT_WIDTH_CM / max_physical_cols
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(cell_width)

    current_row = 0
    for line_idx, line in enumerate(layout_lines):
        # 构建单元格序列：题组之间插入 gap
        cells_seq = []
        for g_idx, group in enumerate(line):
            if g_idx > 0:
                cells_seq.append(('gap', ''))
            for num in group:
                cells_seq.append(('question', str(num)))
        while len(cells_seq) < max_physical_cols:
            cells_seq.append(('padding', ''))

        # 题号行
        q_row = table.rows[current_row]
        q_row.height = Cm(ANSWER_TABLE_Q_ROW_HEIGHT_CM)
        for j, (cell_type, text) in enumerate(cells_seq):
            cell = q_row.cells[j]
            _set_answer_cell_border(cell, has_border=(cell_type == 'question'))
            _set_answer_cell_vertical_center(cell)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.0
            if text:
                run = p.add_run(text)
                _set_run_font(run, cn_font='宋体', en_font='Times New Roman', size_pt=12)

        # 答案行
        a_row = table.rows[current_row + 1]
        a_row.height = Cm(ANSWER_TABLE_A_ROW_HEIGHT_CM)
        for j, (cell_type, text) in enumerate(cells_seq):
            cell = a_row.cells[j]
            _set_answer_cell_border(cell, has_border=(cell_type == 'question'))
            _set_answer_cell_vertical_center(cell)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.0

        current_row += 2

        # 块间间隔行
        if line_idx < len(layout_lines) - 1:
            gap_row = table.rows[current_row]
            gap_row.height = Cm(ANSWER_TABLE_GAP_ROW_HEIGHT_CM)
            for j in range(max_physical_cols):
                cell = gap_row.cells[j]
                _set_answer_cell_border(cell, has_border=False)
                cell.text = ''
            current_row += 1

    logger.info(f'  答题表: {total_rows}行 x {max_physical_cols}列, {len(layout_lines)}个答题块')
    return table


def add_choice_answer_table(doc, sections, logger):
    """在文档末尾添加选择题答题表。

    答题表插入位置：主观题最后一道题空行之后（即文档当前末尾）。
    无选择题时不生成。

    Args:
        doc: Word Document 对象
        sections: 试卷分区列表
        logger: 日志记录器

    Returns:
        Table 对象，无选择题时返回 None
    """
    choice_questions = []
    for section in sections:
        if section.get('type') == '选择题':
            for q in section.get('questions', []):
                if q.get('question_type') == '选择题':
                    choice_questions.append(q)

    if not choice_questions:
        logger.info('无选择题，跳过答题表生成')
        return None

    groups = _identify_choice_groups(choice_questions)
    logger.info(f'选择题题组: {groups}')

    total_q = sum(len(g) for g in groups)
    target_lines = _calc_target_lines(total_q)
    target_per_line, max_cols = _auto_calc_target_per_line(groups, target_lines, PAGE_CONTENT_WIDTH_CM)
    logger.info(f'选择题 {total_q}题, 目标 {target_lines}行, 每行上限 {target_per_line}题, {max_cols}列')

    layout_lines = _greedy_layout_lines(groups, target_per_line=target_per_line, max_physical_cols=max_cols)
    logger.info(f'答题表分行: {layout_lines}')

    # 答题表标题
    p_title = doc.add_paragraph()
    apply_style(p_title, '题型标题')
    add_mixed_text(p_title, '选择题答题卡', cn_font='黑体')

    table = _build_choice_answer_table(doc, layout_lines, logger)
    return table


# ============================================================================
# 版式二封面生成
# ============================================================================

def _generate_cover(doc, meta, sections, logger):
    """生成版式二的独立封面。

    封面包含：
      - 练习日期标签
      - 考试名称（动态）
      - 科目名称（动态）
      - 选择题答题卡（动态）
      - 姓名/班级/分数标签
    """
    title = _strip_placeholder_tokens(meta.get('title', ''))
    subtitle = _strip_placeholder_tokens(meta.get('subtitle', ''))
    subject = meta.get('subject', '')

    # --- 练习日期 ---
    practice_table = doc.add_table(rows=1, cols=2)
    try:
        practice_table.style = 'Table Normal'
    except KeyError:
        pass
    practice_table.autofit = True
    # 设置练习日期单元格
    date_cell = practice_table.rows[0].cells[0]
    date_cell.text = ''
    p = date_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('练习日期：')
    _set_run_font(run, cn_font='宋体', en_font='Times New Roman', size_pt=14)
    # 右侧空白单元格供书写日期
    blank_cell = practice_table.rows[0].cells[1]
    blank_cell.text = ''
    # 设置练习日期表格宽度适应
    date_cell.width = Cm(3.0)
    blank_cell.width = Cm(10.0)

    # --- 空行留白（参考封面模板的上部 spacing） ---
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.keep_with_next = True

    # --- 考试名称（动态生成） ---
    # subtitle 与 subject 通常相同（都是科目名称），避免重复渲染：
    # subtitle → 在"考试名称"行中以宋体显示
    # subject  → 在"科目名称"行中以黑体显示
    # 若两者相同，则 subtitle 从考试名称行中排除，只保留 subject 的黑体渲染
    # 此外，title 末尾可能已包含科目名称（如"浙江省选考地理"），
    # 需去除末尾的 subject 后缀以避免科目名在考试名称行和科目名称行重复显示
    exam_title_lines = []
    if title:
        # 支持 title 中的 \n 作为建议性换行标记（由 tag_structure 的 AI 判断插入）
        for line in title.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue
            # 若该行以 subject 结尾，去除科目后缀，避免与科目名称行重复
            if subject and stripped.endswith(subject):
                stripped = stripped[:-len(subject)].strip()
            if stripped:
                exam_title_lines.append(stripped)
    if subtitle and subtitle.strip() != subject.strip():
        exam_title_lines.append(subtitle)

    for line_text in exam_title_lines:
        p = doc.add_paragraph()
        apply_style(p, '考试名称')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        add_mixed_text(p, line_text, cn_font='宋体', size_pt=28)
        logger.info(f'  封面考试名: {line_text[:50]}')

    # --- 科目名称 ---
    if subject:
        p = doc.add_paragraph()
        apply_style(p, '科目名称')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        add_mixed_text(p, subject, cn_font='黑体', size_pt=48)
        logger.info(f'  封面科目: {subject}')

    # --- 考前信息行 ---
    # 空行
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.keep_with_next = True

    # --- 选择题答题卡（封面专用） ---
    choice_questions = []
    for section in sections:
        if section.get('type') == '选择题':
            for q in section.get('questions', []):
                if q.get('question_type') == '选择题':
                    choice_questions.append(q)

    if choice_questions:
        # 答题卡标题
        p_title = doc.add_paragraph()
        apply_style(p_title, '题型标题')
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.keep_with_next = True
        add_mixed_text(p_title, '选择题答题卡', cn_font='黑体')

        groups = _identify_choice_groups(choice_questions)

        total_q = sum(len(g) for g in groups)
        target_lines = _calc_target_lines(total_q)
        target_per_line, max_cols = _auto_calc_target_per_line(groups, target_lines, PAGE_CONTENT_WIDTH_CM)

        layout_lines = _greedy_layout_lines(groups, target_per_line=target_per_line, max_physical_cols=max_cols)
        logger.info(f'  封面答题表: {len(choice_questions)}题, {len(layout_lines)}行')

        table = _build_choice_answer_table(doc, layout_lines, logger)
        if table is not None:
            # 防止表格行跨页拆分，确保封面内容在一页内
            for row in table.rows:
                tr = row._tr
                trPr = tr.get_or_add_trPr()
                cant_split = OxmlElement('w:cantSplit')
                trPr.append(cant_split)

    # --- 答题卡与考生信息之间的空行 ---
    p_gap = doc.add_paragraph()
    p_gap.paragraph_format.space_before = Pt(0)
    p_gap.paragraph_format.space_after = Pt(0)
    p_gap.paragraph_format.line_spacing = 1.15
    p_gap.paragraph_format.keep_with_next = True

    # --- 考生信息行：姓名 / 班级 / 分数 ---
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.paragraph_format.space_before = Pt(0)
    p_info.paragraph_format.line_spacing = 1.5

    # "姓名：" + 留白 + "班级：" + 留白 + "分数："
    run_name = p_info.add_run('姓名：__________')
    _set_run_font(run_name, cn_font='宋体', en_font='Times New Roman', size_pt=14)

    run_spacer1 = p_info.add_run('        ')
    _set_run_font(run_spacer1, cn_font='宋体', size_pt=14)

    run_class = p_info.add_run('班级：__________')
    _set_run_font(run_class, cn_font='宋体', en_font='Times New Roman', size_pt=14)

    run_spacer2 = p_info.add_run('        ')
    _set_run_font(run_spacer2, cn_font='宋体', size_pt=14)

    run_score = p_info.add_run('分数：__________')
    _set_run_font(run_score, cn_font='宋体', en_font='Times New Roman', size_pt=14)

    # --- 封面与正文之间的分页 ---
    doc.add_page_break()
    logger.info('  封面→正文分页已添加')


# ============================================================================
# 主排版流程
# ============================================================================

def typeset_exam(json_path, template_path, output_path, images_dir=None, log_path=None, report_dir=None, format='v1'):
    """主排版入口函数。

    Args:
        json_path:   final_exam.json (v3.0 Schema)
        template_path: template.dotx
        output_path:  输出 docx 路径
        images_dir:   图片目录
        log_path:     日志路径
        report_dir:   质检报告输出目录（默认与 output_path 同目录）
        format:       版式类型 'v1' (默认) 或 'v2'

    Returns:
        bool: 排版是否成功
    """
    json_path = os.path.abspath(json_path)
    template_path = os.path.abspath(template_path)
    output_path = os.path.abspath(output_path)

    # 输入校验
    errors = []
    if not os.path.exists(json_path):
        errors.append(f'JSON 文件不存在: {json_path}')
    if not os.path.exists(template_path):
        errors.append(f'模板文件不存在: {template_path}')
    if errors:
        for e in errors:
            print(f'错误: {e}', file=sys.stderr)
        return False

    if images_dir is None:
        images_dir = os.path.join(os.path.dirname(json_path), '..', '清洗产物', 'images')
    images_dir = os.path.abspath(images_dir)

    if log_path is None:
        log_path = os.path.join(os.path.dirname(output_path), 'typeset_log.txt')

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    logger = setup_logger(log_path)

    logger.info('=' * 60)
    logger.info('地理试卷排版 v3.0 (typeset_exam)')
    logger.info(f'JSON:   {json_path}')
    logger.info(f'模板:   {template_path}')
    logger.info(f'图片:   {images_dir}')
    logger.info(f'输出:   {output_path}')
    logger.info('=' * 60)

    try:
        # 1. 加载 JSON
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        meta = data.get('meta', {})
        document = data.get('document', {})
        images_data = data.get('images', [])
        image_mapping = data.get('image_mapping', [])
        validation = data.get('validation', {})
        sections = document.get('sections', [])

        # 2. 构建图片映射解析器
        image_resolver = ImageResolver(images_data, image_mapping, images_dir)

        # 3. 初始质检数据
        quality = {
            'exam_name': _strip_placeholder_tokens(meta.get('title', '')),
            'sections': 0,
            'total_questions': 0,
            'choice_questions': 0,
            'non_choice_questions': 0,
            'images_inserted': 0,
            'tables_inserted': 0,
            'option_rules': {},
            'missing_images': [],
            'fill_in_blank_count': 0,
            'warnings': [],
        }

        # 注入校验警告
        if validation.get('has_unmapped_placeholders'):
            quality['warnings'].append('存在未映射的占位符：' + ', '.join(
                validation.get('unmapped_placeholders', [])))
        if validation.get('has_unused_images'):
            quality['warnings'].append('存在未使用的图片：' + ', '.join(
                validation.get('unused_images', [])))
        for w in validation.get('warnings', []):
            if w not in quality['warnings']:
                quality['warnings'].append(w)

        # 4. 加载模板
        logger.info('加载模板...')
        doc = load_template(template_path)
        logger.info('模板加载成功')

        # v2 格式：生成封面
        if format == 'v2':
            logger.info('')
            logger.info('--- 生成封面（版式二） ---')
            _generate_cover(doc, meta, sections, logger)
            logger.info('封面生成完成')

        # 5. 排版：考试信息（v1 显示，v2 封面已含此内容故跳过）
        if format != 'v2':
            logger.info('')
            logger.info('--- 考试信息 ---')
            _format_exam_header(doc, meta, image_resolver, logger, quality)

        # 6. 排版：注意事项
        _format_meta_notes(doc, meta, logger)

        # 7. 排版：各分区
        for section in sections:
            _format_section(doc, section, image_resolver, logger, quality)

        # 8. 未归类块（如有）
        unclassified = document.get('unclassified_blocks', [])
        if unclassified:
            logger.warning(f'存在 {len(unclassified)} 个未归类块，渲染为灰色警告标注')
            quality['warnings'].append(f'存在 {len(unclassified)} 个未归类文本块，已以灰色标注渲染')
            # 添加分隔标记
            p_sep = doc.add_paragraph()
            apply_style(p_sep, '题型标题')
            add_mixed_text(p_sep, '【未归类内容】', cn_font='黑体')
            for block in unclassified:
                text = block.get('text', '')
                reason = block.get('reason', '未知原因')
                if text:
                    p = doc.add_paragraph()
                    apply_style(p, 'Body Text')
                    run_label = p.add_run(f'[{reason}] ')
                    run_label.font.color.rgb = RGBColor(180, 180, 180)
                    run_label.font.size = Pt(10)
                    add_mixed_text(p, text, cn_font='宋体')

        # 8.5. 选择题答题表（v2 封面已含，v1 插入到主观题最后一道题空行之后）
        if format != 'v2':
            logger.info('')
            logger.info('--- 选择题答题表 ---')
            add_choice_answer_table(doc, sections, logger)

        # 9. 添加页脚页码
        logger.info('')
        logger.info('--- 添加页脚页码 ---')
        _add_page_number_footer(doc)
        logger.info('页脚页码已添加')

        # 10. 保存
        logger.info('')
        logger.info('--- 保存文档 ---')
        doc.save(output_path)
        logger.info(f'文档已保存: {output_path}')

        # 11. 质检报告
        _generate_report(quality, output_path, logger, report_dir)

        # 打印摘要
        logger.info('')
        logger.info('=' * 60)
        logger.info(f'排版完成!')
        logger.info(f'  分区: {quality["sections"]}  题目: {quality["total_questions"]}')
        logger.info(f'  选择题: {quality["choice_questions"]}  非选择题: {quality["non_choice_questions"]}')
        logger.info(f'  图片: {quality["images_inserted"]}  表格: {quality["tables_inserted"]}')
        logger.info(f'  缺失: {len(quality["missing_images"])}  警告: {len(quality["warnings"])}')
        logger.info('=' * 60)

        return True

    except json.JSONDecodeError as e:
        logger.error(f'JSON 解析失败: {e}')
        return False
    except Exception as e:
        logger.error(f'排版失败: {e}', exc_info=True)
        return False


# ============================================================================
# CLI
# ============================================================================

def main():
    parser = argparse.ArgumentParser(
        description='地理试卷排版 v3.0 - 基于 final_exam.json + template.dotx 生成排版文档',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
示例:
  python typeset_exam.py --json {工作目录}/试卷数据/final_exam.json \\
                         --template assets/template.dotx \\
                         --images {工作目录}/清洗产物/images/ \\
                         --output {工作目录}/{试卷名称}-排版后.docx \\
                         --report-dir {工作目录}/排版文档/
        '''
    )
    parser.add_argument('--json', '-j', required=True, help='final_exam.json 路径 (v3.0 Schema)')
    parser.add_argument('--template', '-t', required=True, help='template.dotx 路径')
    parser.add_argument('--output', '-o', required=True, help='输出 docx 路径')
    parser.add_argument('--images', '-i', help='图片目录 (默认 {工作目录}/清洗产物/images/)')
    parser.add_argument('--log', '-l', help='日志文件路径')
    parser.add_argument('--report-dir', '-r', help='质检报告输出目录（默认与 output 同目录）')
    parser.add_argument('--format', '-f', default='v1', choices=['v1', 'v2'],
                        help='版式类型: v1=标准版式, v2=封面版式 (默认 v1)')

    args = parser.parse_args()

    success = typeset_exam(
        json_path=args.json,
        template_path=args.template,
        output_path=args.output,
        images_dir=args.images,
        log_path=args.log,
        report_dir=args.report_dir,
        format=args.format,
    )

    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
